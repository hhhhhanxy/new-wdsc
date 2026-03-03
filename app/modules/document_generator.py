from flask_restful import Resource, reqparse
import os
from docx import Document

class DocumentGenerator(Resource):
    def post(self):
        parser = reqparse.RequestParser()
        parser.add_argument('template_path', type=str, required=True, help='模板路径')
        parser.add_argument('data', type=dict, required=True, help='填充数据')
        parser.add_argument('output_path', type=str, required=True, help='输出路径')
        args = parser.parse_args()
        
        template_path = args['template_path']
        data = args['data']
        output_path = args['output_path']
        
        # 确保路径是绝对路径
        if not os.path.isabs(template_path):
            template_path = os.path.abspath(template_path)
        if not os.path.isabs(output_path):
            output_path = os.path.abspath(output_path)
        
        if not os.path.exists(template_path):
            return {'error': '模板文件不存在'}, 404
        
        # 确保输出目录存在
        output_dir = os.path.dirname(output_path)
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        try:
            self.generate_document(template_path, data, output_path)
            return {'status': 'success', 'output_path': output_path}, 200
        except Exception as e:
            return {'error': str(e)}, 500
    
    def generate_document(self, template_path, data, output_path):
        """基于模板生成文档"""
        doc = Document(template_path)
        
        # 替换文本内容
        for paragraph in doc.paragraphs:
            for key, value in data.items():
                if f'{{{{{key}}}}}' in paragraph.text:
                    paragraph.text = paragraph.text.replace(f'{{{{{key}}}}}', str(value))
        
        # 替换表格内容
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in data.items():
                        if f'{{{{{key}}}}}' in cell.text:
                            cell.text = cell.text.replace(f'{{{{{key}}}}}', str(value))
        
        doc.save(output_path)