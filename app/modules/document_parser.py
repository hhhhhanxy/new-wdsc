from flask_restful import Resource, reqparse
import os
import PyPDF2
from docx import Document

class DocumentParser(Resource):
    def post(self):
        parser = reqparse.RequestParser()
        parser.add_argument('file_path', type=str, required=True, help='文档路径')
        args = parser.parse_args()
        
        file_path = args['file_path']
        # 确保路径是绝对路径
        if not os.path.isabs(file_path):
            file_path = os.path.abspath(file_path)
        
        if not os.path.exists(file_path):
            return {'error': '文件不存在'}, 404
        
        try:
            if file_path.endswith('.pdf'):
                content = self.parse_pdf(file_path)
            elif file_path.endswith('.docx'):
                content = self.parse_docx(file_path)
            else:
                return {'error': '不支持的文件格式'}, 400
            
            return {'status': 'success', 'content': content}, 200
        except Exception as e:
            return {'error': str(e)}, 500
    
    def parse_pdf(self, file_path):
        """解析PDF文档"""
        content = {
            'text': '',
            'tables': [],
            'images': []
        }
        
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page_num in range(len(reader.pages)):
                page = reader.pages[page_num]
                content['text'] += page.extract_text() + '\n'
        
        return content
    
    def parse_docx(self, file_path):
        """解析Word文档"""
        content = {
            'text': '',
            'tables': [],
            'images': []
        }
        
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            content['text'] += paragraph.text + '\n'
        
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            content['tables'].append(table_data)
        
        return content