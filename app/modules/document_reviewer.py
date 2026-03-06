import os
import sys
import datetime

# 添加项目根目录到Python路径
sys.path.append(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))

from flask_restful import Resource, reqparse
from docx import Document
from app.modules.review_framework import rule_manager

class DocumentReviewer(Resource):
    def post(self):
        parser = reqparse.RequestParser()
        parser.add_argument('file_path', type=str, required=True, help='文档路径')
        parser.add_argument('rules', type=dict, required=False, help='审查规则')
        parser.add_argument('save_result', type=bool, required=False, default=False, help='是否保存审查结果文档')
        parser.add_argument('result_path', type=str, required=False, help='审查结果文档保存路径')
        args = parser.parse_args()
        
        file_path = args['file_path']
        rules = args.get('rules', {})
        save_result = args.get('save_result', False)
        result_path = args.get('result_path', None)
        
        # 确保路径是绝对路径
        if not os.path.isabs(file_path):
            file_path = os.path.abspath(file_path)
        
        if not os.path.exists(file_path):
            return {'error': '文件不存在'}, 404
        
        try:
            review_results = self.review_document(file_path, rules)
            
            # 如果需要保存结果文档
            result_document_path = None
            if save_result:
                result_document_path = self.save_review_result(file_path, review_results, result_path)
            
            response = {'status': 'success', 'results': review_results}
            if result_document_path:
                response['result_document'] = result_document_path
            
            return response, 200
        except Exception as e:
            return {'error': str(e)}, 500
    
    def save_review_result(self, original_file_path, review_results, result_path=None):
        """
        保存审查结果为Word文档
        :param original_file_path: 原始文档路径
        :param review_results: 审查结果
        :param result_path: 结果文档保存路径
        :return: 结果文档保存路径
        """
        # 如果未指定保存路径，生成默认路径
        if not result_path:
            original_name = os.path.splitext(os.path.basename(original_file_path))[0]
            timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
            # 默认保存到results文件夹
            project_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
            result_path = os.path.join(project_root, 'results', f'{original_name}_审查结果_{timestamp}.docx')
        
        # 确保保存目录存在
        result_dir = os.path.dirname(result_path)
        if not os.path.exists(result_dir):
            os.makedirs(result_dir)
        
        # 创建Word文档
        doc = Document()
        
        # 添加标题
        doc.add_heading('文档审查结果', 0)
        
        # 添加基本信息
        doc.add_heading('基本信息', level=1)
        doc.add_paragraph(f'审查时间: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph(f'原始文档: {original_file_path}')
        
        # 添加审查结果
        doc.add_heading('审查结果', level=1)
        
        # 格式检查结果
        if review_results.get('format_check', []):
            doc.add_heading('1. 格式规范检查', level=2)
            for issue in review_results['format_check']:
                doc.add_paragraph(f'- {issue}', style='List Bullet')
        else:
            doc.add_heading('1. 格式规范检查', level=2)
            doc.add_paragraph('✓ 无问题', style='List Bullet')
        
        # 内容检查结果
        if review_results.get('content_check', []):
            doc.add_heading('2. 内容完整性检查', level=2)
            for issue in review_results['content_check']:
                doc.add_paragraph(f'- {issue}', style='List Bullet')
        else:
            doc.add_heading('2. 内容完整性检查', level=2)
            doc.add_paragraph('✓ 无问题', style='List Bullet')
        
        # 术语检查结果
        if review_results.get('terminology_check', []):
            doc.add_heading('3. 术语一致性检查', level=2)
            for issue in review_results['terminology_check']:
                doc.add_paragraph(f'- {issue}', style='List Bullet')
        else:
            doc.add_heading('3. 术语一致性检查', level=2)
            doc.add_paragraph('✓ 无问题', style='List Bullet')
        
        # 合规性检查结果
        if review_results.get('compliance_check', []):
            doc.add_heading('4. 合规性验证', level=2)
            for issue in review_results['compliance_check']:
                doc.add_paragraph(f'- {issue}', style='List Bullet')
        else:
            doc.add_heading('4. 合规性验证', level=2)
            doc.add_paragraph('✓ 无问题', style='List Bullet')
        
        # 语言检查结果
        if review_results.get('language_check', []):
            doc.add_heading('5. 语言风格检查', level=2)
            for issue in review_results['language_check']:
                doc.add_paragraph(f'- {issue}', style='List Bullet')
        else:
            doc.add_heading('5. 语言风格检查', level=2)
            doc.add_paragraph('✓ 无问题', style='List Bullet')
        
        # 添加总结
        doc.add_heading('总结', level=1)
        total_issues = sum(len(issues) for issues in review_results.values())
        if total_issues == 0:
            doc.add_paragraph('✓ 文档审查通过，未发现问题。')
        else:
            doc.add_paragraph(f'⚠️ 文档审查发现 {total_issues} 个问题，建议根据上述结果进行修改。')
        
        # 保存文档
        doc.save(result_path)
        
        return result_path
    
    def review_document(self, file_path, rules):
        """审查文档"""
        # 确保rules是字典
        if rules is None:
            rules = {}
        
        if file_path.endswith('.docx'):
            doc = Document(file_path)
            text = '\n'.join([p.text for p in doc.paragraphs])
        else:
            # 对于其他格式，这里简化处理
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
        
        # 使用规则管理器执行所有审查规则
        results = rule_manager.execute_rules(text, rules)
        
        return results