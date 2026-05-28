"""
航空技术文档模板化生成器。

按模板章节结构，逐章调用 LLM 生成内容，组装为 DOCX。
"""
import logging
from typing import Dict, Any, Optional, List
from pathlib import Path

from generators.base_generator import BaseGenerator
from models.document import DocumentType
from templates.template_manager import TemplateManager, ChapterTemplate

logger = logging.getLogger(__name__)


class TemplateDocxGenerator(BaseGenerator):
    """按模板章节结构生成航空技术文档。"""

    name = "template_docx"
    display_name = "模板文档生成器"
    description = "根据标准模板章节结构，逐章生成航空技术文档"
    supported_doc_types = [t.value for t in DocumentType]

    def __init__(self, template_manager: TemplateManager = None):
        self.template_manager = template_manager or TemplateManager()

    def generate(
        self,
        title: str,
        params: Dict[str, Any],
        llm_client=None,
        output_path: Optional[str] = None,
    ) -> str:
        if not llm_client:
            raise ValueError("TemplateDocxGenerator 需要 LLM 客户端")

        doc_type_str = params.get("doc_type", "requirements")
        doc_type = DocumentType(doc_type_str)
        template = self.template_manager.get_template(doc_type)
        if not template:
            raise ValueError(f"未找到文档类型 {doc_type_str} 的模板")

        description = params.get("description", "")
        tech_params = params.get("technical_params", "")

        # 逐章生成内容
        chapters_content: List[Dict[str, str]] = []
        for chapter in template.chapters:
            content = self._generate_chapter(
                llm_client, title, chapter, description, tech_params
            )
            chapters_content.append({
                "number": chapter.number,
                "title": chapter.title,
                "content": content,
            })

        # 组装 DOCX
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        heading = doc.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for ch in chapters_content:
            doc.add_heading(f"{ch['number']} {ch['title']}", level=1)
            for paragraph in ch['content'].split('\n'):
                paragraph = paragraph.strip()
                if paragraph:
                    doc.add_paragraph(paragraph)

        if output_path is None:
            raise ValueError("output_path is required")

        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        return output_path

    def _generate_chapter(
        self,
        llm_client,
        doc_title: str,
        chapter: ChapterTemplate,
        description: str,
        tech_params: str,
    ) -> str:
        """调用 LLM 生成单个章节内容。"""
        sub_chapter_info = ""
        if chapter.sub_chapters:
            sub_titles = [f"  {sc.number} {sc.title}" for sc in chapter.sub_chapters]
            sub_chapter_info = f"\n\n本章应包含以下子章节：\n" + "\n".join(sub_titles)

        guidance = ""
        if chapter.guidance_prompt:
            guidance = f"\n\n生成指引：{chapter.guidance_prompt}"

        prompt = f"""请为航空作动领域的文档「{doc_title}」生成第 {chapter.number} 章「{chapter.title}」的内容。

产品描述：{description}
关键技术参数：{tech_params}
章节说明：{chapter.description}{sub_chapter_info}{guidance}

要求：
1. 内容专业、准确，符合航空领域规范
2. 使用规范工程术语
3. 内容完整、逻辑清晰
4. 仅输出章节正文内容，不要输出章节标题"""

        try:
            response = llm_client.generate(prompt)
            return response.content
        except Exception as e:
            logger.error("生成章节 %s 失败: %s", chapter.number, e)
            return f"[生成失败：{e}]"
