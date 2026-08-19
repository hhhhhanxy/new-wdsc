"""
Base generator abstraction for document generation.
"""
from abc import ABC, abstractmethod
from typing import Dict, Any, Optional, List, Callable
from pathlib import Path
import logging

from docx.document import Document as DocumentType
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from templates.docx_block_classifier import DocxBlockClassifier
from .prompt_builder import ChapterPromptBuilder

logger = logging.getLogger(__name__)


class BaseGenerator(ABC):
    """文档生成器基类，子类必须实现 generate() 方法。"""

    name: str = ""
    display_name: str = ""
    description: str = ""
    supported_doc_types: List[str] = []

    @abstractmethod
    def generate(
        self,
        title: str,
        params: Dict[str, Any],
        llm_client=None,
        output_path: Optional[str] = None,
        progress_callback: Optional[Callable[[int, int, Any], None]] = None,
        control_callback: Optional[Callable[[], None]] = None,
    ) -> str:
        """
        生成文档。

        Args:
            title: 文档标题
            params: 生成参数（description, technical_params, doc_type 等）
            llm_client: 可选的 LLM 客户端
            output_path: 输出文件路径
            progress_callback: 可选进度回调，参数为当前章节序号、总章节数、章节对象
            control_callback: 可选任务控制回调，用于暂停或停止检查

        Returns:
            生成的文件路径
        """
        ...


class TemplateDocxGenerator(BaseGenerator):
    """基于原始 DOCX 模板的受控填充生成器。"""

    name = "template_docx"
    display_name = "DOCX 模板填充生成器"
    description = "保留原始 DOCX 模板正文与格式，删除说明文字并替换占位内容"
    supported_doc_types = [
        "requirements",
        "general_characteristics",
        "technical_specification",
        "verification",
    ]

    def __init__(self):
        self.classifier = DocxBlockClassifier()
        self.prompt_builder = ChapterPromptBuilder()

    def generate(
        self,
        title: str,
        params: Dict[str, Any],
        llm_client=None,
        output_path: Optional[str] = None,
        progress_callback: Optional[Callable[[int, int, Any], None]] = None,
        control_callback: Optional[Callable[[], None]] = None,
    ) -> str:
        if output_path is None:
            raise ValueError("output_path is required")

        from templates.template_manager import TemplateManager

        template_id = params.get("template_id") or params.get("doc_type")
        template = TemplateManager().get_template(template_id)
        if not template:
            raise ValueError(f"模板不存在: {template_id}")
        self.last_template = template

        inputs = params.get("inputs") or {}
        template_path = self._resolve_template_docx_path(template)
        if template_path and template_path.exists():
            return self._fill_template_document(
                template=template,
                title=title,
                inputs=inputs,
                params=params,
                llm_client=llm_client,
                output_path=output_path,
                progress_callback=progress_callback,
                control_callback=control_callback,
            )

        raise ValueError("该模板缺少原始 DOCX 文件，不能用于保留格式的文档生成；请在生成模板库中上传 DOCX 模板。")

    def _flatten_chapters(self, chapters) -> List[Any]:
        flat = []
        for chapter in chapters:
            flat.append(chapter)
            flat.extend(self._flatten_chapters(chapter.sub_chapters))
        return flat

    def _resolve_template_docx_path(self, template) -> Optional[Path]:
        source_path = (template.metadata or {}).get("source_docx_path") or (template.metadata or {}).get("source_path")
        if not source_path:
            return None
        path = Path(source_path)
        if not path.is_absolute():
            path = Path(__file__).resolve().parent.parent / path
        return path

    def _fill_template_document(
        self,
        template,
        title: str,
        inputs: Dict[str, Any],
        params: Dict[str, Any],
        llm_client,
        output_path: str,
        progress_callback: Optional[Callable[[int, int, Any], None]] = None,
        control_callback: Optional[Callable[[], None]] = None,
    ) -> str:
        from docx import Document

        template_path = self._resolve_template_docx_path(template)
        if not template_path:
            raise ValueError("模板缺少原始 DOCX 文件路径")

        doc = Document(str(template_path))
        replacements = self._build_replacements(title, inputs)
        pending_meta_fields = self._collect_pending_meta_fields(inputs)
        self._replace_document_meta_placeholders(doc, replacements)
        self._replace_document_title_candidates(doc, title, replacements)
        chapters = self._flatten_chapters(template.chapters)
        chapter_targets = self._collect_generation_targets(doc, chapters)
        if "generation_mode" not in inputs:
            chapter_targets = {
                key: values[:1]
                for key, values in chapter_targets.items()
            }
        section_details: Dict[str, dict] = {}
        generated_sections = self._generate_sections(
            template=template,
            title=title,
            inputs=inputs,
            chapters=chapters,
            llm_client=llm_client,
            enabled=str(inputs.get("generation_mode") or "smart") == "smart",
            section_details=section_details,
            chapter_targets=chapter_targets,
            control_callback=control_callback,
        )
        params["_generation_meta"] = {
            "mode": str(inputs.get("generation_mode") or "smart"),
            "llm_enabled": bool(llm_client),
            "generated_sections": len(generated_sections),
            "sections": list(section_details.values()),
            "pending_fields": pending_meta_fields,
        }
        for index, chapter in enumerate(chapters, start=1):
            if control_callback:
                control_callback()
            if progress_callback:
                progress_callback(index, len(chapters), chapter)

        delete_context = ""
        current_chapter_key = ""
        applied_generated_counts: Dict[str, int] = {}
        chapter_needs_target = {
            self._chapter_key(chapter): self._chapter_needs_specific_target(chapter)
            for chapter in chapters
        }
        chapter_strategies = {
            self._chapter_key(chapter): self._chapter_strategy(chapter)
            for chapter in chapters
        }
        for block in list(self._iter_block_items(doc)):
            if isinstance(block, Paragraph):
                text = block.text.strip()
                if not current_chapter_key and self._is_template_document_title(text, title):
                    self._replace_paragraph_text(block, title)
                    continue
                if self.classifier.heading_level(block) and not (
                    delete_context and self.classifier.is_context_continuation(text)
                ):
                    if control_callback:
                        control_callback()
                    current_chapter_key = self._match_chapter_key(text, chapters)
                    delete_context = ""
                    self._replace_paragraph_runs(block, replacements)
                    continue
                block_type = self.classifier.classify_paragraph(block, delete_context)
                if block_type in {"instruction", "example"}:
                    delete_context = self.classifier.next_context(block_type, text, delete_context)
                    self._remove_paragraph(block)
                    continue
                if chapter_strategies.get(current_chapter_key) == "fixed_keep":
                    delete_context = ""
                    continue
                if self._paragraph_has_preserved_object(block):
                    self._replace_paragraph_runs(block, replacements)
                    delete_context = ""
                    continue
                should_apply_generated = (
                    current_chapter_key
                    and current_chapter_key in generated_sections
                    and (
                        not chapter_needs_target.get(current_chapter_key)
                        or self._is_replacement_target(block)
                    )
                )
                if should_apply_generated:
                    generated_items = generated_sections[current_chapter_key]
                    applied_count = applied_generated_counts.get(current_chapter_key, 0)
                    if applied_count >= len(generated_items):
                        self._replace_paragraph_runs(block, replacements)
                        delete_context = ""
                        continue
                    self._replace_paragraph_text(block, generated_items[applied_count])
                    applied_generated_counts[current_chapter_key] = applied_count + 1
                    if current_chapter_key in section_details:
                        section_details[current_chapter_key]["applied"] = True
                        section_details[current_chapter_key]["applied_count"] = applied_count + 1
                        section_details[current_chapter_key]["apply_target"] = (
                            "占位符/斜体/待补充段落"
                            if chapter_needs_target.get(current_chapter_key)
                            else "章节首个正文段落"
                        )
                    delete_context = ""
                    continue
                self._replace_paragraph_runs(block, replacements)
                delete_context = ""
            elif isinstance(block, Table):
                if delete_context == "example" or self._is_red_table(block):
                    self._remove_table(block)
                    delete_context = ""
                    continue
                if chapter_strategies.get(current_chapter_key) == "fixed_keep":
                    delete_context = ""
                    continue
                self._fill_known_table(
                    block,
                    inputs,
                    replacements,
                    strategy=chapter_strategies.get(current_chapter_key, ""),
                )
                self._replace_table_runs(
                    block,
                    replacements,
                    preserve_labels=chapter_strategies.get(current_chapter_key) == "table_fill",
                )
                delete_context = ""

        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)
        return output_path

    def _generate_sections(
        self,
        *,
        template,
        title: str,
        inputs: Dict[str, Any],
        chapters: List[Any],
        llm_client,
        enabled: bool,
        section_details: Dict[str, dict],
        chapter_targets: Dict[str, List[str]],
        control_callback: Optional[Callable[[], None]] = None,
    ) -> Dict[str, List[str]]:
        generated = {}
        for chapter in chapters:
            if control_callback:
                control_callback()
            key = self._chapter_key(chapter)
            strategy = self._chapter_strategy(chapter)
            detail = {
                "chapter_key": key,
                "chapter": self._chapter_label(chapter),
                "strategy": strategy,
                "llm_called": False,
                "success": False,
                "prompt_summary": "",
                "prompt": "",
                "response": "",
                "error": "",
                "applied": False,
                "applied_count": 0,
                "apply_target": "",
                "target_count": len(chapter_targets.get(key) or []),
                "requests": [],
            }
            section_details[key] = detail
            if not enabled or not llm_client or not self._chapter_should_generate(chapter):
                continue
            targets = chapter_targets.get(key) or [""]
            contents = []
            for target_index, target_text in enumerate(targets, start=1):
                prompt = ""
                try:
                    if control_callback:
                        control_callback()
                    prompt = self.prompt_builder.build(
                        title=title,
                        template_name=getattr(template, "name", ""),
                        chapter=chapter,
                        inputs=inputs,
                        target_text=target_text,
                        target_index=target_index,
                        target_total=len(targets),
                    )
                    detail["llm_called"] = True
                    response = llm_client.generate(prompt, system_prompt=self.prompt_builder.system_prompt)
                    if control_callback:
                        control_callback()
                    raw_content = getattr(response, "content", "") or ""
                    content = self._clean_generated_content(raw_content)
                    request_detail = {
                        "target_index": target_index,
                        "target_text": target_text,
                        "prompt": prompt,
                        "response": raw_content,
                        "success": bool(content),
                        "error": "",
                    }
                    detail["requests"].append(request_detail)
                    if content:
                        contents.append(content)
                except Exception as exc:
                    if getattr(exc, "is_task_control", False):
                        raise
                    detail["requests"].append({
                        "target_index": target_index,
                        "target_text": target_text,
                        "prompt": prompt,
                        "response": "",
                        "success": False,
                        "error": str(exc),
                    })
                    logger.warning(
                        "章节智能生成失败，回退模板填充 - chapter=%s target=%s error=%s",
                        key,
                        target_index,
                        exc,
                    )
            if detail["requests"]:
                detail["prompt"] = "\n\n--- 回填位置 ---\n\n".join(
                    item["prompt"] for item in detail["requests"] if item["prompt"]
                )
                detail["response"] = "\n\n--- 模型返回 ---\n\n".join(
                    item["response"] for item in detail["requests"] if item["response"]
                )
                detail["prompt_summary"] = self._prompt_summary(detail["prompt"])
                errors = [item["error"] for item in detail["requests"] if item["error"]]
                detail["error"] = "；".join(errors)
            if contents:
                generated[key] = contents
                detail["success"] = len(contents) == len(targets)
        return generated

    def _collect_generation_targets(self, doc, chapters: List[Any]) -> Dict[str, List[str]]:
        targets: Dict[str, List[str]] = {}
        current_chapter_key = ""
        delete_context = ""
        for block in self._iter_block_items(doc):
            if not isinstance(block, Paragraph):
                continue
            text = block.text.strip()
            if self.classifier.heading_level(block) and not (
                delete_context and self.classifier.is_context_continuation(text)
            ):
                current_chapter_key = self._match_chapter_key(text, chapters)
                delete_context = ""
                continue
            block_type = self.classifier.classify_paragraph(block, delete_context)
            if block_type in {"instruction", "example"}:
                delete_context = self.classifier.next_context(block_type, text, delete_context)
                continue
            delete_context = ""
            if current_chapter_key and not self._paragraph_has_preserved_object(block) and self._is_replacement_target(block):
                targets.setdefault(current_chapter_key, []).append(text)
        return targets

    def _chapter_should_generate(self, chapter) -> bool:
        if self._chapter_strategy(chapter) != "smart_generate":
            return False
        if getattr(chapter, "guidance_prompt", ""):
            return True
        blocks = getattr(chapter, "template_blocks", []) or []
        if any(block.get("type") in {"instruction", "example", "instruction_table", "example_table"} for block in blocks):
            return True
        return bool(getattr(chapter, "placeholders", []) or blocks)

    def _chapter_strategy(self, chapter) -> str:
        value = str(getattr(chapter, "generation_strategy", "") or "smart_generate")
        if value not in {"fixed_keep", "placeholder_replace", "smart_generate", "table_fill"}:
            return "smart_generate"
        return value

    def _chapter_label(self, chapter) -> str:
        return f"{getattr(chapter, 'number', '')} {getattr(chapter, 'title', '')}".strip()

    def _prompt_summary(self, prompt: str) -> str:
        lines = [line.strip() for line in str(prompt or "").splitlines() if line.strip()]
        return " / ".join(lines[:6])[:500]

    def _chapter_needs_specific_target(self, chapter) -> bool:
        return bool(getattr(chapter, "placeholders", []) or self._chapter_strategy(chapter) == "placeholder_replace")

    def _is_replacement_target(self, paragraph) -> bool:
        if self._paragraph_has_preserved_object(paragraph):
            return False
        text = paragraph.text.strip()
        if not text:
            return True
        if any(run.italic and run.text.strip() for run in paragraph.runs):
            return True
        if any(token in text for token in ("NNN", "XXX", "X项目", "X 项目", "待补充", "待填写", "待替换", "____", "＿")):
            return True
        return False

    def _clean_generated_content(self, content: str) -> str:
        lines = [line.strip() for line in str(content or "").splitlines()]
        cleaned = []
        previous = ""
        for line in lines:
            if not line or self.classifier.is_example_marker(line):
                continue
            if self._looks_like_markdown_table_line(line):
                continue
            if line == previous:
                continue
            cleaned.append(line)
            previous = line
        return "\n".join(cleaned).strip()

    def _looks_like_markdown_table_line(self, line: str) -> bool:
        text = str(line or "").strip()
        if text.count("|") < 2:
            return False
        compact = text.replace("|", "").replace(":", "").replace("-", "").strip()
        return not compact or text.startswith("|") or "文件编号" in text

    def _is_template_document_title(self, text: str, target_title: str = "") -> bool:
        normalized = str(text or "").strip().replace(" ", "")
        if not normalized:
            return False
        if target_title and normalized == str(target_title).strip().replace(" ", ""):
            return False
        title_markers = ("文档标题", "文档名称", "文件名称", "试验大纲", "技术文件", "技术文档", "说明书", "规范", "报告", "方案", "项目")
        template_markers = ("XX", "XXX", "某型", "模板", "模版", "项目")
        if "文档标题" in normalized and any(token in normalized for token in title_markers):
            return True
        if len(normalized) <= 60 and any(token in normalized for token in title_markers):
            return any(marker in normalized for marker in template_markers)
        return False

    def _replace_document_title_candidates(self, doc: DocumentType, title: str, replacements: Dict[str, str]):
        title = str(title or "").strip()
        if not title:
            return
        visited = 0
        for block in self._iter_block_items(doc):
            if isinstance(block, Paragraph):
                text = block.text.strip()
                if self.classifier.heading_level(block):
                    break
                if self._is_template_document_title(text, title):
                    self._replace_paragraph_text(block, title)
                    return
                visited += 1
            elif isinstance(block, Table):
                for row in block.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            text = paragraph.text.strip()
                            if self._is_template_document_title(text, title):
                                self._replace_paragraph_text(paragraph, title)
                                return
                            visited += 1
            if visited >= 40:
                break

    def _build_replacements(self, title: str, inputs: Dict[str, Any]) -> Dict[str, str]:
        product_name = str(inputs.get("product_name", "")).strip()
        project_name = str(inputs.get("project_name", "")).strip() or str(inputs.get("background", "")).strip() or title
        product_model = str(inputs.get("product_model", "") or inputs.get("model_number", "")).strip()
        document_code = str(inputs.get("document_code", "") or inputs.get("file_code", "")).strip()
        confidentiality = str(inputs.get("confidentiality", "") or inputs.get("secret_level", "")).strip()
        version = str(inputs.get("version", "") or inputs.get("revision", "")).strip()
        author = str(inputs.get("author", "") or inputs.get("prepared_by", "")).strip()
        reviewer = str(inputs.get("reviewer", "") or inputs.get("checked_by", "")).strip()
        approver = str(inputs.get("approver", "") or inputs.get("approved_by", "")).strip()
        date = str(inputs.get("date", "") or inputs.get("document_date", "")).strip()
        test_item = str(inputs.get("test_item", "")).strip() or product_name
        replacements = {
            "(文件名称)": title,
            "（文件名称）": title,
            "(文档标题)": title,
            "（文档标题）": title,
            "(产品名称)": product_name,
            "（产品名称）": product_name,
            "(产品型号)": product_model,
            "（产品型号）": product_model,
            "(文件代号)": document_code,
            "（文件代号）": document_code,
            "(密级)": confidentiality,
            "（密级）": confidentiality,
            "(版次)": version,
            "（版次）": version,
            "(编制)": author,
            "（编制）": author,
            "(审核)": reviewer,
            "（审核）": reviewer,
            "(批准)": approver,
            "（批准）": approver,
            "(系统年)年(系统月)月": date,
            "（系统年）年（系统月）月": date,
            "NNNNN": product_name,
            "NNNN": product_name,
            "NNN": product_name,
            "XXX": project_name,
            "XX": project_name,
            "X 项目": project_name,
            "X项目": project_name,
            "受试设备": test_item,
        }
        dynamic_values = inputs.get("dynamic_fields") or {}
        for definition in inputs.get("dynamic_field_definitions") or []:
            if not isinstance(definition, dict):
                continue
            value = str(dynamic_values.get(definition.get("key"), "") or "").strip()
            if not value:
                continue
            for token in definition.get("placeholder_tokens") or []:
                token = str(token or "").strip()
                if token:
                    replacements[token] = value
        return {key: value for key, value in replacements.items() if value}

    def _collect_pending_meta_fields(self, inputs: Dict[str, Any]) -> List[dict]:
        field_defs = [
            ("confidentiality", "密级"),
            ("document_code", "文件代号"),
            ("product_model", "产品型号"),
            ("version", "版次"),
            ("author", "编制"),
            ("reviewer", "审核"),
            ("approver", "批准"),
            ("date", "日期"),
        ]
        aliases = {
            "confidentiality": ("confidentiality", "secret_level"),
            "document_code": ("document_code", "file_code"),
            "product_model": ("product_model", "model_number"),
            "version": ("version", "revision"),
            "author": ("author", "prepared_by"),
            "reviewer": ("reviewer", "checked_by"),
            "approver": ("approver", "approved_by"),
            "date": ("date", "document_date"),
        }
        pending = []
        for key, label in field_defs:
            if any(str(inputs.get(alias, "") or "").strip() for alias in aliases.get(key, (key,))):
                continue
            pending.append({
                "field": key,
                "label": label,
                "message": f"{label}未从生成素材中识别，已保留模板占位，请人工确认。",
            })
        return pending

    def _replace_document_meta_placeholders(self, doc: DocumentType, replacements: Dict[str, str]):
        for paragraph in doc.paragraphs:
            self._replace_paragraph_runs(paragraph, replacements)
        for table in doc.tables:
            self._replace_table_runs(table, replacements, preserve_labels=True)
        for section in doc.sections:
            for container in (section.header, section.first_page_header, section.even_page_header,
                              section.footer, section.first_page_footer, section.even_page_footer):
                for paragraph in container.paragraphs:
                    self._replace_paragraph_runs(paragraph, replacements)
                for table in container.tables:
                    self._replace_table_runs(table, replacements, preserve_labels=True)

    def _replace_paragraph_runs(self, paragraph, replacements: Dict[str, str]):
        for run in paragraph.runs:
            original = run.text
            replaced = self._replace_text(original, replacements)
            if run.italic and replaced == original:
                stripped = original.strip()
                if stripped:
                    replacement = self._replacement_for_italic_text(stripped, replacements)
                    if replacement:
                        replaced = original.replace(stripped, replacement)
                        run.italic = False
            if replaced != original:
                run.text = replaced

    def _replace_paragraph_text(self, paragraph, text: str):
        if paragraph.runs:
            paragraph.runs[0].text = text
            for run in paragraph.runs[1:]:
                run.text = ""
            return
        paragraph.add_run(text)

    def _replace_text(self, text: str, replacements: Dict[str, str]) -> str:
        result = text
        for placeholder, replacement in sorted(replacements.items(), key=lambda item: len(item[0]), reverse=True):
            result = result.replace(placeholder, replacement)
        return result

    def _replacement_for_italic_text(self, text: str, replacements: Dict[str, str]) -> str:
        if "项目" in text:
            return replacements.get("XXX") or replacements.get("XX") or ""
        return replacements.get("NNNNN") or replacements.get("受试设备") or ""

    def _is_red_paragraph(self, paragraph) -> bool:
        return self.classifier.is_red_paragraph(paragraph)

    def _is_red_run(self, run) -> bool:
        return self.classifier.is_red_run(run)

    def _is_example_paragraph(self, paragraph) -> bool:
        return self.classifier.is_example_marker(paragraph.text)

    def _paragraph_has_preserved_object(self, paragraph) -> bool:
        return any(
            self._xml_local_name(element.tag) in {"oMath", "oMathPara", "drawing", "pict", "object"}
            for element in paragraph._element.iter()
        )

    def _xml_local_name(self, tag: str) -> str:
        return str(tag or "").rsplit("}", 1)[-1]

    def _is_red_table(self, table) -> bool:
        paragraphs = [
            paragraph
            for row in table.rows
            for cell in row.cells
            for paragraph in cell.paragraphs
            if paragraph.text.strip()
        ]
        if not paragraphs:
            return False
        red_count = sum(1 for paragraph in paragraphs if self._is_red_paragraph(paragraph))
        return red_count / len(paragraphs) >= 0.5

    def _replace_table_runs(
        self,
        table,
        replacements: Dict[str, str],
        preserve_labels: bool = False,
    ):
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                for paragraph in list(cell.paragraphs):
                    if self._is_red_paragraph(paragraph):
                        self._remove_paragraph(paragraph)
                    elif (
                        not preserve_labels
                        and self.classifier.classify_paragraph(paragraph) in {"instruction", "example"}
                    ):
                        self._remove_paragraph(paragraph)
                    else:
                        self._replace_paragraph_runs(paragraph, replacements)

    def _fill_known_table(
        self,
        table,
        inputs: Dict[str, Any],
        replacements: Dict[str, str],
        strategy: str = "",
    ):
        headers = [cell.text.strip() for cell in table.rows[0].cells] if table.rows else []
        if "文件编号" in headers and "文件名" in headers:
            rows = self._reference_rows(inputs, replacements)
            if rows:
                self._fill_table_rows(table, rows)
            return
        if strategy == "table_fill":
            self._fill_generic_table(table, inputs, replacements)

    def _fill_generic_table(self, table, inputs: Dict[str, Any], replacements: Dict[str, str]):
        if not table.rows or len(table.rows[0].cells) < 2:
            return
        material_rows = self._structured_material_rows(inputs, replacements)
        material_lookup = {
            self._normalize_table_key(key): value
            for key, value in material_rows
            if self._normalize_table_key(key)
        }
        matched_existing = False
        for row in table.rows[1:]:
            label = row.cells[0].text.strip()
            if not label:
                continue
            value = material_lookup.get(self._normalize_table_key(label))
            if value:
                self._set_cell_text_preserving_style(row.cells[1], value)
                matched_existing = True
            elif self._cell_needs_fill(row.cells[1]):
                self._set_cell_text_preserving_style(row.cells[1], "待补充")
        if matched_existing or not material_rows:
            return
        self._fill_table_rows(table, [[key, value] for key, value in material_rows])

    def _structured_material_rows(
        self,
        inputs: Dict[str, Any],
        replacements: Dict[str, str],
    ) -> List[tuple[str, str]]:
        rows = []
        for field in ("technical_params", "additional_context", "supplement_doc_text"):
            raw = str(inputs.get(field, "") or "")
            for line in raw.splitlines():
                text = self._replace_text(line.strip().strip(";；"), replacements)
                if not text or text.startswith("表格"):
                    continue
                pair = self._split_material_pair(text)
                if pair:
                    rows.append(pair)
        deduplicated = {}
        for key, value in rows:
            deduplicated[key] = value or "待补充"
        return list(deduplicated.items())

    def _split_material_pair(self, text: str) -> Optional[tuple[str, str]]:
        import re

        text = re.sub(r"^\s*(?:[-*•]|\d+[\.、])\s*", "", text).strip()
        for separator in ("|", "：", ":", "=", "＝"):
            if separator not in text:
                continue
            key, value = text.split(separator, 1)
            key = key.strip()
            value = value.strip()
            if key:
                return key, value or "待补充"
        return None

    def _normalize_table_key(self, text: str) -> str:
        import re

        return re.sub(r"[\s：:（）()]+", "", str(text or "")).lower()

    def _cell_needs_fill(self, cell) -> bool:
        text = cell.text.strip()
        return not text or self._is_placeholder_text(text)

    def _is_placeholder_text(self, text: str) -> bool:
        return any(token in str(text or "") for token in (
            "NNN", "XXX", "待补充", "待填写", "待替换", "____", "＿",
        ))

    def _set_cell_text_preserving_style(self, cell, value: str):
        paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        self._replace_paragraph_text(paragraph, value)
        for extra in list(cell.paragraphs[1:]):
            self._remove_paragraph(extra)

    def _reference_rows(self, inputs: Dict[str, Any], replacements: Dict[str, str]) -> List[List[str]]:
        raw = str(inputs.get("references", "")).strip()
        if not raw:
            return []
        rows = []
        for line in raw.splitlines():
            text = line.strip().strip(";；")
            if not text:
                continue
            text = self._replace_text(text, replacements)
            parts = [part.strip() for part in text.split("|")]
            if len(parts) >= 2:
                rows.append((parts + [""])[:3])
                continue
            rows.append(self._parse_reference_line(text))
        return rows

    def _parse_reference_line(self, text: str) -> List[str]:
        import re

        text = re.sub(r"^\s*(?:\[\d+\]|\d+[\.、])\s*", "", text).strip()
        date_match = re.search(r"(\d{4}(?:[-./]\d{1,2})?(?:[-./]\d{1,2})?)$", text)
        date = date_match.group(1) if date_match else ""
        if date:
            text = text[:date_match.start()].strip()
        parts = text.split(maxsplit=1)
        if len(parts) == 2 and re.search(r"\d", parts[0]):
            return [parts[0], parts[1], date]
        return ["", text, date]

    def _fill_table_rows(self, table, data_rows: List[List[str]]):
        from copy import deepcopy

        while len(table.rows) < len(data_rows) + 1:
            if len(table.rows) > 1:
                template_row = table.rows[-1]._tr
                table._tbl.append(deepcopy(template_row))
            else:
                table.add_row()
        for index, values in enumerate(data_rows, start=1):
            for col_index, value in enumerate(values):
                if col_index >= len(table.rows[index].cells):
                    break
                self._set_cell_text_preserving_style(table.rows[index].cells[col_index], value)
        for row_index in range(len(data_rows) + 1, len(table.rows)):
            for cell in table.rows[row_index].cells:
                self._set_cell_text_preserving_style(cell, "")

    def _remove_paragraph(self, paragraph):
        element = paragraph._element
        parent = element.getparent()
        parent.remove(element)
        paragraph._p = paragraph._element = None

    def _remove_table(self, table):
        element = table._element
        parent = element.getparent()
        parent.remove(element)

    def _chapter_key(self, chapter) -> str:
        return f"{getattr(chapter, 'number', '')}::{getattr(chapter, 'title', '')}".strip(":")

    def _match_chapter_key(self, heading_text: str, chapters: List[Any]) -> str:
        normalized = self._normalize_heading(heading_text)
        for chapter in chapters:
            candidates = [
                f"{getattr(chapter, 'number', '')} {getattr(chapter, 'title', '')}".strip(),
                str(getattr(chapter, "title", "") or ""),
            ]
            if any(self._normalize_heading(candidate) == normalized for candidate in candidates if candidate):
                return self._chapter_key(chapter)
        return ""

    def _normalize_heading(self, text: str) -> str:
        import re

        return re.sub(r"\s+", "", str(text or "")).strip("：:。.、")

    def _iter_block_items(self, parent):
        if isinstance(parent, DocumentType):
            parent_elm = parent.element.body
        else:
            parent_elm = parent._tc
        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

class GeneratorFactory:
    """文档生成器工厂。"""

    _generators: Dict[str, type] = {}

    @classmethod
    def register(cls, name: str, generator_class: type):
        cls._generators[name] = generator_class

    @classmethod
    def create(cls, name: str, **kwargs) -> BaseGenerator:
        gen_class = cls._generators.get(name)
        if gen_class is None:
            raise ValueError(f"Unknown generator: {name}")
        return gen_class(**kwargs)

    @classmethod
    def available_generators(cls) -> List[str]:
        return list(cls._generators.keys())


# 注册内置生成器
GeneratorFactory.register("template_docx", TemplateDocxGenerator)
