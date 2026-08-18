"""Basic quality checks for generated DOCX files."""
from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document


class GeneratedDocxQualityChecker:
    """Run deterministic checks before later LLM review is wired in."""

    placeholder_pattern = re.compile(r"(?:N{3,}|X{2,}|X\s*项目|X项目|＿+|_{2,}|待填写|待替换)")
    guidance_pattern = re.compile(r"(?:【\s*)?(?:说明|举例|示例|填写说明|编写说明|填写要求)(?:\s*】)?\s*[:：]?")

    def check(self, file_path: str, *, template=None, inputs: dict[str, Any] | None = None, title: str = "") -> dict[str, Any]:
        inputs = inputs or {}
        path = Path(file_path)
        checks = []
        if not path.exists():
            return {
                "passed": False,
                "total_issues": 1,
                "checks": [{
                    "code": "file_missing",
                    "label": "生成文件存在性",
                    "passed": False,
                    "message": "生成文件不存在",
                }],
            }

        doc = Document(str(path))
        body_text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
        table_text = "\n".join(
            cell.text.strip()
            for table in doc.tables
            for row in table.rows
            for cell in row.cells
            if cell.text.strip()
        )
        all_text = "\n".join(part for part in (body_text, table_text) if part)
        text_locations = self._collect_text_locations(doc)
        placeholder_locations = self._matching_locations(text_locations, self.placeholder_pattern)
        guidance_locations = self._matching_locations(text_locations, self.guidance_pattern)

        checks.append(self._text_check(
            "placeholder_residue",
            "占位符残留",
            not placeholder_locations,
            "未发现明显占位符残留",
            "生成文档中仍存在明显占位符，请补充素材或检查模板解析结果",
            placeholder_locations,
        ))
        checks.append(self._text_check(
            "guidance_residue",
            "说明/示例残留",
            not guidance_locations,
            "未发现明显说明或示例残留",
            "生成文档中可能仍存在说明、举例或填写要求文字",
            guidance_locations,
        ))
        checks.append(self._text_check(
            "content_non_empty",
            "正文内容",
            bool(all_text.strip()),
            "生成文档包含正文内容",
            "生成文档正文为空",
        ))
        empty_table_locations = self._empty_table_locations(doc)
        empty_tables = len(empty_table_locations)
        checks.append(self._text_check(
            "table_non_empty",
            "表格填写",
            empty_tables == 0,
            "未发现完全空白表格",
            f"发现 {empty_tables} 个空白表格，请确认是否需要补充",
            empty_table_locations,
        ))
        product_name = str(inputs.get("product_name", "") or "").strip()
        if product_name:
            product_locations = self._text_locations_containing(text_locations, product_name)
            checks.append(self._text_check(
                "product_name_present",
                "产品名称一致性",
                bool(product_locations),
                "文档中包含产品名称",
                "文档中未发现用户输入的产品名称",
                [],
            ))
        if title:
            title_locations = self._text_locations_containing(text_locations, str(title).strip())
            checks.append(self._text_check(
                "title_present",
                "标题一致性",
                bool(title_locations),
                "文档中包含生成标题",
                "文档中未明显发现生成标题，请确认模板标题是否需要替换",
                self._first_locations(text_locations, 3) if not title_locations else title_locations[:3],
            ))
        missing_required = self._missing_required_chapters(doc, template)
        checks.append(self._text_check(
            "required_chapters",
            "必填章节",
            not missing_required,
            "未发现必填章节缺失",
            "可能缺少必填章节：" + "、".join(missing_required[:8]),
        ))
        pending_count = all_text.count("待补充")
        pending_locations = self._text_locations_containing(text_locations, "待补充")
        checks.append(self._text_check(
            "pending_content",
            "待补充内容",
            pending_count == 0,
            "未发现待补充内容",
            f"发现 {pending_count} 处“待补充”，请确认是否允许保留",
            pending_locations,
        ))

        total_issues = sum(1 for item in checks if not item["passed"])
        return {
            "passed": total_issues == 0,
            "total_issues": total_issues,
            "checks": checks,
        }

    def _empty_table_locations(self, doc) -> list[dict[str, str]]:
        locations = []
        for index, table in enumerate(doc.tables, start=1):
            values = [
                cell.text.strip()
                for row in table.rows
                for cell in row.cells
                if cell.text.strip()
            ]
            if not values:
                locations.append({
                    "location": f"表格 {index}",
                    "excerpt": "完全空白表格",
                })
        return locations

    def _collect_text_locations(self, doc) -> list[dict[str, str]]:
        locations = []
        for index, paragraph in enumerate(doc.paragraphs, start=1):
            text = paragraph.text.strip()
            if text:
                locations.append({
                    "location": f"段落 {index}",
                    "excerpt": self._excerpt(text),
                })
        for table_index, table in enumerate(doc.tables, start=1):
            for row_index, row in enumerate(table.rows, start=1):
                for cell_index, cell in enumerate(row.cells, start=1):
                    text = cell.text.strip()
                    if text:
                        locations.append({
                            "location": f"表格 {table_index} 第 {row_index} 行第 {cell_index} 列",
                            "excerpt": self._excerpt(text),
                        })
        for section_index, section in enumerate(doc.sections, start=1):
            containers = [
                ("页眉", section.header),
                ("首页页眉", section.first_page_header),
                ("偶数页页眉", section.even_page_header),
                ("页脚", section.footer),
                ("首页页脚", section.first_page_footer),
                ("偶数页页脚", section.even_page_footer),
            ]
            for label, container in containers:
                for paragraph_index, paragraph in enumerate(container.paragraphs, start=1):
                    text = paragraph.text.strip()
                    if text:
                        locations.append({
                            "location": f"第 {section_index} 节{label}段落 {paragraph_index}",
                            "excerpt": self._excerpt(text),
                        })
                for table_index, table in enumerate(container.tables, start=1):
                    for row_index, row in enumerate(table.rows, start=1):
                        for cell_index, cell in enumerate(row.cells, start=1):
                            text = cell.text.strip()
                            if text:
                                locations.append({
                                    "location": f"第 {section_index} 节{label}表格 {table_index} 第 {row_index} 行第 {cell_index} 列",
                                    "excerpt": self._excerpt(text),
                                })
        return locations

    def _matching_locations(self, locations: list[dict[str, str]], pattern: re.Pattern) -> list[dict[str, str]]:
        return [
            location
            for location in locations
            if pattern.search(location.get("excerpt", ""))
        ][:8]

    def _text_locations_containing(self, locations: list[dict[str, str]], needle: str) -> list[dict[str, str]]:
        value = str(needle or "").strip()
        if not value:
            return []
        return [
            location
            for location in locations
            if value in location.get("excerpt", "")
        ][:8]

    def _first_locations(self, locations: list[dict[str, str]], limit: int) -> list[dict[str, str]]:
        return locations[:limit]

    def _excerpt(self, text: str, limit: int = 96) -> str:
        value = re.sub(r"\s+", " ", str(text or "")).strip()
        if len(value) <= limit:
            return value
        return value[:limit].rstrip() + "..."

    def _missing_required_chapters(self, doc, template) -> list[str]:
        if not template:
            return []
        existing = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
        missing = []
        for chapter in self._flatten_chapters(getattr(template, "chapters", []) or []):
            if not getattr(chapter, "required", True):
                continue
            title = str(getattr(chapter, "title", "") or "").strip()
            number = str(getattr(chapter, "number", "") or "").strip()
            candidates = [f"{number} {title}".strip(), title]
            if title and not any(candidate and candidate in existing for candidate in candidates):
                missing.append(f"{number} {title}".strip())
        return missing

    def _flatten_chapters(self, chapters) -> list[Any]:
        flat = []
        for chapter in chapters:
            flat.append(chapter)
            flat.extend(self._flatten_chapters(getattr(chapter, "sub_chapters", []) or []))
        return flat

    def _text_check(
        self,
        code: str,
        label: str,
        passed: bool,
        ok: str,
        fail: str,
        locations: list[dict[str, str]] | None = None,
    ) -> dict[str, Any]:
        item = {
            "code": code,
            "label": label,
            "passed": bool(passed),
            "message": ok if passed else fail,
        }
        if not passed and locations:
            item["locations"] = locations
        return item
