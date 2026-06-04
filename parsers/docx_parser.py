import re
from abc import ABC, abstractmethod
from typing import Optional
from docx import Document
from models.document import ParsedDocument, DocumentRegionType, DocumentSection, ContentType
from parsers.structure_parser import DocumentStructureParser


class BaseParser(ABC):
    @abstractmethod
    def parse(self, file_path: str) -> ParsedDocument:
        pass


class DocxParser(BaseParser):
    def __init__(self, chunk_size: int = 2500):
        self.section_counter = 0
        self.chunk_size = chunk_size
    
    def parse(self, file_path: str) -> ParsedDocument:
        self.section_counter = 0
        doc = Document(file_path)
        atomic_sections = []
        raw_text_parts = []
        
        title = self._extract_title(doc)
        
        for element in doc.element.body:
            if element.tag.endswith('p'):
                section = self._parse_paragraph(element, doc)
                if section:
                    atomic_sections.append(section)
            elif element.tag.endswith('tbl'):
                section = self._parse_table(element, doc)
                if section:
                    atomic_sections.append(section)

        structure = DocumentStructureParser().parse(atomic_sections)
        regions_by_section = {
            section_id: region.region_type.value
            for region in structure.regions
            for section_id in region.section_ids
        }
        sections = self._merge_sections(atomic_sections, regions_by_section)
        raw_text_parts = [section.text for section in sections]
        
        # raw_text = "\n".join(raw_text_parts)
        raw_text = "\n\n".join(raw_text_parts)
        
        return ParsedDocument(
            file_path=file_path,
            title=title,
            sections=sections,
            raw_text=raw_text,
            metadata={"total_sections": len(sections)},
            structure=structure,
        )

    def _merge_sections(self, sections: list[DocumentSection], regions_by_section: dict[str, str] = None) -> list[DocumentSection]:
        """按标题和长度聚合段落，减少后续 LLM 审查请求次数。"""
        merged: list[DocumentSection] = []
        current: list[DocumentSection] = []
        regions_by_section = regions_by_section or {}

        def flush():
            if not current:
                return
            first = current[0]
            text = "\n".join(s.text for s in current if s.text).strip()
            if not text:
                current.clear()
                return
            heading_text = ""
            for source in current:
                if source.content_type == ContentType.HEADING:
                    heading_text = source.text.strip()
                    break
            region = regions_by_section.get(first.section_id, DocumentRegionType.BODY.value)
            table_cells = []
            for source in current:
                table_cells.extend(source.metadata.get("table_cells", []))
            merged.append(DocumentSection(
                section_id=f"section_{len(merged) + 1}",
                content_type=first.content_type,
                text=text,
                level=first.level,
                metadata={
                    "source_sections": [s.section_id for s in current],
                    "source_count": len(current),
                    "style": first.metadata.get("style", ""),
                    "heading_text": heading_text,
                    "region": region,
                    "target_text": text,
                    "table_cells": table_cells,
                }
            ))
            current.clear()

        for section in sections:
            is_heading = section.content_type == ContentType.HEADING
            current_len = sum(len(s.text) for s in current)
            next_len = current_len + len(section.text) + 1

            if current and (is_heading or next_len > self.chunk_size):
                flush()
            current.append(section)

        flush()
        return merged
    
    def _extract_title(self, doc: Document) -> str:
        if doc.paragraphs and doc.paragraphs[0].style.name.startswith('Heading'):
            return doc.paragraphs[0].text
        elif doc.paragraphs:
            return doc.paragraphs[0].text[:50] if doc.paragraphs[0].text else "Untitled"
        return "Untitled"

    def _get_heading_level(self, style_name: str) -> int:
        match = re.search(r'Heading\s*(\d+)', style_name)
        return int(match.group(1)) if match else 1
    
    def _parse_paragraph(self, element, doc: Document) -> Optional[DocumentSection]:
        from docx.text.paragraph import Paragraph
        
        para = Paragraph(element, doc)
        original_text = para.text  # 保留原始文本（包含空格）
        
        # 检查是否为空段落（只包含空白字符）
        if not original_text.strip():
            return None
        
        self.section_counter += 1
        section_id = f"section_{self.section_counter}"
        
        if para.style.name.startswith('Heading'):
            level = self._get_heading_level(para.style.name)
            content_type = ContentType.HEADING
        else:
            level = 0
            content_type = ContentType.PARAGRAPH
        
        return DocumentSection(
            section_id=section_id,
            content_type=content_type,
            text=original_text,  # 使用原始文本
            level=level,
            metadata={"style": para.style.name}
        )
    
    def _parse_table(self, element, doc: Document) -> Optional[DocumentSection]:
        from docx.table import Table
        
        table = Table(element, doc)
        rows_text = []
        table_cells = []
        
        for row_idx, row in enumerate(table.rows):
            cells_text = []
            for col_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                cells_text.append(cell_text)
                table_cells.append({
                    "row": row_idx,
                    "col": col_idx,
                    "text": cell_text,
                })
            # rows_text.append(" | ".join(cells_text))
            rows_text.append(" | ".join([f"[{j}] {cell}" for j, cell in enumerate(cells_text)]))

        text = "\n".join(rows_text)
        
        if not text.strip():
            return None
        
        self.section_counter += 1
        section_id = f"section_{self.section_counter}"
        
        return DocumentSection(
            section_id=section_id,
            content_type=ContentType.TABLE,
            text=text,
            metadata={"rows": len(table.rows), "cols": len(table.columns), "table_cells": table_cells}
        )


class ParserFactory:
    _parsers = {
        ".docx": DocxParser,
    }
    
    @classmethod
    def get_parser(cls, file_extension: str) -> Optional[BaseParser]:
        parser_class = cls._parsers.get(file_extension.lower())
        if parser_class:
            return parser_class()
        return None
    
    @classmethod
    def register_parser(cls, file_extension: str, parser_class: type):
        cls._parsers[file_extension.lower()] = parser_class
