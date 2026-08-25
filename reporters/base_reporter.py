from abc import ABC, abstractmethod
from typing import List, Optional
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from web.time_utils import beijing_now_str

from core.executor import DocumentReviewResult, SectionReviewResult
from rules.base_rule import RuleSeverity


class BaseReporter(ABC):
    @abstractmethod
    def generate(self, result: DocumentReviewResult) -> str:
        pass
    
    @abstractmethod
    def save(self, result: DocumentReviewResult, output_path: str):
        pass


# ----------------------------
# Markdown 报告
# ----------------------------
class MarkdownReporter(BaseReporter):
    def generate(self, result: DocumentReviewResult) -> str:
        lines = []
        
        lines.append("# 文档审查报告\n")
        
        # 基本信息
        lines.append("## 基本信息\n")
        lines.append(f"- **文档标题**: {result.document_title}")
        lines.append(f"- **文档路径**: `{result.document_path}`")
        lines.append(f"- **审查时间**: {result.review_time}")
        status_emoji = "✅" if result.overall_passed else "❌"
        lines.append(f"- **审查结果**: {status_emoji} {'通过' if result.overall_passed else '未通过'}\n")
        
        # 审查统计
        lines.append("## 审查统计\n")
        lines.append("| 指标 | 数量 |")
        lines.append("|------|------|")
        lines.append(f"| 总问题数 | {result.total_issues} |")
        lines.append(f"| 错误 | {result.errors} |")
        lines.append(f"| 警告 | {result.warnings} |\n")
        
        # 审查总结
        if result.summary:
            lines.append("## 审查总结\n")
            lines.append(result.summary + "\n")
        
        # 章节及问题
        lines.append("## 详细结果\n")
        for section in result.section_results:
            lines.append(f"### 章节 {section.section_id}\n")
            lines.append(f"> {section.section_text[:200]}...\n")
            
            # 规则问题
            if any(not r.passed for r in section.rule_results):
                lines.append("#### 规则问题\n")
                lines.append("| 规则 | 来源 | 严重程度 | 复核状态 | 描述 | 建议 |")
                lines.append("|------|------|----------|----------|------|------|")
                for r in section.rule_results:
                    if not r.passed:
                        severity_str = {
                            RuleSeverity.ERROR: "🔴 错误",
                            RuleSeverity.WARNING: "🟡 警告",
                            RuleSeverity.INFO: "🔵 信息"
                        }.get(r.severity, "未知")
                        suggestions = "; ".join(r.suggestions) if r.suggestions else "-"
                        review_status = r.details.get("review_status_label", "待复核")
                        lines.append(f"| {r.rule_name} | {r.rule_source} | {severity_str} | {review_status} | {r.message} | {suggestions} |")
                lines.append("")
            


        lines.append("---")
        lines.append(f"*报告生成时间: {beijing_now_str()}*")
        return "\n".join(lines)

    def save(self, result: DocumentReviewResult, output_path: str):
        content = self.generate(result)
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(content)


# ----------------------------
# JSON 报告
# ----------------------------
class JsonReporter(BaseReporter):
    def generate(self, result: DocumentReviewResult) -> str:
        import json
        data = {
            "document_title": result.document_title,
            "document_path": result.document_path,
            "review_time": result.review_time,
            "overall_passed": result.overall_passed,
            "statistics": {
                "total_issues": result.total_issues,
                "errors": result.errors,
                "warnings": result.warnings
            },
            "summary": result.summary,
            "sections": []
        }
        
        for section in result.section_results:
            sec = {
                "section_id": section.section_id,
                "content": section.section_text,
                "passed": section.passed,
                "rule_issues": []
            }
            
            # 规则问题（包括 RULE 和 LLM 结果）
            for r in section.rule_results:
                if not r.passed:
                    sec["rule_issues"].append({
                        "rule_name": r.rule_name,
                        "rule_source": r.rule_source,
                        "severity": r.severity.name,
                        "review_status": r.details.get("review_status", "pending"),
                        "review_status_label": r.details.get("review_status_label", "待复核"),
                        "message": r.message,
                        "suggestions": r.suggestions or []
                    })
            
            data["sections"].append(sec)
        
        # 全局LLM总结
        if getattr(result, 'llm_summary', None):
            data["llm_summary"] = result.llm_summary
        
        return json.dumps(data, ensure_ascii=False, indent=2)

    def save(self, result: DocumentReviewResult, output_path: str):
        content = self.generate(result)
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(content)


# ----------------------------
# DOCX 报告
# ----------------------------
class DocxReporter(BaseReporter):
    BLUE = RGBColor(31, 78, 121)
    HEADER_BLUE = "1F4E79"
    LABEL_FILL = "D9EAF7"
    GREEN = RGBColor(0, 112, 60)
    RED = RGBColor(192, 0, 0)
    ORANGE = RGBColor(197, 90, 17)

    def _setup_document(self, doc: Document):
        for section in doc.sections:
            section.top_margin = Inches(0.65)
            section.bottom_margin = Inches(0.65)
            section.left_margin = Inches(0.6)
            section.right_margin = Inches(0.6)

        normal = doc.styles["Normal"]
        normal.font.name = "宋体"
        normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        normal.font.size = Pt(10.5)

    def _set_table_style(self, table):
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    def _set_column_widths(self, table, widths):
        table.autofit = False
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = width

    def _shade_cell(self, cell, fill: str):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)

    def _set_cell_text(
        self,
        cell,
        text: str,
        *,
        bold: bool = False,
        color: Optional[RGBColor] = None,
        size: float = 10.5,
        align=None,
    ):
        cell.text = ""
        paragraph = cell.paragraphs[0]
        if align is not None:
            paragraph.alignment = align
        run = paragraph.add_run(str(text or "-"))
        run.bold = bold
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(size)
        if color is not None:
            run.font.color.rgb = color

    def _style_header_row(self, row, size: float = 9):
        for cell in row.cells:
            self._shade_cell(cell, self.HEADER_BLUE)
            self._set_cell_text(
                cell,
                cell.text,
                bold=True,
                color=RGBColor(255, 255, 255),
                size=size,
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )

    def _style_label_cell(self, cell):
        self._shade_cell(cell, self.LABEL_FILL)
        self._set_cell_text(cell, cell.text, bold=True, size=10.5)

    def _add_section_heading(self, doc: Document, text: str, status: Optional[str] = None):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(12)
        paragraph.paragraph_format.space_after = Pt(8)
        run = paragraph.add_run(text)
        run.bold = True
        run.font.name = "黑体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        run.font.size = Pt(16)
        run.font.color.rgb = self.BLUE
        if status:
            status_run = paragraph.add_run(f"【{status}】")
            status_run.bold = True
            status_run.font.name = "黑体"
            status_run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            status_run.font.size = Pt(16)
            status_run.font.color.rgb = self.RED
        return paragraph

    def _add_body_paragraph(self, doc: Document, text: str):
        paragraph = doc.add_paragraph(str(text or ""))
        paragraph.paragraph_format.space_after = Pt(6)
        return paragraph

    def _all_results(self, result: DocumentReviewResult):
        items = []
        for section in result.section_results:
            for rule_result in section.rule_results:
                items.append((section, rule_result))
        return items

    def _all_issues(self, result: DocumentReviewResult):
        issues = []
        for section in result.section_results:
            for rule_result in section.rule_results:
                if not rule_result.passed:
                    issues.append((section, rule_result))
        return issues

    def _review_status(self, result: DocumentReviewResult, issues) -> str:
        if not issues and result.overall_passed:
            return "通过"
        if any(self._severity_rank(issue) <= 1 for _, issue in issues):
            return "不通过"
        return "待确认"

    def _total_rule_count(self, result: DocumentReviewResult) -> int:
        configured_total = getattr(result, "total_rules", None) or getattr(result, "rule_count", None)
        if configured_total:
            return int(configured_total)
        all_results = self._all_results(result)
        return len(all_results) or int(result.total_issues or 0)

    def _pending_count(self, issues) -> int:
        return sum(1 for _, issue in issues if self._issue_result_label(issue) == "待确认")

    def _issue_result_label(self, issue) -> str:
        status = str(issue.details.get("review_status_label") or issue.details.get("review_status") or "")
        if "待" in status or "确认" in status or issue.rule_source == "LLM":
            return "待确认"
        return "不通过"

    def _severity_rank(self, issue) -> int:
        label = self._severity_display(issue)
        return {"严重": 0, "一般": 1, "建议": 2}.get(label, 3)

    def _severity_display(self, issue) -> str:
        if issue.severity == RuleSeverity.ERROR:
            return "严重"
        if issue.severity == RuleSeverity.WARNING:
            return "一般"
        return "建议"

    def _severity_color(self, label: str) -> RGBColor:
        return {"严重": self.RED, "一般": self.ORANGE, "建议": self.BLUE}.get(label, self.BLUE)

    def _result_color(self, label: str) -> RGBColor:
        return {"通过": self.GREEN, "不通过": self.RED, "待确认": self.ORANGE}.get(label, self.BLUE)

    def _issue_category(self, issue) -> str:
        text = " ".join(
            str(value or "")
            for value in [
                issue.rule_name,
                issue.rule_id,
                issue.message,
                issue.rule_reference,
                issue.details.get("rule_category"),
                issue.details.get("category"),
                issue.details.get("rule_set"),
                issue.details.get("rule_logic"),
            ]
        )
        category_keywords = [
            ("文档格式", ["标题", "编号", "格式", "目录", "页眉", "页脚", "字体", "字号", "页面"]),
            ("文档内容", ["章节缺失", "内容缺失", "要求缺失", "完整", "一致", "表述", "描述"]),
            ("技术要求", ["参数", "性能", "寿命", "可靠性", "安全", "指标", "接口", "功能"]),
            ("标准引用", ["标准", "规范", "引用", "依据", "文件"]),
        ]
        for category, keywords in category_keywords:
            if any(keyword in text for keyword in keywords):
                return category
        return "其他"

    def _category_detail_title(self, category: str) -> str:
        return {
            "文档格式": "文档格式规范审查（页面设置、字体字号等）",
            "文档内容": "文档内容审查",
            "技术要求": "技术要求审查",
            "标准引用": "标准引用审查",
            "其他": "其他审查",
        }.get(category, f"{category}审查")

    def _issue_location(self, section: SectionReviewResult, issue) -> str:
        location = issue.section_id or section.section_id or ""
        if location == "document_structure":
            location = "文档结构定位"
        if issue.position and issue.position.line_number:
            location = f"{location or '正文'} / 第 {issue.position.line_number} 行"
        return location or "[章节/页码/段落]"

    def _issue_suggestion(self, issue) -> str:
        if issue.suggestions:
            return "；".join(str(item) for item in issue.suggestions if item)
        if getattr(issue, "fix_suggestion", None) and issue.fix_suggestion.description:
            return issue.fix_suggestion.description
        return "[修改建议及引用规范]"

    def _issue_basis(self, issue) -> str:
        return issue.rule_reference or issue.details.get("rule_reference") or issue.details.get("basis") or "文档管理规范"

    def _issue_check_content(self, issue) -> str:
        return issue.details.get("check_content") or issue.details.get("rule_logic") or issue.rule_name or issue.rule_id or "-"

    def _document_type_checkboxes(self, result: DocumentReviewResult) -> str:
        doc_name = f"{result.document_title} {result.document_path}"
        options = ["技术方案", "设计规范", "试验报告", "工艺文件", "计算书", "说明书"]
        checked = []
        has_match = False
        for option in options:
            mark = "☑" if option in doc_name else "□"
            has_match = has_match or mark == "☑"
            checked.append(f"{mark} {option}")
        other = "□ 其他：____" if has_match else "☑ 其他：____"
        checked.append(other)
        return "  ".join(checked)

    def _document_code(self, result: DocumentReviewResult) -> str:
        for attr in ("document_code", "file_code", "doc_code"):
            value = getattr(result, attr, "")
            if value:
                return str(value)
        return "—"

    def _review_date(self, result: DocumentReviewResult) -> str:
        raw = result.review_time or getattr(result, "completed_at", "") or beijing_now_str()
        text = str(raw)
        if len(text) >= 10 and text[4] == "-" and text[7] == "-":
            year, month, day = text[:10].split("-")
            return f"{int(year)} 年 {int(month)} 月 {int(day)} 日"
        return text or "—"

    def _percent(self, value: int, total: int) -> str:
        if total <= 0:
            return "—"
        return f"{value / total * 100:.0f}%"

    def _sort_issues(self, issues):
        return sorted(issues, key=lambda item: (self._severity_rank(item[1]), self._issue_category(item[1])))

    def _add_basic_info_section(self, doc: Document, result: DocumentReviewResult):
        self._add_section_heading(doc, "一、文档基本信息")
        table = doc.add_table(rows=4, cols=2)
        self._set_table_style(table)
        self._set_column_widths(table, [Inches(1.45), Inches(5.55)])
        rows = [
            ("文件名称", result.document_title or Path(result.document_path).stem or "—"),
            ("文件代号", self._document_code(result)),
            ("文件类型", self._document_type_checkboxes(result)),
            ("审查时间", self._review_date(result)),
        ]
        for index, (label, value) in enumerate(rows):
            table.cell(index, 0).text = label
            self._style_label_cell(table.cell(index, 0))
            table.cell(index, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._set_cell_text(table.cell(index, 1), value, size=10.5)

    def _add_conclusion_section(self, doc: Document, result: DocumentReviewResult, issues, review_status: str):
        self._add_section_heading(doc, "二、审查结论", review_status)
        total = self._total_rule_count(result)
        pending = self._pending_count(issues)
        not_pass = max(len(issues) - pending, 0)
        passed = max(total - not_pass - pending, 0)
        dominant_category = self._dominant_category(issues)
        priority_category = self._priority_category(issues)
        self._add_body_paragraph(
            doc,
            f"本次审查共检查 {total} 项审查规则，其中通过 {passed} 项，不通过 {not_pass} 项，"
            f"待确认 {pending} 项。主要问题集中在 {dominant_category} 方面，建议优先整改 {priority_category} 类问题。"
        )
        table = doc.add_table(rows=1, cols=4)
        self._set_table_style(table)
        headers = ["统计指标", "数量", "占比", "状态"]
        for idx, header in enumerate(headers):
            table.rows[0].cells[idx].text = header
        self._style_header_row(table.rows[0], size=10)
        rows = [
            ("审查规则总数", total, "—", "—"),
            ("通过项", passed, self._percent(passed, total), "√ 通过"),
            ("不通过项", not_pass, self._percent(not_pass, total), "× 不通过"),
            ("待确认项", pending, self._percent(pending, total), "△ 待确认"),
        ]
        for label, count, percent, status in rows:
            cells = table.add_row().cells
            values = [label, str(count), percent, status]
            for idx, value in enumerate(values):
                color = self._result_color(status.replace("√ ", "").replace("× ", "").replace("△ ", "")) if idx == 3 else None
                self._set_cell_text(cells[idx], value, bold=idx == 3 and status != "—", color=color, align=WD_ALIGN_PARAGRAPH.CENTER)

    def _dominant_category(self, issues) -> str:
        if not issues:
            return "无明显问题"
        counts = {}
        for _, issue in issues:
            category = self._issue_category(issue)
            counts[category] = counts.get(category, 0) + 1
        return max(counts.items(), key=lambda item: item[1])[0]

    def _priority_category(self, issues) -> str:
        if not issues:
            return "无"
        return self._issue_category(self._sort_issues(issues)[0][1])

    def _add_issue_list_section(self, doc: Document, issues):
        self._add_section_heading(doc, "三、不符合项/问题清单")
        if not issues:
            self._add_body_paragraph(doc, "本次审查未发现不符合项。")
        else:
            self._add_body_paragraph(doc, "以下列出本次审查中发现的所有不符合项，按严重程度由高到低排列：")
            table = doc.add_table(rows=1, cols=6)
            self._set_table_style(table)
            headers = ["序号", "问题类别", "位置定位", "不符合项/问题描述", "修改建议", "严重程度"]
            for idx, header in enumerate(headers):
                table.rows[0].cells[idx].text = header
            self._style_header_row(table.rows[0], size=8.5)
            for idx, (section, issue) in enumerate(self._sort_issues(issues), start=1):
                cells = table.add_row().cells
                severity = self._severity_display(issue)
                row_values = [
                    str(idx),
                    self._issue_category(issue),
                    self._issue_location(section, issue),
                    issue.message or "[具体问题描述]",
                    self._issue_suggestion(issue),
                    severity,
                ]
                for col_idx, value in enumerate(row_values):
                    color = self._severity_color(severity) if col_idx == 5 else None
                    self._set_cell_text(cells[col_idx], value, bold=col_idx == 5, color=color, size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER if col_idx in (0, 5) else None)

        self._add_body_paragraph(doc, "严重程度分级说明：")
        for text in [
            "严重：存在技术性错误、数据不准确、标准引用错误等，影响文档的技术可靠性和可用性，必须整改后方可发布。",
            "一般：格式不规范、表述不统一、内容不完整等，影响文档的专业性和可读性，应在发布前完成整改。",
            "建议：优化性建议，不影响文档发布，建议在后续修订中改进。",
        ]:
            self._add_body_paragraph(doc, text)

    def _add_review_detail_section(self, doc: Document, issues):
        self._add_section_heading(doc, "四、本次审查详情")
        if not issues:
            self._add_body_paragraph(doc, "无不符合项审查详情。")
            return

        category_order = ["文档格式", "文档内容", "技术要求", "标准引用", "其他"]
        grouped = {category: [] for category in category_order}
        for item in self._sort_issues(issues):
            grouped.setdefault(self._issue_category(item[1]), []).append(item)

        detail_index = 1
        for category in category_order:
            items = grouped.get(category) or []
            if not items:
                continue
            heading = doc.add_paragraph()
            heading.paragraph_format.space_before = Pt(8)
            heading.paragraph_format.space_after = Pt(4)
            run = heading.add_run(f"1.{detail_index} {self._category_detail_title(category)}")
            run.bold = True
            run.font.name = "黑体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            run.font.size = Pt(12)
            table = doc.add_table(rows=1, cols=8)
            self._set_table_style(table)
            headers = ["序号", "位置定位", "审查内容", "审查依据", "审查结果", "问题描述", "修改建议", "备注"]
            for idx, header in enumerate(headers):
                table.rows[0].cells[idx].text = header
            self._style_header_row(table.rows[0], size=7.5)
            for row_idx, (section, issue) in enumerate(items, start=1):
                cells = table.add_row().cells
                result_label = self._issue_result_label(issue)
                values = [
                    str(row_idx),
                    self._issue_location(section, issue),
                    self._issue_check_content(issue),
                    self._issue_basis(issue),
                    result_label,
                    issue.message or "[具体问题描述]",
                    self._issue_suggestion(issue),
                    issue.details.get("remark") or "-",
                ]
                for col_idx, value in enumerate(values):
                    color = self._result_color(result_label) if col_idx == 4 else None
                    self._set_cell_text(cells[col_idx], value, bold=col_idx == 4, color=color, size=7.5, align=WD_ALIGN_PARAGRAPH.CENTER if col_idx in (0, 4) else None)
            detail_index += 1

    def generate(self, result: DocumentReviewResult) -> Document:
        doc = Document()
        self._setup_document(doc)
        
        # 标题
        title = doc.add_heading('文档审查报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.name = "黑体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            run.font.color.rgb = self.BLUE
        subtitle = doc.add_paragraph('面向用户的审查结论、问题清单与规则依据')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        issues = self._all_issues(result)
        review_status = self._review_status(result, issues)
        self._add_basic_info_section(doc, result)
        self._add_conclusion_section(doc, result, issues, review_status)
        self._add_issue_list_section(doc, issues)
        self._add_review_detail_section(doc, issues)
        
        # 页脚
        for section in doc.sections:
            footer = section.footer
            footer.add_paragraph(f"报告生成时间: {beijing_now_str()}")
        
        return doc

    def save(self, result: DocumentReviewResult, output_path: str):
        doc = self.generate(result)
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)


# ----------------------------
# Reporter Factory
# ----------------------------
class ReporterFactory:
    _reporters = {
        "md": MarkdownReporter,
        "markdown": MarkdownReporter,
        "json": JsonReporter,
        "docx": DocxReporter,
    }

    @classmethod
    def create_reporter(cls, format_type: str) -> BaseReporter:
        reporter_class = cls._reporters.get(format_type.lower())
        if reporter_class:
            return reporter_class()
        raise ValueError(f"Unknown reporter format: {format_type}")

    @classmethod
    def register_reporter(cls, format_type: str, reporter_class: type):
        cls._reporters[format_type.lower()] = reporter_class

    @classmethod
    def get_supported_formats(cls) -> List[str]:
        return list(cls._reporters.keys())
