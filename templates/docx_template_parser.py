"""DOCX template structure parser."""
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocumentType
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from .docx_block_classifier import DocxBlockClassifier


class DocxTemplateParser:
    """Parse DOCX headings into an editable generation template chapter tree."""

    heading_number_pattern = re.compile(r"^\s*(\d+(?:\.\d+)*)\s*[\.\、．]?\s*(.+?)\s*$")
    placeholder_pattern = re.compile(r"(?:N{3,}|X{2,}|X\s*项目|X项目|＿+|_{2,})")
    toc_page_pattern = re.compile(r".+(?:\t| {2,}|…{2,}|\.{2,})\s*(?:[IVXLCDM]+|\d+)\s*$", re.IGNORECASE)
    appendix_pattern = re.compile(r"^\s*附录\s*([A-ZＡ-Ｚ])(?:\s+|[（(]|$)(.*)$", re.IGNORECASE)
    figure_table_caption_pattern = re.compile(r"^\s*(?:表|图)\s*\d+(?:[-－—]\d+)?\s*")
    document_title_pattern = re.compile(r"(?:规范|报告|大纲|说明书|计划|方案|规程|手册|细则|要求)$")
    list_item_pattern = re.compile(
        r"^\s*(?P<marker>(?:[a-zA-Z]|[一二三四五六七八九十]+|\d+)[\)）\.、]|[（(]\d+[）)])\s*(?P<text>.+?)\s*$"
    )
    expandable_marker_pattern = re.compile(r"^\s*(?:…{2,}|\.{3,}|……+)\s*[。；;]?\s*$")

    def __init__(self):
        self.classifier = DocxBlockClassifier()

    def parse(self, file_path: str) -> dict[str, Any]:
        doc = Document(file_path)
        headings = []
        current_heading = None
        current_context = ""
        region = "front_matter"

        for block_item in self._iter_block_items(doc):
            if isinstance(block_item, Paragraph):
                text = block_item.text.strip()
                if not text:
                    continue
                style_name = block_item.style.name if block_item.style else ""

                if self._is_toc_start(text):
                    region = "toc"
                    current_context = ""
                    continue

                if region == "toc":
                    if self._is_toc_entry(text, style_name):
                        continue
                    region = "body" if self._is_body_start(block_item, text) else "front_matter"

                if region == "front_matter" and not self._is_body_start(block_item, text):
                    continue

                heading = self._detect_heading(block_item, current_context)
                if heading:
                    region = "body"
                    current_heading = heading
                    headings.append(current_heading)
                    current_context = ""
                    continue
                if self._looks_like_document_title(text):
                    continue

                if current_heading:
                    block = self._paragraph_block(block_item, current_context)
                    if self._is_empty_marker_block(block):
                        current_context = f"{block['type']}_marker"
                        continue
                    block = self._append_paragraph_block(current_heading, block)
                    current_heading["placeholders"].extend(block.get("placeholders", []))
                    if block["type"] in ("instruction", "example"):
                        current_heading["_guidance_parts"].append(
                            self.classifier.guidance_display_text(block["type"], text)
                        )
                    if not current_heading.get("body_style_name") and block["type"] in {"template_text", "template_list"}:
                        current_heading["body_style_name"] = style_name
                    current_context = self._next_context(block, current_context)
            elif isinstance(block_item, Table) and current_heading:
                block = self._table_block(block_item, current_context)
                current_heading["template_blocks"].append(block)
                current_heading["placeholders"].extend(block.get("placeholders", []))
                if block["type"] in ("instruction_table", "example_table"):
                    current_heading["_guidance_parts"].append(f"{block['label']}：{self._table_text(block)}")

        if not headings:
            headings = self._fallback_headings(doc)

        return {
            "name": Path(file_path).stem,
            "description": "由 DOCX 模板自动解析生成",
            "chapters": self._build_tree(headings),
        }

    def _detect_heading(self, paragraph, current_context: str = "") -> dict | None:
        text = paragraph.text.strip()
        if not text or self._is_non_chapter_numbered_text(paragraph, text):
            return None
        if current_context and self.classifier.is_context_continuation(text):
            return None
        if self.classifier.is_example_marker(text) or self.classifier.is_instruction_marker(text):
            return None

        style_name = paragraph.style.name if paragraph.style else ""
        if self._is_toc_entry(text, style_name) or self._looks_like_document_title(text):
            return None

        if self._is_preface_heading(text, style_name):
            return self._make_heading("", "前言", 1, style_name)

        appendix_match = self.appendix_pattern.match(text)
        if appendix_match and len(text) <= 80:
            number = f"附录{appendix_match.group(1).upper()}"
            title = appendix_match.group(2).strip(" （()")
            return self._make_heading(number, title, 1, style_name, required=False)

        level = self.classifier.heading_level(paragraph)
        number_match = self.heading_number_pattern.match(text)
        if not number_match:
            return None

        number, title = self.classifier.heading_parts(text)
        if not self._is_valid_heading_title(number, title, text):
            return None
        return self._make_heading(number, title, number.count(".") + 1 if number else max(1, level or 1), style_name)

    def _make_heading(self, number: str, title: str, level: int, style_name: str, required: bool = True) -> dict:
        return {
            "number": number,
            "title": title,
            "level": max(1, level or 1),
            "style_name": style_name,
            "body_style_name": "",
            "description": "",
            "required": required,
            "guidance_prompt": "",
            "template_blocks": [],
            "placeholders": [],
            "sub_chapters": [],
            "_guidance_parts": [],
        }

    def _is_body_start(self, paragraph, text: str) -> bool:
        return bool(self._is_preface_heading(text, paragraph.style.name if paragraph.style else "") or self._detect_heading(paragraph, ""))

    def _is_toc_start(self, text: str) -> bool:
        normalized = re.sub(r"\s+", "", text or "")
        return normalized in {"目次", "目录"}

    def _is_toc_entry(self, text: str, style_name: str = "") -> bool:
        style = (style_name or "").strip().lower()
        if style.startswith("toc"):
            return True
        if self.toc_page_pattern.match(text or ""):
            return True
        if "（章标题）" in text or "（条标题）" in text:
            return True
        return False

    def _is_preface_heading(self, text: str, style_name: str = "") -> bool:
        normalized = re.sub(r"\s+", "", text or "")
        if normalized == "前言":
            return True
        return "前言" in (style_name or "") and len(normalized) <= 4

    def _looks_like_document_title(self, text: str) -> bool:
        clean = re.sub(r"\s+", "", text or "")
        if len(clean) > 40 or self.heading_number_pattern.match(clean):
            return False
        return bool(self.document_title_pattern.search(clean))

    def _is_non_chapter_numbered_text(self, paragraph, text: str) -> bool:
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name == "List Paragraph":
            return True
        if self.figure_table_caption_pattern.match(text):
            return True
        if text.startswith(("式中", "注：", "注:", "ID号", "需求内容", "需求来源", "需求类型", "类型：")):
            return True
        return False

    def _is_valid_heading_title(self, number: str, title: str, text: str) -> bool:
        title = (title or "").strip()
        if not number or not title or len(text) > 100:
            return False
        if len(title) > 55:
            return False
        if self.placeholder_pattern.fullmatch(title.replace("（章标题）", "").replace("（条标题）", "").strip()):
            return False
        if title.endswith(("。", "；", ";")) and len(title) > 24:
            return False
        if re.search(r"[:：]", title) and not re.search(r"(目标|目的|范围|依据|方法|结果|建议|要求|说明)$", title):
            return False
        return True

    def _fallback_headings(self, doc: Document) -> list[dict]:
        result = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                result.append({
                    "number": str(len(result) + 1),
                    "title": text[:60],
                    "level": 1,
                    "style_name": paragraph.style.name if paragraph.style else "",
                    "body_style_name": paragraph.style.name if paragraph.style else "",
                    "description": "",
                    "required": True,
                    "guidance_prompt": "",
                    "template_blocks": [],
                    "placeholders": [],
                    "sub_chapters": [],
                })
            if len(result) >= 8:
                break
        return result

    def _build_tree(self, headings: list[dict]) -> list[dict]:
        roots = []
        stack = []
        for item in headings:
            chapter = {
                "number": item.get("number", ""),
                "title": item.get("title", ""),
                "description": item.get("description", ""),
                "required": item.get("required", True),
                "guidance_prompt": item.get("guidance_prompt", "") or self._make_guidance(item),
                "style_name": item.get("style_name", ""),
                "body_style_name": item.get("body_style_name", ""),
                "template_blocks": item.get("template_blocks", []),
                "placeholders": sorted(set(item.get("placeholders", []))),
                "generation_strategy": item.get("generation_strategy", "") or self._infer_generation_strategy(item),
                "sub_chapters": [],
            }
            level = int(item.get("level") or 1)
            while stack and stack[-1]["level"] >= level:
                stack.pop()
            if stack:
                stack[-1]["chapter"]["sub_chapters"].append(chapter)
            else:
                roots.append(chapter)
            stack.append({"level": level, "chapter": chapter})
        return roots

    def _make_guidance(self, item: dict) -> str:
        parts = [part.strip() for part in item.get("_guidance_parts", []) if part.strip()]
        if not parts:
            return item.get("description", "")
        text = "\n".join(parts)
        return text[:2000]

    def _infer_generation_strategy(self, item: dict) -> str:
        blocks = item.get("template_blocks", []) or []
        block_types = {block.get("type") for block in blocks}
        if "template_table" in block_types and not item.get("placeholders"):
            return "table_fill"
        if item.get("placeholders"):
            return "smart_generate"
        if {"instruction", "example", "instruction_table", "example_table"} & block_types:
            return "smart_generate"
        return "placeholder_replace"

    def _paragraph_block(self, paragraph, current_context: str = "") -> dict[str, Any]:
        text = paragraph.text.strip()
        italic_texts = [run.text.strip() for run in paragraph.runs if run.italic and run.text.strip()]
        placeholders = self.placeholder_pattern.findall(text)
        placeholders.extend(italic_texts)
        block_type = self._paragraph_block_type(paragraph, current_context)
        list_item = self._parse_template_list_item(text) if block_type == "template_text" else None
        if list_item:
            return {
                "type": "template_list",
                "label": self._block_label("template_list"),
                "text": text,
                "style_name": paragraph.style.name if paragraph.style else "",
                "list_style": list_item["list_style"],
                "can_expand": list_item["can_expand"],
                "items": [list_item],
                "has_italic": bool(italic_texts),
                "italic_texts": italic_texts,
                "placeholders": sorted(set(placeholders)),
            }
        return {
            "type": block_type,
            "label": self._block_label(block_type),
            "text": text,
            "style_name": paragraph.style.name if paragraph.style else "",
            "has_italic": bool(italic_texts),
            "italic_texts": italic_texts,
            "placeholders": sorted(set(placeholders)),
        }

    def _paragraph_block_type(self, paragraph, current_context: str = "") -> str:
        return self.classifier.classify_paragraph(paragraph, current_context)

    def _is_empty_marker_block(self, block: dict) -> bool:
        text = str(block.get("text") or "").strip()
        if block.get("type") == "instruction" and self.classifier.is_instruction_marker(text):
            return bool(re.fullmatch(r"【?\s*(?:说明|注|注意|填写说明|编写说明|编制说明|生成说明|要求|填写要求|编写要求|表格填写要求)\s*】?\s*[:：]?", text))
        if block.get("type") == "example" and self.classifier.is_example_marker(text):
            return bool(re.fullmatch(r"【?\s*(?:举例|示例|例|参考示例|样例)\s*】?\s*[:：]?", text))
        return False

    def _append_paragraph_block(self, heading: dict, block: dict) -> dict:
        blocks = heading["template_blocks"]
        if block.get("type") == "template_list":
            previous = blocks[-1] if blocks else None
            if previous and previous.get("type") == "template_list" and (
                previous.get("list_style") == block.get("list_style") or block.get("list_style") == "expandable"
            ):
                previous.setdefault("items", []).extend(block.get("items") or [])
                previous["text"] = self._list_block_text(previous)
                previous["can_expand"] = bool(previous.get("can_expand") or block.get("can_expand"))
                previous["placeholders"] = sorted(set((previous.get("placeholders") or []) + (block.get("placeholders") or [])))
                return previous
            block["text"] = self._list_block_text(block)
        blocks.append(block)
        return block

    def _parse_template_list_item(self, text: str) -> dict | None:
        if self.expandable_marker_pattern.match(text or ""):
            return {"marker": "……", "text": "", "list_style": "expandable", "can_expand": True}
        match = self.list_item_pattern.match(text or "")
        if not match:
            return None
        marker = match.group("marker").strip()
        value = match.group("text").strip()
        if not value:
            return None
        return {
            "marker": marker,
            "text": value,
            "list_style": self._list_style(marker),
            "can_expand": False,
        }

    def _list_style(self, marker: str) -> str:
        if re.match(r"^[a-zA-Z][\)）\.、]$", marker or ""):
            return "lower_alpha_cn" if "）" in marker or "、" in marker else "lower_alpha"
        if re.match(r"^[（(]\d+[）)]$", marker or ""):
            return "number_parentheses"
        if re.match(r"^\d+[\)）\.、]$", marker or ""):
            return "number_cn" if "）" in marker or "、" in marker else "number_dot"
        return "cn_number"

    def _list_block_text(self, block: dict) -> str:
        lines = []
        for item in block.get("items") or []:
            marker = item.get("marker") or ""
            text = item.get("text") or ""
            lines.append(f"{marker}{text}" if text else marker)
        return "\n".join(lines)

    def _table_block(self, table, current_context: str = "") -> dict[str, Any]:
        rows = []
        placeholders = []
        for row in table.rows:
            row_values = []
            for cell in row.cells:
                text = "\n".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                row_values.append(text)
                placeholders.extend(self.placeholder_pattern.findall(text))
            rows.append(row_values)
        if current_context in {"example", "example_marker"}:
            block_type = "example_table"
        elif current_context in {"instruction", "instruction_marker"}:
            block_type = "instruction_table"
        else:
            block_type = "template_table"
        return {
            "type": block_type,
            "label": self._block_label(block_type),
            "rows": rows,
            "text": self._table_text({"rows": rows}),
            "placeholders": sorted(set(placeholders)),
        }

    def _table_text(self, block: dict) -> str:
        return "\n".join(" | ".join(cell for cell in row if cell) for row in block.get("rows", []))

    def _next_context(self, block: dict, current_context: str = "") -> str:
        return self.classifier.next_context(block["type"], block.get("text", ""), current_context)

    def _block_label(self, block_type: str) -> str:
        return {
            "template_text": "模板文字",
            "template_list": "模板列表",
            "template_table": "模板表格",
            "instruction": "说明",
            "instruction_table": "说明表格",
            "example": "举例",
            "example_table": "举例表格",
        }.get(block_type, "模板内容")

    def _iter_block_items(self, parent):
        if isinstance(parent, DocumentType):
            parent_elm = parent.element.body
        else:
            parent_elm = parent._tc
        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)
