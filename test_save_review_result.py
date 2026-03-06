#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
测试保存审查结果为文档的功能（独立脚本）
"""

import os
import sys
import datetime
from docx import Document

# 添加项目根目录到Python路径
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

# 模拟审查结果
review_results = {
    'format_check': [
        '第4行：标题符号后缺少空格',
        '第7行：标题符号后缺少空格',
        '第10行：标题符号后缺少空格',
        '第13行：标题符号后缺少空格',
        '第16行：标题符号后缺少空格',
        '第19行：标题符号后缺少空格',
        '第24行：标题符号后缺少空格',
        '第20行：数字与单位之间缺少空格',
        '第21行：数字与单位之间缺少空格',
        '第22行：数字与单位之间缺少空格',
        '第27行：数字与单位之间缺少空格'
    ],
    'content_check': [
        '文档中没有引用'
    ],
    'terminology_check': [],
    'compliance_check': [
        '缺少合规性声明：本系统符合相关标准要求'
    ],
    'language_check': [
        '拼写错误：控制系统 应该为 控制系统'
    ]
}

# 原始文档路径
original_file_path = os.path.abspath('test_document.txt')

# 生成默认保存路径
original_name = os.path.splitext(os.path.basename(original_file_path))[0]
timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
# 默认保存到results文件夹
project_root = os.path.dirname(os.path.abspath(__file__))
default_save_path = os.path.join(project_root, 'results', f'{original_name}_审查结果_{timestamp}.docx')

# 确保保存目录存在
result_dir = os.path.dirname(default_save_path)
if not os.path.exists(result_dir):
    os.makedirs(result_dir)

# 创建Word文档
doc = Document()

# 添加标题
doc.add_heading('文档审查结果', 0)

# 添加基本信息
doc.add_heading('基本信息', level=1)
doc.add_paragraph(f'审查时间: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
doc.add_paragraph(f'原始文档: {original_file_path}')

# 添加审查结果
doc.add_heading('审查结果', level=1)

# 格式检查结果
if review_results.get('format_check', []):
    doc.add_heading('1. 格式规范检查', level=2)
    for issue in review_results['format_check']:
        doc.add_paragraph(f'- {issue}', style='List Bullet')
else:
    doc.add_heading('1. 格式规范检查', level=2)
    doc.add_paragraph('✓ 无问题', style='List Bullet')

# 内容检查结果
if review_results.get('content_check', []):
    doc.add_heading('2. 内容完整性检查', level=2)
    for issue in review_results['content_check']:
        doc.add_paragraph(f'- {issue}', style='List Bullet')
else:
    doc.add_heading('2. 内容完整性检查', level=2)
    doc.add_paragraph('✓ 无问题', style='List Bullet')

# 术语检查结果
if review_results.get('terminology_check', []):
    doc.add_heading('3. 术语一致性检查', level=2)
    for issue in review_results['terminology_check']:
        doc.add_paragraph(f'- {issue}', style='List Bullet')
else:
    doc.add_heading('3. 术语一致性检查', level=2)
    doc.add_paragraph('✓ 无问题', style='List Bullet')

# 合规性检查结果
if review_results.get('compliance_check', []):
    doc.add_heading('4. 合规性验证', level=2)
    for issue in review_results['compliance_check']:
        doc.add_paragraph(f'- {issue}', style='List Bullet')
else:
    doc.add_heading('4. 合规性验证', level=2)
    doc.add_paragraph('✓ 无问题', style='List Bullet')

# 语言检查结果
if review_results.get('language_check', []):
    doc.add_heading('5. 语言风格检查', level=2)
    for issue in review_results['language_check']:
        doc.add_paragraph(f'- {issue}', style='List Bullet')
else:
    doc.add_heading('5. 语言风格检查', level=2)
    doc.add_paragraph('✓ 无问题', style='List Bullet')

# 添加总结
doc.add_heading('总结', level=1)
total_issues = sum(len(issues) for issues in review_results.values())
if total_issues == 0:
    doc.add_paragraph('✓ 文档审查通过，未发现问题。')
else:
    doc.add_paragraph(f'⚠️ 文档审查发现 {total_issues} 个问题，建议根据上述结果进行修改。')

# 保存文档
doc.save(default_save_path)

print(f'审查结果文档已保存到: {default_save_path}')
print('您可以打开该文档查看详细的审查结果。')
