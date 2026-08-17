from docx import Document

from web.routes.generate import (
    _apply_docx_revision_operation,
    _extract_docx_revision_context,
)


def _texts(doc):
    return [p.text for p in doc.paragraphs if p.text.strip()]


def test_revision_context_selects_relevant_late_section_for_long_doc():
    doc = Document()
    doc.add_heading("1 概述", level=1)
    for index in range(80):
        doc.add_paragraph(f"普通段落 {index}，用于拉长文档。")
    doc.add_heading("9 电源输入特性", level=1)
    doc.add_paragraph("电源输入应满足 28VDC，纹波不大于 1V。")

    context = _extract_docx_revision_context(doc, query="修改电源输入纹波要求", max_chars=1200)

    assert "9 电源输入特性" in context["excerpt"]
    assert "纹波不大于 1V" in context["excerpt"]


def test_revision_replace_uses_target_paragraph_index():
    doc = Document()
    doc.add_heading("2 试验条件", level=1)
    doc.add_paragraph("温度范围为 -40℃～+60℃。")
    target_index = next(i for i, p in enumerate(doc.paragraphs) if "温度范围" in p.text)

    applied = _apply_docx_revision_operation(doc, {
        "operation": "replace",
        "target_heading": "2 试验条件",
        "target_paragraph_index": target_index,
        "new_text": "温度范围为 -55℃～+70℃。",
    })

    assert applied
    assert "温度范围为 -55℃～+70℃。" in _texts(doc)
    assert "温度范围为 -40℃～+60℃。" not in _texts(doc)


def test_revision_replace_can_replace_adjacent_paragraph_range():
    doc = Document()
    doc.add_heading("3 合格判据", level=1)
    doc.add_paragraph("试验中通信应保持正常。")
    doc.add_paragraph("试验后外观应无损伤。")

    applied = _apply_docx_revision_operation(doc, {
        "operation": "replace",
        "target_heading": "3 合格判据",
        "multi_old_text_hints": [
            "试验中通信应保持正常。",
            "试验后外观应无损伤。",
        ],
        "new_text": "试验中和试验后，模块应无物理损伤，通信正常，精度漂移不超过满量程的±0.2%。",
    })

    text = "\n".join(_texts(doc))
    assert applied
    assert "精度漂移不超过满量程的±0.2%" in text
    assert "试验后外观应无损伤" not in text


def test_revision_replace_uses_target_index_when_heading_is_inexact():
    doc = Document()
    doc.add_heading("1 概述", level=1)
    doc.add_paragraph("这里是概述内容。")
    doc.add_heading("9 电源输入特性", level=1)
    doc.add_paragraph("电源纹波不大于 1V。")
    target_index = next(i for i, p in enumerate(doc.paragraphs) if "电源纹波" in p.text)

    applied = _apply_docx_revision_operation(doc, {
        "operation": "replace",
        "target_heading": "电源章节",
        "target_paragraph_index": target_index,
        "new_text": "电源纹波不大于 0.5V。",
    })

    text = "\n".join(_texts(doc))
    assert applied
    assert "电源纹波不大于 0.5V。" in text
    assert "这里是概述内容。" in text
