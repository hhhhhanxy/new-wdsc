from docx import Document

from core.executor import ReviewExecutor, ReviewMode
from models.document import ContentType, DocumentSection, ParsedDocument
from parsers.docx_parser import DocxParser
from rules.base_rule import Rule, RuleCategory, RuleRegistry, RuleResult, RuleSeverity, ReviewType


class ExplodingLLMClient:
    def generate(self, prompt: str, system_prompt=None):
        raise RuntimeError("LLM service unavailable")


def test_docx_parser_merges_paragraphs_into_review_chunks(tmp_path):
    path = tmp_path / "many_paragraphs.docx"
    doc = Document()
    doc.add_heading("1. 标题", level=1)
    for index in range(20):
        doc.add_paragraph(f"这是第 {index + 1} 个普通段落，用于验证审查分块。")
    doc.save(path)

    parsed = DocxParser(chunk_size=500).parse(str(path))

    assert len(parsed.sections) < 21
    assert parsed.metadata["total_sections"] == len(parsed.sections)
    assert any(section.metadata.get("source_count", 0) > 1 for section in parsed.sections)


def test_rule_only_review_does_not_call_llm():
    class RuleOnlyLLMClient:
        def generate(self, prompt: str, system_prompt=None):
            raise AssertionError("LLM should not be called for rule-only reviews")

    registry = RuleRegistry()
    registry.register(Rule(
        rule_id="rule_only",
        name="规则检查",
        description="仅规则引擎检查",
        category=RuleCategory.CUSTOM,
        severity=RuleSeverity.WARNING,
        review_type=ReviewType.RULE,
        check_func=lambda section, context: RuleResult(
            rule_id="rule_only",
            rule_name="规则检查",
            passed=True,
            severity=RuleSeverity.WARNING,
            message="ok",
            section_id=section.section_id,
        ),
    ))
    document = ParsedDocument(
        file_path="test.docx",
        title="测试文档",
        sections=[DocumentSection("s1", ContentType.PARAGRAPH, "测试内容")],
        raw_text="测试内容",
    )

    executor = ReviewExecutor(registry, llm_client=RuleOnlyLLMClient(), mode=ReviewMode.RULE_ONLY)
    result = executor.review_document(document)

    assert result.total_issues == 0


def test_llm_failure_is_recorded_without_interrupting_review():
    registry = RuleRegistry()
    registry.register(Rule(
        rule_id="llm_rule",
        name="LLM规则",
        description="需要 LLM 审查",
        category=RuleCategory.CUSTOM,
        severity=RuleSeverity.WARNING,
        review_type=ReviewType.LLM,
    ))
    document = ParsedDocument(
        file_path="test.docx",
        title="测试文档",
        sections=[DocumentSection("s1", ContentType.PARAGRAPH, "测试内容")],
        raw_text="测试内容",
    )

    executor = ReviewExecutor(registry, llm_client=ExplodingLLMClient(), mode=ReviewMode.BOTH)
    result = executor.review_document(document)

    assert result.total_issues == 1
    assert result.section_results[0].rule_results[0].rule_id == "llm_rule"
    assert "LLM审查调用失败" in result.section_results[0].rule_results[0].message
