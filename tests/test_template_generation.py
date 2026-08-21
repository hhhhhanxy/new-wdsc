import json

from docx import Document


class FakeLLMResponse:
    def __init__(self, content):
        self.content = content


class FakeLLMClient:
    def __init__(self):
        self.prompts = []

    def generate(self, prompt, system_prompt=None):
        self.prompts.append(prompt)
        return FakeLLMResponse("本章节内容基于用户输入生成。")


def test_template_manager_serializes_templates():
    from templates.template_manager import TemplateManager

    templates = TemplateManager().list_template_dicts()

    assert templates
    assert any(t["id"] == "technical_specification" for t in templates)
    assert all("chapters" in t for t in templates)


def test_template_manager_persists_custom_templates(tmp_path):
    from templates.template_manager import TemplateManager

    store_path = tmp_path / "generation_templates.json"
    manager = TemplateManager(custom_store_path=str(store_path))
    template = manager.create_template(
        name="用户试验模板",
        description="用户上传模板",
        chapters=[{"number": "1", "title": "试验目的", "sub_chapters": []}],
    )

    reloaded = TemplateManager(custom_store_path=str(store_path))
    loaded = reloaded.get_template(template.template_id)

    assert loaded is not None
    assert loaded.name == "用户试验模板"
    assert loaded.doc_type is None
    assert loaded.source_type == "uploaded_docx"


def test_template_manager_persists_chapter_generation_strategy(tmp_path):
    from templates.template_manager import TemplateManager

    store_path = tmp_path / "generation_templates.json"
    manager = TemplateManager(custom_store_path=str(store_path))
    template = manager.create_template(
        name="策略模板",
        description="",
        chapters=[{
            "number": "1",
            "title": "固定章节",
            "generation_strategy": "fixed_keep",
            "sub_chapters": [],
        }],
    )

    reloaded = TemplateManager(custom_store_path=str(store_path)).get_template(template.template_id)

    assert reloaded.chapters[0].generation_strategy == "fixed_keep"
    assert TemplateManager(custom_store_path=str(store_path)).serialize_template(reloaded)["chapters"][0]["generation_strategy"] == "fixed_keep"


def test_template_manager_persists_dynamic_input_fields(tmp_path):
    from templates.template_manager import TemplateManager

    store_path = tmp_path / "generation_templates.json"
    manager = TemplateManager(custom_store_path=str(store_path))
    template = manager.create_template(
        name="动态字段模板",
        description="",
        chapters=[{"number": "1", "title": "范围", "sub_chapters": []}],
        input_fields=[{
            "key": "rated_voltage",
            "label": "额定电压",
            "type": "number",
            "required": True,
            "default_value": "28",
            "example": "28 V",
            "chapter_keys": ["1::范围"],
            "placeholder_tokens": ["VVVV"],
        }],
    )

    reloaded = TemplateManager(custom_store_path=str(store_path)).get_template(template.template_id)
    field = reloaded.input_fields[0]

    assert field.key == "rated_voltage"
    assert field.label == "额定电压"
    assert field.required is True
    assert field.chapter_keys == ["1::范围"]
    assert field.placeholder_tokens == ["VVVV"]


def test_template_manager_can_hide_builtin_templates(tmp_path):
    from templates.template_manager import TemplateManager

    store_path = tmp_path / "generation_templates.json"
    manager = TemplateManager(custom_store_path=str(store_path))

    assert manager.get_template("requirements") is not None
    assert manager.delete_template("requirements") is True

    reloaded = TemplateManager(custom_store_path=str(store_path))
    assert reloaded.get_template("requirements") is None
    assert all(t["id"] != "requirements" for t in reloaded.list_template_dicts())


def test_template_docx_generator_rejects_template_without_source_docx(tmp_path):
    from generators.base_generator import GeneratorFactory

    output_path = tmp_path / "generated.docx"
    generator = GeneratorFactory.create("template_docx")

    try:
        generator.generate(
            title="某型作动系统技术说明书",
            params={"template_id": "technical_specification", "inputs": {"product_name": "某型电动作动器"}},
            llm_client=FakeLLMClient(),
            output_path=str(output_path),
        )
    except ValueError as exc:
        assert "缺少原始 DOCX" in str(exc)
    else:
        raise AssertionError("无原始 DOCX 的模板不应允许生成")


def test_template_docx_generator_reports_fill_progress(tmp_path, monkeypatch):
    from generators.base_generator import GeneratorFactory
    from templates.template_manager import ChapterTemplate, DocumentTemplate

    template_path = tmp_path / "progress_template.docx"
    doc = Document()
    doc.add_heading("1 范围", level=1)
    doc.add_paragraph("NNNNN 模板正文")
    doc.save(template_path)
    output_path = tmp_path / "generated.docx"
    progress_events = []
    template = DocumentTemplate(
        template_id="progress_template",
        name="进度模板",
        description="",
        source_type="uploaded_docx",
        metadata={"source_path": str(template_path)},
        chapters=[ChapterTemplate(number="1", title="范围")],
    )

    class FakeTemplateManager:
        def get_template(self, template_id):
            return template if template_id == "progress_template" else None

    monkeypatch.setattr("templates.template_manager.TemplateManager", FakeTemplateManager)

    GeneratorFactory.create("template_docx").generate(
        title="某型作动系统需求文档",
        params={
            "template_id": "progress_template",
            "inputs": {"product_name": "某型电动作动器"},
        },
        llm_client=None,
        output_path=str(output_path),
        progress_callback=lambda current, total, chapter: progress_events.append(
            (current, total, chapter.title)
        ),
    )

    assert output_path.exists()
    assert progress_events
    assert progress_events[0][0] == 1
    assert progress_events[-1][0] == progress_events[-1][1]
    assert "某型电动作动器" in "\n".join(p.text for p in Document(output_path).paragraphs)


def test_template_docx_generator_uses_uploaded_template_styles(tmp_path, monkeypatch):
    from docx.enum.style import WD_STYLE_TYPE
    from generators.base_generator import GeneratorFactory
    from templates.template_manager import ChapterTemplate, DocumentTemplate

    template_path = tmp_path / "styled_template.docx"
    template_doc = Document()
    styles = template_doc.styles
    styles["Normal"].font.name = "Arial"
    custom_heading = styles.add_style("CustomHeadingOne", WD_STYLE_TYPE.PARAGRAPH)
    custom_heading.base_style = styles["Heading 1"]
    body_style = styles.add_style("CustomBodyText", WD_STYLE_TYPE.PARAGRAPH)
    body_style.base_style = styles["Normal"]
    template_doc.add_paragraph("1 范围", style="CustomHeadingOne")
    template_doc.add_paragraph("模板正文占位", style="CustomBodyText")
    template_doc.save(template_path)

    template = DocumentTemplate(
        template_id="styled_template",
        name="带格式模板",
        description="用于验证样式继承",
        source_type="uploaded_docx",
        metadata={"source_path": str(template_path)},
        chapters=[
            ChapterTemplate(number="1", title="范围", style_name="CustomHeadingOne", body_style_name="CustomBodyText"),
        ],
    )

    class FakeTemplateManager:
        def get_template(self, template_id):
            return template if template_id == "styled_template" else None

    monkeypatch.setattr("templates.template_manager.TemplateManager", FakeTemplateManager)

    output_path = tmp_path / "generated.docx"
    generator = GeneratorFactory.create("template_docx")
    generator.generate(
        title="样式继承验证",
        params={
            "template_id": "styled_template",
            "inputs": {"product_name": "某型电动作动器"},
        },
        llm_client=FakeLLMClient(),
        output_path=str(output_path),
    )

    generated_doc = Document(output_path)
    assert generated_doc.styles["Normal"].font.name == "Arial"
    assert generated_doc.paragraphs[0].style.name == "CustomHeadingOne"
    assert generated_doc.paragraphs[1].style.name == "CustomBodyText"
    assert "样式继承验证" not in "\n".join(p.text for p in generated_doc.paragraphs)


def test_dynamic_input_field_replaces_configured_placeholder(tmp_path, monkeypatch):
    from generators.base_generator import GeneratorFactory
    from templates.template_manager import ChapterTemplate, DocumentTemplate

    template_path = tmp_path / "dynamic_placeholder.docx"
    doc = Document()
    doc.add_heading("1 参数", level=1)
    doc.add_paragraph("额定电压为 VVVV。")
    doc.save(template_path)
    template = DocumentTemplate(
        template_id="dynamic_placeholder",
        name="动态占位符模板",
        description="",
        source_type="uploaded_docx",
        metadata={"source_path": str(template_path)},
        chapters=[
            ChapterTemplate(
                number="1",
                title="参数",
                generation_strategy="placeholder_replace",
                placeholders=["VVVV"],
            ),
        ],
    )

    class FakeTemplateManager:
        def get_template(self, template_id):
            return template if template_id == "dynamic_placeholder" else None

    monkeypatch.setattr("templates.template_manager.TemplateManager", FakeTemplateManager)
    output_path = tmp_path / "dynamic_placeholder_generated.docx"
    GeneratorFactory.create("template_docx").generate(
        title="动态字段替换",
        params={
            "template_id": "dynamic_placeholder",
            "inputs": {
                "product_name": "某产品",
                "generation_mode": "template_fill",
                "dynamic_fields": {"rated_voltage": "28 V"},
                "dynamic_field_definitions": [{
                    "key": "rated_voltage",
                    "label": "额定电压",
                    "placeholder_tokens": ["VVVV"],
                }],
            },
        },
        llm_client=None,
        output_path=str(output_path),
    )

    text = "\n".join(p.text for p in Document(output_path).paragraphs)
    assert "额定电压为 28 V" in text
    assert "VVVV" not in text


def test_template_docx_generator_fills_source_docx_and_removes_red_instructions(tmp_path, monkeypatch):
    from docx.shared import RGBColor
    from generators.base_generator import GeneratorFactory
    from templates.template_manager import ChapterTemplate, DocumentTemplate

    template_path = tmp_path / "fill_template.docx"
    template_doc = Document()
    template_doc.add_heading("1 目的", level=1)
    p1 = template_doc.add_paragraph()
    p1.add_run("本试验大纲规定了 XXX 项目 NNNNN 的鉴定试验程序。")
    p2 = template_doc.add_paragraph()
    p2.add_run("模板中")
    italic = p2.add_run("X 项目")
    italic.italic = True
    p2.add_run("按需修改。")
    red = template_doc.add_paragraph()
    red_run = red.add_run("【说明】红色字体说明在正式文件中删除。")
    red_run.font.color.rgb = RGBColor(255, 0, 0)
    template_doc.save(template_path)

    template = DocumentTemplate(
        template_id="fill_template",
        name="填充模板",
        description="用于验证模板原文填充",
        source_type="uploaded_docx",
        metadata={"source_path": str(template_path)},
        chapters=[
            ChapterTemplate(
                number="1",
                title="目的",
                guidance_prompt="红色字体删除，NNNNN 替换为设备名称，斜体项目按需修改。",
                template_blocks=[
                    {"type": "template_text", "text": "本试验大纲规定了 XXX 项目 NNNNN 的鉴定试验程序。"},
                    {"type": "instruction", "text": "【说明】红色字体说明在正式文件中删除。"},
                ],
                placeholders=["XXX", "NNNNN", "X 项目"],
            ),
        ],
    )

    class FakeTemplateManager:
        def get_template(self, template_id):
            return template if template_id == "fill_template" else None

    monkeypatch.setattr("templates.template_manager.TemplateManager", FakeTemplateManager)

    output_path = tmp_path / "generated.docx"
    generator = GeneratorFactory.create("template_docx")
    generator.generate(
        title="正式试验大纲",
        params={
            "template_id": "fill_template",
            "inputs": {
                "product_name": "某型电动作动器",
                "project_name": "某型电动作动器鉴定试验",
            },
        },
        llm_client=FakeLLMClient(),
        output_path=str(output_path),
    )

    generated_doc = Document(output_path)
    text = "\n".join(p.text for p in generated_doc.paragraphs)
    assert "红色字体说明" not in text
    assert "NNNNN" not in text
    assert "XXX" not in text
    assert "某型电动作动器" in text
    assert "某型电动作动器鉴定试验" in text


def test_template_docx_generator_replaces_fixed_template_title(tmp_path, monkeypatch):
    from generators.base_generator import GeneratorFactory
    from templates.template_manager import ChapterTemplate, DocumentTemplate

    template_path = tmp_path / "title_template.docx"
    template_doc = Document()
    template_doc.add_heading("XX型飞机航电系统国产化升级项目", level=0)
    template_doc.add_heading("1 目的", level=1)
    template_doc.add_paragraph("NNNNN 的试验目的待补充。")
    template_doc.save(template_path)

    template = DocumentTemplate(
        template_id="title_template",
        name="标题替换模板",
        description="用于验证固定模板标题替换",
        source_type="uploaded_docx",
        metadata={"source_path": str(template_path)},
        chapters=[
            ChapterTemplate(number="1", title="目的", placeholders=["NNNNN", "待补充"]),
        ],
    )

    class FakeTemplateManager:
        def get_template(self, template_id):
            return template if template_id == "title_template" else None

    monkeypatch.setattr("templates.template_manager.TemplateManager", FakeTemplateManager)

    output_path = tmp_path / "generated.docx"
    GeneratorFactory.create("template_docx").generate(
        title="某型飞机大气数据模块环境适应性试验大纲",
        params={
            "template_id": "title_template",
            "inputs": {"product_name": "ADM-200型大气数据模块"},
        },
        llm_client=FakeLLMClient(),
        output_path=str(output_path),
    )

    text = "\n".join(p.text for p in Document(output_path).paragraphs)
    assert "某型飞机大气数据模块环境适应性试验大纲" in text
    assert "XX型飞机航电系统国产化升级项目" not in text


def test_template_docx_generator_fills_cover_header_footer_placeholders(tmp_path, monkeypatch):
    from generators.base_generator import GeneratorFactory
    from templates.template_manager import ChapterTemplate, DocumentTemplate

    template_path = tmp_path / "cover_template.docx"
    template_doc = Document()
    template_doc.add_paragraph("(密级)")
    template_doc.add_paragraph("（产品型号）（产品名称）")
    template_doc.add_paragraph("(文件名称)")
    template_doc.sections[0].header.paragraphs[0].text = "页眉：(文件名称)"
    template_doc.sections[0].footer.paragraphs[0].text = "页脚：（产品型号）（产品名称）"
    template_doc.add_heading("1 范围", level=1)
    template_doc.add_paragraph("NNNNN 适用于本文件。")
    template_doc.save(template_path)

    template = DocumentTemplate(
        template_id="cover_template",
        name="封面页眉页脚模板",
        description="",
        source_type="uploaded_docx",
        metadata={"source_path": str(template_path)},
        chapters=[ChapterTemplate(number="1", title="范围", placeholders=["NNNNN"])],
    )

    class FakeTemplateManager:
        def get_template(self, template_id):
            return template if template_id == "cover_template" else None

    monkeypatch.setattr("templates.template_manager.TemplateManager", FakeTemplateManager)

    params = {
        "template_id": "cover_template",
        "inputs": {
            "product_name": "ADM-200型大气数据模块",
            "product_model": "ADM-200",
        },
    }
    output_path = tmp_path / "cover_generated.docx"
    GeneratorFactory.create("template_docx").generate(
        title="某型飞机大气数据模块可靠性分配报告",
        params=params,
        llm_client=None,
        output_path=str(output_path),
    )

    generated = Document(output_path)
    body_text = "\n".join(p.text for p in generated.paragraphs)
    header_text = "\n".join(p.text for p in generated.sections[0].header.paragraphs)
    footer_text = "\n".join(p.text for p in generated.sections[0].footer.paragraphs)
    assert "某型飞机大气数据模块可靠性分配报告" in body_text
    assert "ADM-200ADM-200型大气数据模块" in body_text
    assert "某型飞机大气数据模块可靠性分配报告" in header_text
    assert "ADM-200ADM-200型大气数据模块" in footer_text
    assert any(item["label"] == "密级" for item in params["_generation_meta"]["pending_fields"])


def test_template_docx_generator_uses_llm_for_smart_generation(tmp_path, monkeypatch):
    from generators.base_generator import GeneratorFactory
    from templates.template_manager import ChapterTemplate, DocumentTemplate

    template_path = tmp_path / "smart_template.docx"
    template_doc = Document()
    template_doc.add_heading("1 范围", level=1)
    template_doc.add_paragraph("本章说明 NNNNN 的适用范围。")
    template_doc.save(template_path)

    template = DocumentTemplate(
        template_id="smart_template",
        name="智能生成模板",
        description="",
        source_type="uploaded_docx",
        metadata={"source_path": str(template_path)},
        chapters=[
            ChapterTemplate(
                number="1",
                title="范围",
                guidance_prompt="根据用户素材生成适用范围。",
                template_blocks=[{"type": "template_text", "text": "本章说明 NNNNN 的适用范围。"}],
                placeholders=["NNNNN"],
            ),
        ],
    )

    class FakeTemplateManager:
        def get_template(self, template_id):
            return template if template_id == "smart_template" else None

    class SmartLLM:
        def __init__(self):
            self.prompts = []

        def generate(self, prompt, system_prompt=None):
            self.prompts.append((prompt, system_prompt))
            return FakeLLMResponse("本文件适用于某型电动作动器的设计、验证和交付过程。")

    monkeypatch.setattr("templates.template_manager.TemplateManager", FakeTemplateManager)
    llm = SmartLLM()
    output_path = tmp_path / "generated.docx"
    params = {
        "template_id": "smart_template",
        "inputs": {
            "product_name": "某型电动作动器",
            "generation_mode": "smart",
            "generation_requirements": "正文正式、准确。",
        },
    }

    GeneratorFactory.create("template_docx").generate(
        title="智能生成测试",
        params=params,
        llm_client=llm,
        output_path=str(output_path),
    )

    text = "\n".join(p.text for p in Document(output_path).paragraphs)
    assert "本文件适用于某型电动作动器" in text
    assert llm.prompts
    assert params["_generation_meta"]["generated_sections"] == 1


def test_smart_generation_targets_placeholder_paragraph_instead_of_first_body(tmp_path, monkeypatch):
    from generators.base_generator import GeneratorFactory
    from templates.template_manager import ChapterTemplate, DocumentTemplate

    template_path = tmp_path / "target_template.docx"
    template_doc = Document()
    template_doc.add_heading("1 范围", level=1)
    template_doc.add_paragraph("本段为固定说明，应保留。")
    template_doc.add_paragraph("NNNNN 的适用范围待补充。")
    template_doc.save(template_path)

    template = DocumentTemplate(
        template_id="target_template",
        name="目标段落模板",
        description="",
        source_type="uploaded_docx",
        metadata={"source_path": str(template_path)},
        chapters=[
            ChapterTemplate(
                number="1",
                title="范围",
                guidance_prompt="生成适用范围。",
                generation_strategy="smart_generate",
                placeholders=["NNNNN"],
                template_blocks=[{"type": "template_text", "text": "NNNNN 的适用范围待补充。"}],
            ),
        ],
    )

    class FakeTemplateManager:
        def get_template(self, template_id):
            return template if template_id == "target_template" else None

    class SmartLLM:
        def generate(self, prompt, system_prompt=None):
            return FakeLLMResponse("某型产品适用于试验验证阶段。")

    monkeypatch.setattr("templates.template_manager.TemplateManager", FakeTemplateManager)
    output_path = tmp_path / "generated.docx"
    params = {
        "template_id": "target_template",
        "inputs": {"product_name": "某型产品", "generation_mode": "smart"},
    }

    GeneratorFactory.create("template_docx").generate(
        title="目标段落测试",
        params=params,
        llm_client=SmartLLM(),
        output_path=str(output_path),
    )

    paragraphs = [p.text for p in Document(output_path).paragraphs if p.text.strip()]
    assert "本段为固定说明，应保留。" in paragraphs
    assert "某型产品适用于试验验证阶段。" in paragraphs
    assert paragraphs.index("本段为固定说明，应保留。") < paragraphs.index("某型产品适用于试验验证阶段。")
    section = params["_generation_meta"]["sections"][0]
    assert section["apply_target"] == "占位符/斜体/待补充段落"
    assert "完整" not in section["prompt_summary"]
    assert "【用户输入素材】" in section["prompt"]
    assert section["response"] == "某型产品适用于试验验证阶段。"


def test_smart_generation_fills_multiple_targets_in_one_chapter(tmp_path, monkeypatch):
    from generators.base_generator import GeneratorFactory
    from templates.template_manager import ChapterTemplate, DocumentTemplate

    template_path = tmp_path / "multiple_targets.docx"
    template_doc = Document()
    template_doc.add_heading("1 范围", level=1)
    template_doc.add_paragraph("本段为固定正文，应保持不变。")
    template_doc.add_paragraph("NNNNN 的适用对象待补充。")
    template_doc.add_paragraph("NNNNN 的使用限制待补充。")
    template_doc.save(template_path)

    template = DocumentTemplate(
        template_id="multiple_targets",
        name="多目标模板",
        description="",
        source_type="uploaded_docx",
        metadata={"source_path": str(template_path)},
        chapters=[
            ChapterTemplate(
                number="1",
                title="范围",
                guidance_prompt="分别填写适用对象和使用限制。",
                generation_strategy="smart_generate",
                placeholders=["NNNNN", "待补充"],
            ),
        ],
    )

    class FakeTemplateManager:
        def get_template(self, template_id):
            return template if template_id == "multiple_targets" else None

    class SequenceLLM:
        def __init__(self):
            self.responses = iter([
                "本文件适用于某型产品的设计验证。",
                "本文件不适用于未经批准的改型产品。",
            ])
            self.prompts = []

        def generate(self, prompt, system_prompt=None):
            self.prompts.append(prompt)
            return FakeLLMResponse(next(self.responses))

    monkeypatch.setattr("templates.template_manager.TemplateManager", FakeTemplateManager)
    llm = SequenceLLM()
    output_path = tmp_path / "multiple_targets_generated.docx"
    params = {
        "template_id": "multiple_targets",
        "inputs": {"product_name": "某型产品", "generation_mode": "smart"},
    }

    GeneratorFactory.create("template_docx").generate(
        title="多目标生成测试",
        params=params,
        llm_client=llm,
        output_path=str(output_path),
    )

    paragraphs = [p.text for p in Document(output_path).paragraphs if p.text.strip()]
    assert "本段为固定正文，应保持不变。" in paragraphs
    assert "本文件适用于某型产品的设计验证。" in paragraphs
    assert "本文件不适用于未经批准的改型产品。" in paragraphs
    assert len(llm.prompts) == 2
    assert "第 1/2 处" in llm.prompts[0]
    assert "第 2/2 处" in llm.prompts[1]
    section = params["_generation_meta"]["sections"][0]
    assert section["target_count"] == 2
    assert section["applied_count"] == 2
    assert len(section["requests"]) == 2


def test_fixed_keep_strategy_does_not_replace_placeholders(tmp_path, monkeypatch):
    from generators.base_generator import GeneratorFactory
    from templates.template_manager import ChapterTemplate, DocumentTemplate

    template_path = tmp_path / "fixed_template.docx"
    template_doc = Document()
    template_doc.add_heading("1 固定", level=1)
    template_doc.add_paragraph("固定正文 NNNNN 不替换。")
    template_doc.save(template_path)

    template = DocumentTemplate(
        template_id="fixed_template",
        name="固定模板",
        description="",
        source_type="uploaded_docx",
        metadata={"source_path": str(template_path)},
        chapters=[
            ChapterTemplate(
                number="1",
                title="固定",
                generation_strategy="fixed_keep",
                placeholders=["NNNNN"],
            ),
        ],
    )

    class FakeTemplateManager:
        def get_template(self, template_id):
            return template if template_id == "fixed_template" else None

    monkeypatch.setattr("templates.template_manager.TemplateManager", FakeTemplateManager)
    output_path = tmp_path / "generated.docx"

    GeneratorFactory.create("template_docx").generate(
        title="固定测试",
        params={
            "template_id": "fixed_template",
            "inputs": {"product_name": "某产品", "generation_mode": "smart"},
        },
        llm_client=FakeLLMClient(),
        output_path=str(output_path),
    )

    text = "\n".join(p.text for p in Document(output_path).paragraphs)
    assert "NNNNN" in text
    assert "某产品" not in text


def test_table_fill_strategy_maps_structured_material_and_expands_rows(tmp_path, monkeypatch):
    from generators.base_generator import GeneratorFactory
    from templates.template_manager import ChapterTemplate, DocumentTemplate

    template_path = tmp_path / "generic_table.docx"
    template_doc = Document()
    template_doc.add_heading("3 技术参数", level=1)
    table = template_doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "参数"
    table.cell(0, 1).text = "要求"
    table.cell(1, 0).text = ""
    table.cell(1, 1).text = "待填写"
    template_doc.save(template_path)

    template = DocumentTemplate(
        template_id="generic_table",
        name="通用参数表模板",
        description="",
        source_type="uploaded_docx",
        metadata={"source_path": str(template_path)},
        chapters=[
            ChapterTemplate(
                number="3",
                title="技术参数",
                generation_strategy="table_fill",
            ),
        ],
    )

    class FakeTemplateManager:
        def get_template(self, template_id):
            return template if template_id == "generic_table" else None

    monkeypatch.setattr("templates.template_manager.TemplateManager", FakeTemplateManager)
    output_path = tmp_path / "generic_table_generated.docx"

    GeneratorFactory.create("template_docx").generate(
        title="参数表生成测试",
        params={
            "template_id": "generic_table",
            "inputs": {
                "product_name": "某型产品",
                "generation_mode": "template_fill",
                "technical_params": "额定电压：28 V\n最大电流：5 A\n工作温度：-55 ℃～70 ℃",
            },
        },
        llm_client=None,
        output_path=str(output_path),
    )

    generated_table = Document(output_path).tables[0]
    rows = [[cell.text for cell in row.cells] for row in generated_table.rows]
    assert rows[0] == ["参数", "要求"]
    assert rows[1] == ["额定电压", "28 V"]
    assert rows[2] == ["最大电流", "5 A"]
    assert rows[3] == ["工作温度", "-55 ℃～70 ℃"]


def test_template_docx_generator_keeps_template_table_and_removes_example_table(tmp_path, monkeypatch):
    from docx.shared import RGBColor
    from generators.base_generator import GeneratorFactory
    from templates.template_manager import ChapterTemplate, DocumentTemplate

    template_path = tmp_path / "table_template.docx"
    template_doc = Document()
    template_doc.add_heading("3 依据", level=1)
    template_doc.add_paragraph("表3-1：依据文件")
    table = template_doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "文件编号"
    table.cell(0, 1).text = "文件名"
    table.cell(0, 2).text = "发布日期"
    table.cell(1, 0).text = "ABCC-12-R2"
    table.cell(1, 1).text = "XXX标准"
    table.cell(1, 2).text = ""
    red = template_doc.add_paragraph()
    red_run = red.add_run("【举例】")
    red_run.font.color.rgb = RGBColor(255, 0, 0)
    red_title = template_doc.add_paragraph()
    red_title_run = red_title.add_run("表3-1：依据文件")
    red_title_run.font.color.rgb = RGBColor(255, 0, 0)
    example = template_doc.add_table(rows=2, cols=3)
    example.cell(0, 0).text = "文件编号"
    example.cell(0, 1).text = "文件名"
    example.cell(0, 2).text = "发布日期"
    example.cell(1, 0).text = "D121-XXXX"
    example.cell(1, 1).text = "常见设备厂商文件"
    example.cell(1, 2).text = ""
    template_doc.save(template_path)

    template = DocumentTemplate(
        template_id="table_template",
        name="表格模板",
        description="用于验证表格处理",
        source_type="uploaded_docx",
        metadata={"source_path": str(template_path)},
        chapters=[ChapterTemplate(number="3", title="依据")],
    )

    class FakeTemplateManager:
        def get_template(self, template_id):
            return template if template_id == "table_template" else None

    monkeypatch.setattr("templates.template_manager.TemplateManager", FakeTemplateManager)

    output_path = tmp_path / "generated.docx"
    GeneratorFactory.create("template_docx").generate(
        title="正式试验大纲",
        params={
            "template_id": "table_template",
            "inputs": {
                "project_name": "某型电动作动器鉴定试验",
                "references": "GJB150A-2009 军用装备实验室环境试验方法\nQ/ABC-001 某型电动作动器产品技术条件 2026",
            },
        },
        llm_client=None,
        output_path=str(output_path),
    )

    generated_doc = Document(output_path)
    all_text = "\n".join(p.text for p in generated_doc.paragraphs)
    table_text = "\n".join(
        cell.text
        for table in generated_doc.tables
        for row in table.rows
        for cell in row.cells
    )
    assert len(generated_doc.tables) == 1
    assert "【举例】" not in all_text
    assert "常见设备厂商文件" not in table_text
    assert "GJB150A-2009" in table_text
    assert "军用装备实验室环境试验方法" in table_text
    assert "Q/ABC-001" in table_text
    assert "某型电动作动器产品技术条件" in table_text


def test_template_docx_generator_removes_plain_instruction_and_example_blocks(tmp_path, monkeypatch):
    from generators.base_generator import GeneratorFactory
    from templates.template_manager import ChapterTemplate, DocumentTemplate

    template_path = tmp_path / "plain_guidance_template.docx"
    template_doc = Document()
    template_doc.add_heading("1 目的", level=1)
    template_doc.add_paragraph("本文件规定了 NNNNN 的试验目的。")
    template_doc.add_paragraph("【说明】正式文件中删除本段。")
    template_doc.add_paragraph("1. 文中 NNNNN 替换为产品名称。")
    template_doc.add_paragraph("举例：某型设备鉴定试验大纲。")
    template_doc.add_heading("2 范围", level=1)
    template_doc.add_paragraph("NNNNN 适用于航空机载作动系统。")
    template_doc.save(template_path)

    template = DocumentTemplate(
        template_id="plain_guidance_template",
        name="非红色说明模板",
        description="用于验证非红色说明和举例",
        source_type="uploaded_docx",
        metadata={"source_path": str(template_path)},
        chapters=[
            ChapterTemplate(number="1", title="目的"),
            ChapterTemplate(number="2", title="范围"),
        ],
    )

    class FakeTemplateManager:
        def get_template(self, template_id):
            return template if template_id == "plain_guidance_template" else None

    monkeypatch.setattr("templates.template_manager.TemplateManager", FakeTemplateManager)

    output_path = tmp_path / "generated.docx"
    GeneratorFactory.create("template_docx").generate(
        title="正式试验大纲",
        params={
            "template_id": "plain_guidance_template",
            "inputs": {"product_name": "某型电动作动器"},
        },
        llm_client=None,
        output_path=str(output_path),
    )

    generated_doc = Document(output_path)
    text = "\n".join(p.text for p in generated_doc.paragraphs)
    assert "正式文件中删除" not in text
    assert "替换为产品名称" not in text
    assert "某型设备鉴定试验大纲" not in text
    assert "2 范围" in text
    assert "某型电动作动器 适用于航空机载作动系统" in text


def test_generate_document_only_skips_review_and_creates_docx(tmp_path, monkeypatch):
    from core.pipeline import generate_document_only
    from templates.template_manager import ChapterTemplate, DocumentTemplate

    template_path = tmp_path / "pipeline_template.docx"
    doc = Document()
    doc.add_paragraph("NNNNN 正式模板正文")
    doc.save(template_path)
    template = DocumentTemplate(
        template_id="pipeline_template",
        name="流水线模板",
        description="",
        source_type="uploaded_docx",
        metadata={"source_path": str(template_path)},
        chapters=[ChapterTemplate(number="1", title="范围")],
    )

    class FakeTemplateManager:
        def get_template(self, template_id):
            return template if template_id == "pipeline_template" else None

    monkeypatch.setattr("templates.template_manager.TemplateManager", FakeTemplateManager)
    progress_events = []

    result = generate_document_only(
        doc_type="pipeline_template",
        title="需求文档:测试/下载",
        params={
            "template_id": "pipeline_template",
            "inputs": {"product_name": "某型电动作动器"},
            "generator": "template_docx",
        },
        llm_client=None,
        output_dir=str(tmp_path),
        progress_callback=lambda current, total, chapter: progress_events.append(current),
    )

    assert result.status == "generated"
    assert result.generated_path
    assert result.quality_result
    assert "checks" in result.quality_result
    assert "需求文档_测试_下载.docx" in result.generated_path
    assert result.review_result is None
    assert progress_events
    assert "某型电动作动器" in "\n".join(p.text for p in Document(result.generated_path).paragraphs)


def test_docx_template_parser_detects_heading_tree(tmp_path):
    from docx import Document
    from templates.docx_template_parser import DocxTemplateParser

    path = tmp_path / "template.docx"
    doc = Document()
    doc.add_heading("1 范围", level=1)
    doc.add_heading("1.1 适用范围", level=2)
    doc.add_heading("2 引用文件", level=1)
    doc.save(path)

    parsed = DocxTemplateParser().parse(str(path))

    assert parsed["name"] == "template"
    assert parsed["chapters"][0]["number"] == "1"
    assert parsed["chapters"][0]["title"] == "范围"
    assert parsed["chapters"][0]["sub_chapters"][0]["number"] == "1.1"
    assert parsed["chapters"][1]["title"] == "引用文件"


def test_docx_template_parser_extracts_guidance_and_styles(tmp_path):
    from docx.enum.style import WD_STYLE_TYPE
    from docx.shared import RGBColor
    from templates.docx_template_parser import DocxTemplateParser

    path = tmp_path / "template_guidance.docx"
    doc = Document()
    styles = doc.styles
    custom_body = styles.add_style("TemplateBody", WD_STYLE_TYPE.PARAGRAPH)
    custom_body.base_style = styles["Normal"]
    doc.add_heading("1 范围", level=1)
    doc.add_paragraph("本章应说明文档适用对象、边界和使用限制。", style="TemplateBody")
    red = doc.add_paragraph()
    red_run = red.add_run("【说明】红色字体在正式文件中删除。")
    red_run.font.color.rgb = RGBColor(255, 0, 0)
    doc.add_heading("2 引用文件", level=1)
    doc.save(path)

    parsed = DocxTemplateParser().parse(str(path))

    first = parsed["chapters"][0]
    assert first["style_name"] in ("Heading 1", "标题 1")
    assert first["body_style_name"] == ""
    assert first["template_blocks"][0]["type"] == "instruction"
    assert "红色字体" in first["guidance_prompt"]
    assert any(block["type"] == "instruction" for block in first["template_blocks"])


def test_docx_template_parser_extracts_example_and_table_blocks(tmp_path):
    from docx.shared import RGBColor
    from templates.docx_template_parser import DocxTemplateParser

    path = tmp_path / "template_table.docx"
    doc = Document()
    doc.add_heading("3 依据", level=1)
    doc.add_paragraph("表3-1：依据文件")
    table = doc.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "文件编号"
    table.cell(0, 1).text = "文件名"
    table.cell(0, 2).text = "发布日期"
    table.cell(1, 0).text = "ABCC-12-R2"
    table.cell(1, 1).text = "XXX标准"
    red = doc.add_paragraph()
    run = red.add_run("【举例】")
    run.font.color.rgb = RGBColor(255, 0, 0)
    example_table = doc.add_table(rows=1, cols=2)
    example_table.cell(0, 0).text = "D121-XXXX"
    example_table.cell(0, 1).text = "常见设备厂商文件"
    doc.save(path)

    parsed = DocxTemplateParser().parse(str(path))
    blocks = parsed["chapters"][0]["template_blocks"]

    assert any(block["type"] == "template_table" for block in blocks)
    assert any(block["type"] == "example" for block in blocks)
    assert any(block["type"] == "example_table" for block in blocks)
    assert "举例" in parsed["chapters"][0]["guidance_prompt"]
    assert "XXX" in parsed["chapters"][0]["placeholders"]


def test_docx_template_parser_classifies_plain_guidance_markers(tmp_path):
    from templates.docx_template_parser import DocxTemplateParser

    path = tmp_path / "plain_guidance.docx"
    doc = Document()
    doc.add_heading("1 目的", level=1)
    doc.add_paragraph("本试验大纲规定了 NNNNN 的鉴定试验程序。")
    doc.add_paragraph("说明：正式文档中删除该说明段。")
    doc.add_paragraph("1. 文中 NNNNN 应替换为受试设备名称。")
    doc.add_paragraph("示例：某型产品鉴定试验大纲。")
    doc.save(path)

    parsed = DocxTemplateParser().parse(str(path))
    blocks = parsed["chapters"][0]["template_blocks"]

    assert [block["type"] for block in blocks] == [
        "template_text",
        "instruction",
        "instruction",
        "example",
    ]
    assert "正式文档中删除" in parsed["chapters"][0]["guidance_prompt"]
    assert "某型产品鉴定试验大纲" in parsed["chapters"][0]["guidance_prompt"]


def test_docx_template_parser_groups_template_list_items(tmp_path):
    from templates.docx_template_parser import DocxTemplateParser

    path = tmp_path / "template_list.docx"
    doc = Document()
    doc.add_heading("4.1 产品主要功能", level=2)
    doc.add_paragraph("XXX产品用于……。")
    doc.add_paragraph("XXX产品具有以下功能：")
    doc.add_paragraph("a）XXXXXX；")
    doc.add_paragraph("b）XXXXXX；")
    doc.add_paragraph("c）XXXXXX；")
    doc.add_paragraph("……")
    doc.save(path)

    parsed = DocxTemplateParser().parse(str(path))
    blocks = parsed["chapters"][0]["template_blocks"]
    list_block = next(block for block in blocks if block["type"] == "template_list")

    assert list_block["label"] == "模板列表"
    assert list_block["list_style"] == "lower_alpha_cn"
    assert list_block["can_expand"] is True
    assert [item["marker"] for item in list_block["items"]] == ["a）", "b）", "c）", "……"]
    assert [item["text"] for item in list_block["items"][:3]] == ["XXXXXX；", "XXXXXX；", "XXXXXX；"]


def test_docx_template_parser_treats_preface_writing_notes_as_guidance(tmp_path):
    from templates.docx_template_parser import DocxTemplateParser

    path = tmp_path / "preface_guidance.docx"
    doc = Document()
    doc.add_paragraph("前言")
    doc.add_paragraph("XXXXXXXXXXXXXXXXXXXXXXXXXXXX。")
    doc.add_paragraph("前言一般应说明下列内容：")
    doc.add_paragraph("a）说明文件编制依据或背景；")
    doc.add_paragraph("b）说明文件废止或代替其它文件的情况；")
    doc.add_paragraph("注：前言属概述要素，可根据需要作适当剪裁。")
    doc.save(path)

    parsed = DocxTemplateParser().parse(str(path))
    blocks = parsed["chapters"][0]["template_blocks"]
    guidance = parsed["chapters"][0]["guidance_prompt"]

    assert blocks[0]["type"] == "template_text"
    assert all(block["type"] == "instruction" for block in blocks[1:])
    assert "前言一般应说明下列内容" in guidance
    assert "说明文件编制依据或背景" in guidance
    assert "适当剪裁" in guidance
    assert "说明：注：" not in guidance
    assert "注：前言属概述要素" in guidance


def test_docx_template_parser_drops_empty_guidance_markers(tmp_path):
    from templates.docx_template_parser import DocxTemplateParser

    path = tmp_path / "empty_markers.docx"
    doc = Document()
    doc.add_heading("1 范围", level=1)
    doc.add_paragraph("【说明】")
    doc.add_paragraph("本章应说明文件适用范围。")
    doc.add_paragraph("【示例】")
    doc.add_paragraph("某型产品适用于试验阶段。")
    doc.save(path)

    parsed = DocxTemplateParser().parse(str(path))
    blocks = parsed["chapters"][0]["template_blocks"]
    texts = [block["text"] for block in blocks]

    assert "【说明】" not in texts
    assert "【示例】" not in texts
    assert [block["type"] for block in blocks] == ["instruction", "example"]


def test_docx_template_parser_classifies_content_after_guidance_markers(tmp_path):
    from templates.docx_template_parser import DocxTemplateParser

    path = tmp_path / "guidance_marker_content.docx"
    doc = Document()
    doc.add_heading("1 范围", level=1)
    doc.add_paragraph("【说明】")
    doc.add_paragraph("本节简单介绍本文件编制的目的和适用范围。")
    doc.add_paragraph("【示例】")
    doc.add_paragraph("本文件描述了 XXX 产品的可靠性分配。")
    doc.save(path)

    parsed = DocxTemplateParser().parse(str(path))
    blocks = parsed["chapters"][0]["template_blocks"]

    assert [block["type"] for block in blocks] == ["instruction", "example"]
    assert blocks[0]["text"] == "本节简单介绍本文件编制的目的和适用范围。"
    assert blocks[1]["text"] == "本文件描述了 XXX 产品的可靠性分配。"


def test_docx_template_parser_keeps_numbered_examples_as_example_context(tmp_path):
    from templates.docx_template_parser import DocxTemplateParser

    path = tmp_path / "numbered_examples.docx"
    doc = Document()
    doc.add_heading("2 引用文件", level=1)
    doc.add_paragraph("【说明】")
    doc.add_paragraph("本节应列出文件引用的文献或资料。")
    doc.add_paragraph("应给出文件的编号、名称、版本等信息。")
    doc.add_paragraph("【示例1】")
    doc.add_paragraph("下列文件对于本文件的应用是必不可少的。")
    doc.add_paragraph("表2-1 引用文件")
    table = doc.add_table(rows=2, cols=4)
    table.rows[0].cells[0].text = "序号"
    table.rows[0].cells[1].text = "文件编号"
    table.rows[0].cells[2].text = "文件名"
    table.rows[0].cells[3].text = "备注"
    table.rows[1].cells[0].text = "1"
    table.rows[1].cells[1].text = "GJB 450"
    table.rows[1].cells[2].text = "装备可靠性工作通用要求"
    doc.add_paragraph("【示例2】")
    doc.add_paragraph("GJB 451 可靠性维修性保障性术语")
    doc.add_paragraph("GJB 813 可靠性模型的建立与可靠性预计")
    doc.save(path)

    parsed = DocxTemplateParser().parse(str(path))
    blocks = parsed["chapters"][0]["template_blocks"]
    types = [block["type"] for block in blocks]
    texts = [block.get("text", "") for block in blocks]

    assert "【示例1】" not in texts
    assert "【示例2】" not in texts
    assert types == [
        "instruction",
        "instruction",
        "example",
        "example_table_caption",
        "example_table",
        "example",
        "example",
    ]
    assert "表2-1 引用文件" in texts
    assert "GJB 813" in texts[-1]


def test_docx_template_parser_marks_example_table_captions_and_expandable_lists(tmp_path):
    from templates.docx_template_parser import DocxTemplateParser

    path = tmp_path / "abbreviations.docx"
    doc = Document()
    doc.add_heading("3.2 缩略语", level=2)
    doc.add_paragraph("【说明】")
    doc.add_paragraph("本节应列出文中引用的缩略语及全称。")
    doc.add_paragraph("【示例】")
    doc.add_paragraph("下列缩略语适用于本文件。")
    doc.add_paragraph("表3-2 缩略语及定义表")
    table = doc.add_table(rows=2, cols=3)
    table.rows[0].cells[0].text = "缩略语"
    table.rows[0].cells[1].text = "英文全称"
    table.rows[0].cells[2].text = "中文全称"
    table.rows[1].cells[0].text = "MTBF"
    table.rows[1].cells[1].text = "Mean Time Between Failures"
    table.rows[1].cells[2].text = "平均失效间隔时间"
    doc.add_paragraph("……")
    doc.save(path)

    parsed = DocxTemplateParser().parse(str(path))
    blocks = parsed["chapters"][0]["template_blocks"]
    types = [block["type"] for block in blocks]

    assert types == ["instruction", "example", "example_table_caption", "example_table", "example_list"]
    assert blocks[2]["label"] == "举例表题"
    assert blocks[-1]["can_expand"] is True
    assert "分类层级" in blocks[-1]["expansion_hint"]


def test_docx_template_parser_keeps_word_formula_blocks(tmp_path):
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    from templates.docx_template_parser import DocxTemplateParser

    path = tmp_path / "formula_template.docx"
    doc = Document()
    doc.add_heading("5.1 可靠性数学模型", level=2)
    doc.add_paragraph("其数学模型为：")
    formula_paragraph = doc.add_paragraph()
    formula_paragraph._p.append(parse_xml(
        f'<m:oMath {nsdecls("m")}>'
        '<m:r><m:t>λs=Σniλi</m:t></m:r>'
        '</m:oMath>'
    ))
    doc.add_paragraph("式中：λs为系统故障率（1/h）；")
    doc.save(path)

    parsed = DocxTemplateParser().parse(str(path))
    blocks = parsed["chapters"][0]["template_blocks"]
    formula = next(block for block in blocks if block["type"] == "formula")

    assert formula["label"] == "模板公式"
    assert formula["text"] == "λs=Σniλi"


def test_docx_template_parser_skips_front_matter_toc_and_captions(tmp_path):
    from docx.enum.style import WD_STYLE_TYPE
    from templates.docx_template_parser import DocxTemplateParser

    path = tmp_path / "standard_template.docx"
    doc = Document()
    for style_name in ("TOC 1", "TOC 2", "TOC 3"):
        if style_name not in doc.styles:
            doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    doc.add_paragraph("(密级)")
    doc.add_paragraph("更改记录")
    doc.add_paragraph("目  次")
    doc.add_paragraph("前言\tII", style="TOC 1")
    doc.add_paragraph("1 XXXXXXXX（章标题）\t1", style="TOC 2")
    doc.add_paragraph("2 XXXXXXXX（章标题）\t1", style="TOC 2")
    doc.add_paragraph("2.1 XXXXXXXX（条标题）\t1", style="TOC 3")
    doc.add_paragraph("前   言")
    doc.add_paragraph("前言一般应说明文件编制依据或背景。")
    doc.add_paragraph("XX可靠性分配报告")
    doc.add_paragraph("1范围")
    doc.add_paragraph("【说明】")
    doc.add_paragraph("本节简单介绍本文件编制目的和适用范围。")
    doc.add_paragraph("表1-1 引用文件")
    doc.add_paragraph("图1-1 产品结构图")
    doc.add_paragraph("2 引用文件")
    doc.add_paragraph("2.1 标准文件")
    doc.add_paragraph("附录A")
    doc.save(path)

    parsed = DocxTemplateParser().parse(str(path))
    flat_titles = []

    def walk(chapters):
        for chapter in chapters:
            flat_titles.append((chapter["number"], chapter["title"]))
            walk(chapter.get("sub_chapters", []))

    walk(parsed["chapters"])

    assert ("", "前言") == flat_titles[0]
    assert ("1", "范围") in flat_titles
    assert ("2", "引用文件") in flat_titles
    assert ("2.1", "标准文件") in flat_titles
    assert ("附录A", "") in flat_titles
    appendix = next(chapter for chapter in parsed["chapters"] if chapter["number"] == "附录A")
    assert appendix["required"] is False
    assert all("章标题" not in title and "条标题" not in title for _, title in flat_titles)
    assert all(not title.startswith(("表", "图")) for _, title in flat_titles)


def test_chapter_prompt_builder_combines_template_guidance_and_inputs():
    from generators.prompt_builder import ChapterPromptBuilder
    from templates.template_manager import ChapterTemplate

    chapter = ChapterTemplate(
        number="1",
        title="范围",
        guidance_prompt="说明文档适用边界。",
        template_blocks=[
            {"type": "template_text", "label": "模板文字", "text": "NNNNN 适用于待补充。"},
            {
                "type": "template_list",
                "label": "模板列表",
                "can_expand": True,
                "items": [
                    {"marker": "a）", "text": "XXXXXX；"},
                    {"marker": "b）", "text": "XXXXXX；"},
                    {"marker": "……", "text": ""},
                ],
            },
            {"type": "formula", "label": "模板公式", "text": "λs=Σniλi"},
            {"type": "example", "label": "举例", "text": "示例内容仅作参考。"},
        ],
    )

    prompt = ChapterPromptBuilder().build(
        title="测试文档",
        template_name="测试模板",
        chapter=chapter,
        inputs={"product_name": "某产品", "technical_params": "参数A=1"},
    )

    assert "当前章节：1 范围" in prompt
    assert "某产品" in prompt
    prompt_with_extra = ChapterPromptBuilder().build(
        title="测试文档",
        template_name="测试模板",
        chapter=chapter,
        inputs={"product_name": "某产品", "additional_context": "补充约束：按A类任务处理"},
    )
    assert "补充约束：按A类任务处理" in prompt_with_extra
    assert "说明文档适用边界" in prompt
    assert "只输出当前章节" in prompt
    assert "模板列表" in prompt
    assert "a）XXXXXX；" in prompt
    assert "保留原列表编号形式" in prompt
    assert "模板公式：λs=Σniλi" in prompt
    assert "公式、变量定义、单位或符号关系" in prompt

    prompt_with_doc = ChapterPromptBuilder().build(
        title="测试文档",
        template_name="测试模板",
        chapter=chapter,
        inputs={"product_name": "某产品", "supplement_doc_text": "上传材料：环境试验范围包含温度和湿热。"},
    )
    assert "上传材料：环境试验范围包含温度和湿热。" in prompt_with_doc

    prompt_with_dynamic = ChapterPromptBuilder().build(
        title="测试文档",
        template_name="测试模板",
        chapter=chapter,
        inputs={
            "product_name": "某产品",
            "dynamic_fields": {
                "range_limit": "仅用于鉴定试验",
                "other_value": "不应出现在本章",
            },
            "dynamic_field_definitions": [
                {"key": "range_limit", "label": "适用限制", "chapter_keys": ["1::范围"]},
                {"key": "other_value", "label": "其他章节字段", "chapter_keys": ["2::引用文件"]},
            ],
        },
    )
    assert "适用限制：仅用于鉴定试验" in prompt_with_dynamic
    assert "不应出现在本章" not in prompt_with_dynamic


def test_chapter_prompt_builder_treats_saved_instruction_like_text_as_guidance():
    from generators.prompt_builder import ChapterPromptBuilder
    from templates.template_manager import ChapterTemplate

    chapter = ChapterTemplate(
        number="",
        title="前言",
        template_blocks=[
            {"type": "template_text", "label": "模板文字", "text": "XXXXXXXXXXXXXXXXXXXXXXXX。"},
            {"type": "template_text", "label": "模板文字", "text": "前言一般应说明下列内容："},
            {"type": "template_list", "label": "模板列表", "text": "a）说明文件编制依据或背景；"},
            {"type": "template_text", "label": "模板文字", "text": "注：前言属概述要素，可根据需要剪裁。"},
        ],
    )

    prompt = ChapterPromptBuilder().build(
        title="可靠性分配报告",
        template_name="可靠性分配报告模板",
        chapter=chapter,
        inputs={},
    )

    template_part = prompt.split("【模板正文/结构参考】", 1)[1].split("【模板说明和示例】", 1)[0]
    guidance_part = prompt.split("【模板说明和示例】", 1)[1].split("【本次回填位置】", 1)[0]
    assert "XXXXXXXXXXXXXXXXXXXXXXXX" in template_part
    assert "前言一般应说明下列内容" not in template_part
    assert "前言一般应说明下列内容" in guidance_part
    assert "a）说明文件编制依据或背景" in guidance_part
    assert "注：前言属概述要素" in guidance_part
    assert "说明：注：" not in guidance_part


def test_generate_template_api_returns_templates():
    from web.app import app

    response = app.test_client().get("/generate/api/templates")
    assert response.status_code == 200
    data = response.get_json()
    assert data["templates"]
    assert all("can_generate" in template for template in data["templates"])
    assert any(template["can_generate"] for template in data["templates"])
    assert any(not template["can_generate"] for template in data["templates"])


def test_template_library_page_and_api_load():
    from web.app import app

    page = app.test_client().get("/template-library/")
    response = app.test_client().get("/template-library/api/templates")

    assert page.status_code == 200
    assert "生成模板库".encode("utf-8") in page.data
    assert response.status_code == 200
    assert "templates" in response.get_json()


def test_template_library_rejects_invalid_dynamic_input_fields():
    from web.app import app

    response = app.test_client().post(
        "/template-library/api/templates",
        data=json.dumps({
            "name": "无效字段模板",
            "chapters": [{"number": "1", "title": "范围"}],
            "input_fields": [{
                "key": "",
                "label": "额定电压",
                "type": "text",
            }],
        }),
        content_type="application/json",
    )

    assert response.status_code == 400
    assert "字段键和名称不能为空" in response.get_json()["error"]


def test_generate_page_default_prompt_is_template_agnostic():
    from web.app import app

    page = app.test_client().get("/generate/")

    assert page.status_code == 200
    assert "某型电动作动器鉴定试验项目".encode("utf-8") not in page.data
    assert "温度、高度、温度变化、湿热等自然环境类鉴定试验".encode("utf-8") not in page.data
    assert "collectTemplateGenerationSignals".encode("utf-8") in page.data
    assert "模板专属输入".encode("utf-8") in page.data


def test_generate_supplement_doc_upload_extracts_text_and_tables():
    from io import BytesIO
    from web.app import app

    doc = Document()
    doc.add_paragraph("补充背景：用于鉴定试验。")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "参数"
    table.cell(0, 1).text = "要求"
    table.cell(1, 0).text = "温度"
    table.cell(1, 1).text = "待补充"
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = app.test_client().post(
        "/generate/api/supplement-doc",
        data={"file": (buffer, "reference.docx")},
        content_type="multipart/form-data",
    )

    assert response.status_code == 200
    data = response.get_json()
    assert data["filename"] == "reference.docx"
    assert "补充背景" in data["text"]
    assert "参数 | 要求" in data["text"]


def test_template_library_parse_saves_source_docx_asset(tmp_path):
    from io import BytesIO
    from pathlib import Path
    from web.app import app

    doc = Document()
    doc.add_heading("1 范围", level=1)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = app.test_client().post(
        "/template-library/api/parse",
        data={"file": (buffer, "asset_template.docx")},
        content_type="multipart/form-data",
    )

    assert response.status_code == 200
    data = response.get_json()
    assert data["id"].startswith("template_")
    source_path = Path(app.root_path).parent / data["metadata"]["source_docx_path"]
    assert source_path.exists()
    assert source_path.name == "source.docx"


def test_template_library_replace_source_reparses_template(tmp_path):
    from io import BytesIO
    from pathlib import Path
    from web.app import app

    first = Document()
    first.add_heading("1 范围", level=1)
    first_buffer = BytesIO()
    first.save(first_buffer)
    first_buffer.seek(0)

    client = app.test_client()
    parsed_response = client.post(
        "/template-library/api/parse",
        data={"file": (first_buffer, "first_template.docx")},
        content_type="multipart/form-data",
    )
    parsed = parsed_response.get_json()
    create_response = client.post(
        "/template-library/api/templates",
        data=json.dumps({
            "id": parsed["id"],
            "name": "替换源文件测试模板",
            "description": "",
            "chapters": parsed["chapters"],
            "metadata": parsed["metadata"],
            "source_type": "uploaded_docx",
        }),
        content_type="application/json",
    )
    template = create_response.get_json()["template"]

    second = Document()
    second.add_heading("1 新范围", level=1)
    second.add_heading("2 新依据", level=1)
    second_buffer = BytesIO()
    second.save(second_buffer)
    second_buffer.seek(0)

    replace_response = client.post(
        f"/template-library/api/templates/{template['id']}/source",
        data={"file": (second_buffer, "second_template.docx")},
        content_type="multipart/form-data",
    )

    assert replace_response.status_code == 200
    updated = replace_response.get_json()["template"]
    assert updated["metadata"]["version"] == 2
    assert updated["metadata"]["source_filename"] == "second_template.docx"
    assert updated["chapters"][0]["title"] == "新范围"
    assert updated["chapters"][1]["title"] == "新依据"
    assert (Path(app.root_path).parent / updated["metadata"]["source_docx_path"]).exists()

    client.delete(f"/template-library/api/templates/{template['id']}")


def test_generated_docx_quality_checker_flags_placeholder_and_guidance(tmp_path):
    from generators.quality_checker import GeneratedDocxQualityChecker

    path = tmp_path / "quality.docx"
    doc = Document()
    doc.add_paragraph("本文包含 NNNNN 占位符。")
    doc.add_paragraph("【说明】正式文档中删除。")
    doc.save(path)

    result = GeneratedDocxQualityChecker().check(str(path))

    assert result["passed"] is False
    codes = {item["code"] for item in result["checks"] if not item["passed"]}
    assert "placeholder_residue" in codes
    assert "guidance_residue" in codes
    failed = {item["code"]: item for item in result["checks"] if not item["passed"]}
    assert failed["placeholder_residue"]["locations"][0]["location"].startswith("段落")
    assert "NNNNN" in failed["placeholder_residue"]["locations"][0]["excerpt"]


def test_generate_start_requires_product_name():
    from web.app import app

    response = app.test_client().post(
        "/generate/start",
        data=json.dumps({
            "template_id": "requirements",
            "title": "需求文档",
            "inputs": {},
        }),
        content_type="application/json",
    )

    assert response.status_code == 400
    assert "产品名称" in response.get_json()["error"]


def test_generate_record_delete_api():
    from web.app import app

    db = app.db
    task_id = db.create_generate_task(
        doc_type="template_test",
        template_name="测试模板",
        params={"title": "测试生成记录"},
    )
    db.update_generate_task(task_id, status="completed", progress=100, result_path="")

    response = app.test_client().delete(f"/generate/delete/{task_id}")

    assert response.status_code == 200
    assert response.get_json()["deleted"] == 1
    assert db.get_generate_task(task_id) is None


def test_generate_recent_api_paginates_and_decorates_status():
    from web.app import app

    db = app.db
    task_ids = []
    for index in range(6):
        task_id = db.create_generate_task(
            doc_type=f"template_page_{index}",
            template_name=f"分页模板 {index}",
            params={"title": f"分页记录 {index}"},
        )
        db.update_generate_task(
            task_id,
            status="completed" if index % 2 == 0 else "failed",
            progress=100,
            result_path=f"generated_{index}.docx" if index % 2 == 0 else "",
            error="生成失败" if index % 2 else "",
        )
        task_ids.append(task_id)

    response = app.test_client().get("/generate/api/recent?page=2")

    assert response.status_code == 200
    data = response.get_json()
    assert data["pagination"]["page_size"] == 5
    assert data["pagination"]["page"] >= 1
    assert data["pagination"]["total"] >= 6
    assert len(data["tasks"]) <= 5
    assert all("status_label" in task for task in data["tasks"])
    assert all("badge_class" in task for task in data["tasks"])

    for task_id in task_ids:
        db.delete_generate_task(task_id)


def test_generate_record_rerun_creates_new_task(tmp_path):
    from io import BytesIO
    from web.app import app

    client = app.test_client()
    doc = Document()
    doc.add_heading("1 范围", level=1)
    doc.add_paragraph("NNNNN 正式正文")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    parsed_response = client.post(
        "/template-library/api/parse",
        data={"file": (buffer, "rerun_template.docx")},
        content_type="multipart/form-data",
    )
    parsed = parsed_response.get_json()
    create_response = client.post(
        "/template-library/api/templates",
        data=json.dumps({
            "id": parsed["id"],
            "name": "重跑测试模板",
            "chapters": parsed["chapters"],
            "metadata": parsed["metadata"],
            "source_type": "uploaded_docx",
        }),
        content_type="application/json",
    )
    template = create_response.get_json()["template"]

    db = app.db
    task_id = db.create_generate_task(
        doc_type=template["id"],
        template_name=template["name"],
        params={
            "title": "重跑文档",
            "template_id": template["id"],
            "template_name": template["name"],
            "inputs": {"product_name": "某产品", "generation_mode": "template_fill"},
        },
    )
    db.update_generate_task(task_id, status="failed", progress=100, error="测试失败")

    response = client.post(f"/generate/rerun/{task_id}")

    assert response.status_code == 200
    new_task_id = response.get_json()["task_id"]
    assert new_task_id != task_id
    assert db.get_generate_task(new_task_id) is not None

    db.delete_generate_task(task_id)
    db.delete_generate_task(new_task_id)
    client.delete(f"/template-library/api/templates/{template['id']}")


def test_running_generate_record_cannot_be_deleted():
    from web.app import app

    db = app.db
    task_id = db.create_generate_task(
        doc_type="template_test",
        template_name="测试模板",
        params={"title": "运行中生成记录"},
    )
    db.update_generate_task(task_id, status="processing", progress=20)

    response = app.test_client().delete(f"/generate/delete/{task_id}")

    assert response.status_code == 400
    assert "仍在执行" in response.get_json()["error"]
    assert db.get_generate_task(task_id) is not None
    db.delete_generate_task(task_id)


def test_generate_task_can_pause_resume_and_cancel():
    from web.app import app

    db = app.db
    task_id = db.create_generate_task(
        doc_type="template_control",
        template_name="生成控制测试模板",
        params={"title": "生成控制测试"},
    )
    db.update_generate_task(task_id, status="processing", progress=35)
    client = app.test_client()

    pause_response = client.post(f"/generate/pause/{task_id}")
    assert pause_response.status_code == 200
    assert db.get_generate_task(task_id)["status"] == "paused"

    resume_response = client.post(f"/generate/resume/{task_id}")
    assert resume_response.status_code == 200
    assert db.get_generate_task(task_id)["status"] == "processing"

    cancel_response = client.post(f"/generate/cancel/{task_id}")
    assert cancel_response.status_code == 200
    canceled = db.get_generate_task(task_id)
    assert canceled["status"] == "canceled"
    assert canceled["progress_stage"] == "canceled"
    assert canceled["completed_at"]
    db.delete_generate_task(task_id)


def test_paused_generate_record_cannot_be_deleted_or_rerun():
    from web.app import app

    db = app.db
    task_id = db.create_generate_task(
        doc_type="template_paused",
        template_name="暂停任务测试模板",
        params={"title": "暂停任务测试"},
    )
    db.update_generate_task(task_id, status="paused", progress=40)
    client = app.test_client()

    delete_response = client.delete(f"/generate/delete/{task_id}")
    rerun_response = client.post(f"/generate/rerun/{task_id}")

    assert delete_response.status_code == 400
    assert rerun_response.status_code == 400
    assert db.get_generate_task(task_id)["status"] == "paused"
    db.delete_generate_task(task_id)


def test_template_library_document_kind_management(tmp_path, monkeypatch):
    from templates.template_manager import TemplateManager as RealTemplateManager
    from web.app import app
    from web.routes import template_library

    store_path = tmp_path / "generation_templates.json"
    kinds_path = tmp_path / "document_kinds.json"
    monkeypatch.setattr(template_library, "DOCUMENT_KINDS_FILE", kinds_path)
    monkeypatch.setattr(
        template_library,
        "TemplateManager",
        lambda: RealTemplateManager(custom_store_path=str(store_path)),
    )

    client = app.test_client()
    created = client.post(
        "/template-library/api/document-kinds",
        data=json.dumps({"name": "可靠性分配临时"}),
        content_type="application/json",
    )
    assert created.status_code == 200
    assert any(item["name"] == "可靠性分配临时" for item in created.get_json()["document_kinds"])

    duplicate = client.post(
        "/template-library/api/document-kinds",
        data=json.dumps({"name": "可靠性分配临时"}),
        content_type="application/json",
    )
    assert duplicate.status_code == 400

    manager = RealTemplateManager(custom_store_path=str(store_path))
    manager.create_template(
        name="可靠性分配模板",
        description="",
        chapters=[{"number": "1", "title": "范围"}],
        metadata={"document_kind_name": "可靠性分配临时"},
    )

    protected = client.post(
        "/template-library/api/document-kinds/delete",
        data=json.dumps({"name": "可靠性分配临时"}),
        content_type="application/json",
    )
    assert protected.status_code == 400
    assert "模板使用" in protected.get_json()["error"]

    renamed = client.post(
        "/template-library/api/document-kinds/rename",
        data=json.dumps({"old_name": "可靠性分配临时", "new_name": "可靠性分配报告修订"}),
        content_type="application/json",
    )
    assert renamed.status_code == 200
    templates = client.get("/template-library/api/templates").get_json()["templates"]
    assert any(
        template["metadata"].get("document_kind_name") == "可靠性分配报告修订"
        for template in templates
    )

    unused = client.post(
        "/template-library/api/document-kinds",
        data=json.dumps({"name": "未使用类型"}),
        content_type="application/json",
    )
    assert unused.status_code == 200
    deleted = client.post(
        "/template-library/api/document-kinds/delete",
        data=json.dumps({"name": "未使用类型"}),
        content_type="application/json",
    )
    assert deleted.status_code == 200
