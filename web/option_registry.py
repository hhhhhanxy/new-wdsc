"""Shared option registries for models, specialties and reference cases."""
from __future__ import annotations

import json
from pathlib import Path
from typing import Any

from config.model_service import get_model_service_value


MODEL_OPTIONS_PATH = Path("config/model_options.json")
PROFESSIONAL_CASES_PATH = Path("config/professional_cases.json")
SYSTEM_DEFAULT_MODEL_ID = "system_default"


def get_model_options(usage: str = "both") -> list[dict[str, Any]]:
    usage = str(usage or "both")
    default_model_name, default_provider, default_base_url = _runtime_default_model()
    options = [{
        "id": SYSTEM_DEFAULT_MODEL_ID,
        "name": f"系统默认（{default_model_name}）" if default_model_name else "系统默认",
        "provider": default_provider,
        "model": default_model_name,
        "base_url": default_base_url,
        "usage": ["review", "generate"],
        "is_default": True,
    }]
    options.extend(_load_deployed_models_from_service())
    options.extend(_load_model_options())

    filtered = []
    seen = set()
    for option in options:
        usages = _normalize_usage(option.get("usage"))
        if "both" not in usages and usage not in usages:
            continue
        item = dict(option)
        model_name = item.get("model") or item.get("model_name")
        item["model"] = model_name
        item.setdefault("provider", default_provider)
        item.setdefault("base_url", default_base_url)
        item.setdefault("id", _model_id(item))
        item.setdefault("name", item.get("display_name") or model_name or item.get("id") or "未命名模型")
        item["is_default"] = item.get("id") == SYSTEM_DEFAULT_MODEL_ID or bool(item.get("is_default"))
        dedupe_key = (item.get("id"), item.get("provider"), item.get("model"), item.get("base_url"))
        if dedupe_key in seen:
            continue
        seen.add(dedupe_key)
        filtered.append(item)
    return filtered


def resolve_model_option(model_id: str | None, usage: str = "both") -> dict[str, Any]:
    options = get_model_options(usage)
    default = next((item for item in options if item.get("id") == SYSTEM_DEFAULT_MODEL_ID), options[0])
    if not model_id:
        return default
    return next((item for item in options if item.get("id") == model_id), default)


def _runtime_default_model() -> tuple[str, str, str]:
    try:
        from config.settings import settings

        return settings.llm_model, settings.llm_provider, settings.llm_base_url
    except Exception:
        return (
            get_model_service_value("model_name", "系统默认模型"),
            get_model_service_value("provider", "system"),
            get_model_service_value("base_url", ""),
        )


def get_specialties(usage: str | None = None) -> list[dict[str, Any]]:
    usage = str(usage or "").strip()
    specialties = [item for item in _load_professional_cases().get("specialties", []) if isinstance(item, dict)]
    if not usage:
        return specialties
    return [
        item for item in specialties
        if usage in (item.get("usage") or []) or "both" in (item.get("usage") or [])
    ]


def get_specialty_groups(usage: str | None = None) -> list[dict[str, Any]]:
    specialties = get_specialties(usage)
    config = _load_professional_cases()
    configured_groups = [
        item for item in config.get("specialty_groups", [])
        if isinstance(item, dict)
    ]
    if not configured_groups:
        configured_groups = [
            {"id": "core", "name": "一核"},
            {"id": "wings", "name": "多翼"},
        ]
    grouped = []
    used_ids = set()
    for group in configured_groups:
        group_id = str(group.get("id") or "").strip()
        children = [
            item for item in specialties
            if str(item.get("group_id") or "").strip() == group_id
        ]
        used_ids.update(str(item.get("id") or "") for item in children)
        grouped.append({
            "id": group_id,
            "name": group.get("name") or group_id,
            "specialties": children,
        })
    ungrouped = [
        item for item in specialties
        if str(item.get("id") or "") not in used_ids
    ]
    if ungrouped:
        grouped.append({
            "id": "ungrouped",
            "name": "未分组",
            "specialties": ungrouped,
        })
    return grouped


def get_specialty(specialty_id: str | None) -> dict[str, Any] | None:
    if not specialty_id:
        return None
    return next((item for item in get_specialties() if item.get("id") == specialty_id), None)


def get_reference_cases(specialty_id: str | None) -> list[dict[str, Any]]:
    specialty = get_specialty(specialty_id)
    if not specialty:
        return []
    return [item for item in specialty.get("reference_cases", []) if isinstance(item, dict)]


def add_reference_case(specialty_id: str, case: dict[str, Any]) -> dict[str, Any]:
    config = _load_professional_cases()
    specialty = _find_specialty_in_config(config, specialty_id)
    if not specialty:
        raise ValueError("专业不存在")
    cases = [item for item in specialty.get("reference_cases", []) if isinstance(item, dict)]
    case = dict(case)
    cases.append(case)
    specialty["reference_cases"] = cases
    _save_professional_cases(config)
    return case


def delete_reference_case(specialty_id: str, case_id: str) -> dict[str, Any] | None:
    config = _load_professional_cases()
    specialty = _find_specialty_in_config(config, specialty_id)
    if not specialty:
        return None
    cases = [item for item in specialty.get("reference_cases", []) if isinstance(item, dict)]
    target = next((item for item in cases if str(item.get("id")) == str(case_id)), None)
    specialty["reference_cases"] = [item for item in cases if str(item.get("id")) != str(case_id)]
    _save_professional_cases(config)
    return target


def resolve_reference_cases(specialty_id: str | None, case_ids: list[str] | None) -> list[dict[str, Any]]:
    selected = set(str(item) for item in (case_ids or []) if item)
    if not selected:
        return []
    return [case for case in get_reference_cases(specialty_id) if str(case.get("id")) in selected]


def build_reference_case_context(cases: list[dict[str, Any]]) -> str:
    parts = []
    for case in cases:
        features = case.get("features") or {}
        block = [
            f"参考案例：{case.get('name') or case.get('id')}",
            f"文档类型：{case.get('doc_type') or '未标明'}",
            f"适用场景：{case.get('scenario') or '未标明'}",
            "章节结构：" + _join_feature(features.get("chapter_structure")),
            "术语口径：" + _join_feature(features.get("terminology")),
            "写法特点：" + _join_feature(features.get("writing_style")),
            "格式特点：" + _join_feature(features.get("format_features")),
        ]
        docx_excerpt = _extract_case_docx_excerpt(case.get("file_path"))
        if docx_excerpt:
            block.extend(["案例文档摘录：", docx_excerpt])
        parts.append("\n".join(block))
    return "\n\n".join(parts)


def _load_model_options() -> list[dict[str, Any]]:
    if not MODEL_OPTIONS_PATH.exists():
        return []
    with MODEL_OPTIONS_PATH.open("r", encoding="utf-8") as fh:
        data = json.load(fh)
    if not isinstance(data, list):
        raise ValueError(f"模型选项配置格式错误: {MODEL_OPTIONS_PATH}")
    return [item for item in data if isinstance(item, dict) and item.get("enabled", True)]


def _load_deployed_models_from_service() -> list[dict[str, Any]]:
    models = (
        get_model_service_value("deployed_models")
        or get_model_service_value("available_models")
        or get_model_service_value("models")
        or []
    )
    if not isinstance(models, list):
        return []
    return [_normalize_model_option(item) for item in models if isinstance(item, (dict, str))]


def _normalize_model_option(item: Any) -> dict[str, Any]:
    if isinstance(item, str):
        return {"id": _safe_model_id(item), "name": item, "model": item}
    model_name = item.get("model") or item.get("model_name") or item.get("name")
    option = dict(item)
    option["model"] = model_name
    option.setdefault("id", _model_id(option))
    option.setdefault("name", option.get("display_name") or model_name or option.get("id"))
    option.setdefault("usage", ["review", "generate"])
    return option


def _model_id(option: dict[str, Any]) -> str:
    raw = option.get("id") or option.get("model") or option.get("name") or "model"
    return _safe_model_id(str(raw))


def _safe_model_id(value: str) -> str:
    return "".join(ch if ch.isalnum() else "_" for ch in value).strip("_").lower() or "model"


def _load_professional_cases() -> dict[str, Any]:
    if not PROFESSIONAL_CASES_PATH.exists():
        return {"specialties": []}
    with PROFESSIONAL_CASES_PATH.open("r", encoding="utf-8") as fh:
        data = json.load(fh)
    if not isinstance(data, dict):
        raise ValueError(f"专业案例配置格式错误: {PROFESSIONAL_CASES_PATH}")
    data.setdefault("specialties", [])
    return data


def _save_professional_cases(data: dict[str, Any]) -> None:
    PROFESSIONAL_CASES_PATH.parent.mkdir(parents=True, exist_ok=True)
    with PROFESSIONAL_CASES_PATH.open("w", encoding="utf-8") as fh:
        json.dump(data, fh, ensure_ascii=False, indent=2)


def _find_specialty_in_config(config: dict[str, Any], specialty_id: str | None) -> dict[str, Any] | None:
    target = str(specialty_id or "").strip()
    if not target:
        return None
    return next(
        (
            item for item in config.get("specialties", [])
            if isinstance(item, dict) and str(item.get("id") or "").strip() == target
        ),
        None,
    )


def _normalize_usage(value: Any) -> set[str]:
    if isinstance(value, str):
        return {value}
    if isinstance(value, list):
        return {str(item) for item in value if item}
    return {"both"}


def _join_feature(value: Any) -> str:
    if isinstance(value, list):
        return "；".join(str(item) for item in value if item) or "未配置"
    return str(value or "未配置")


def _extract_case_docx_excerpt(file_path: Any, max_chars: int = 6000) -> str:
    if not file_path:
        return ""
    path = Path(str(file_path))
    if not path.exists() or path.suffix.lower() != ".docx":
        return ""
    try:
        from docx import Document

        doc = Document(str(path))
        parts = []
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)
        for table_index, table in enumerate(doc.tables, start=1):
            rows = []
            for row in table.rows[:8]:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                if any(cells):
                    rows.append(" | ".join(cells))
            if rows:
                parts.append(f"表格{table_index}：\n" + "\n".join(rows))
        text = "\n".join(parts)
        if len(text) > max_chars:
            return text[:max_chars] + "\n...（案例文档较长，已截取前6000字）"
        return text
    except Exception:
        return ""
