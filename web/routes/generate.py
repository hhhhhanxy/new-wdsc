"""
Generate routes for the web application.
"""
import os
import json
import threading
import time
import uuid
from pathlib import Path
from flask import Blueprint, render_template, request, jsonify, send_file, current_app
from werkzeug.utils import secure_filename
from web.time_utils import beijing_now_str

bp = Blueprint('generate', __name__)
RUNNING_STATUSES = {'pending', 'processing', 'paused'}


class GenerationTaskCanceled(Exception):
    """Raised when a generation task is stopped by the user."""

    is_task_control = True


@bp.route('/')
def index():
    return render_template('generate.html', active_page='generate')


@bp.route('/api/templates')
def templates():
    """Return all templates and identify which ones have a source DOCX."""
    from templates.template_manager import TemplateManager

    manager = TemplateManager()
    templates = manager.list_template_dicts()
    for template in templates:
        template["can_generate"] = _has_source_docx(template)
        template["unavailable_reason"] = "" if template["can_generate"] else "缺少原始 DOCX 模板文件"
    return jsonify({"templates": templates})


@bp.route('/api/templates/<template_id>')
def template_detail(template_id):
    """Return one template with its chapter structure."""
    from templates.template_manager import TemplateManager

    manager = TemplateManager()
    template = manager.get_template(template_id)
    if not template:
        return jsonify({"error": f"模板不存在: {template_id}"}), 404
    return jsonify(manager.serialize_template(template))


@bp.route('/api/supplement-doc', methods=['POST'])
@bp.route('/api/reference-doc', methods=['POST'])
def upload_supplement_doc():
    """Upload and extract text from a DOCX supplement material file."""
    file = request.files.get('file')
    if not file or not file.filename:
        return jsonify({'error': '请选择要上传的 Word 文档'}), 400

    filename = secure_filename(file.filename)
    if not filename.lower().endswith('.docx'):
        return jsonify({'error': '当前仅支持 .docx 格式的 Word 文档'}), 400

    reference_dir = Path(current_app.config['UPLOAD_FOLDER']) / 'generation_references'
    reference_dir.mkdir(parents=True, exist_ok=True)
    saved_name = f"{uuid.uuid4().hex}_{filename}"
    saved_path = reference_dir / saved_name
    file.save(saved_path)

    try:
        extracted = _extract_docx_reference_text(saved_path)
    except Exception as exc:
        saved_path.unlink(missing_ok=True)
        return jsonify({'error': f'补充材料解析失败：{exc}'}), 400

    if not extracted.strip():
        return jsonify({'error': '补充材料未解析到可用文本'}), 400

    return jsonify({
        'filename': filename,
        'saved_path': str(saved_path),
        'text': extracted,
        'chars': len(extracted),
    })


@bp.route('/start', methods=['POST'])
def start():
    data = request.get_json()
    return _create_generation_task(data)


@bp.route('/status/<int:task_id>')
def status(task_id):
    task = current_app.db.get_generate_task(task_id)
    if not task:
        return jsonify({'error': '任务不存在'}), 404
    return jsonify(dict(task))


@bp.route('/pause/<int:task_id>', methods=['POST'])
def pause(task_id):
    db = current_app.db
    task = db.get_generate_task(task_id)
    if not task:
        return jsonify({'error': '生成任务不存在'}), 404
    if task.get('status') != 'processing':
        return jsonify({'error': '只有生成中的任务可以暂停'}), 400
    db.update_generate_task(
        task_id,
        status='paused',
        progress_stage='paused',
        progress_message='生成已暂停，点击继续后恢复',
    )
    return jsonify({'ok': True, 'status': 'paused', 'progress': task.get('progress', 0)})


@bp.route('/resume/<int:task_id>', methods=['POST'])
def resume(task_id):
    db = current_app.db
    task = db.get_generate_task(task_id)
    if not task:
        return jsonify({'error': '生成任务不存在'}), 404
    if task.get('status') != 'paused':
        return jsonify({'error': '只有已暂停的生成任务可以继续'}), 400
    db.update_generate_task(
        task_id,
        status='processing',
        progress_stage='filling',
        progress_message='正在恢复文档生成',
    )
    return jsonify({'ok': True, 'status': 'processing', 'progress': task.get('progress', 0)})


@bp.route('/cancel/<int:task_id>', methods=['POST'])
def cancel(task_id):
    db = current_app.db
    task = db.get_generate_task(task_id)
    if not task:
        return jsonify({'error': '生成任务不存在'}), 404
    if task.get('status') not in RUNNING_STATUSES:
        return jsonify({'error': '只有等待中、生成中或已暂停的任务可以停止'}), 400
    db.update_generate_task(
        task_id,
        status='canceled',
        progress_stage='canceled',
        progress_message='用户手动停止生成',
        error='用户手动停止生成任务',
        completed_at=beijing_now_str(),
    )
    return jsonify({'ok': True, 'status': 'canceled', 'progress': task.get('progress', 0)})


@bp.route('/api/recent')
def recent():
    """Return recent generation task records."""
    page_size = 5
    page = request.args.get('page', 1, type=int)
    total = current_app.db.count_generate_tasks()
    total_pages = max(1, (total + page_size - 1) // page_size)
    page = min(max(page, 1), total_pages)
    offset = (page - 1) * page_size
    tasks = [
        _decorate_generate_task(task)
        for task in current_app.db.get_recent_generate_tasks(limit=page_size, offset=offset)
    ]
    return jsonify({
        'tasks': tasks,
        'pagination': {
            'page': page,
            'page_size': page_size,
            'total': total,
            'total_pages': total_pages,
            'has_prev': page > 1,
            'has_next': page < total_pages,
            'start': offset + 1 if total else 0,
            'end': min(offset + len(tasks), total),
        },
    })


@bp.route('/rerun/<int:task_id>', methods=['POST'])
def rerun(task_id):
    task = current_app.db.get_generate_task(task_id)
    if not task:
        return jsonify({'error': '生成记录不存在'}), 404
    if task.get('status') in RUNNING_STATUSES:
        return jsonify({'error': '该生成任务仍在执行，不能重新生成'}), 400
    try:
        params_data = json.loads(task.get('params') or '{}')
    except (TypeError, json.JSONDecodeError):
        return jsonify({'error': '原始生成参数不可用，不能重新生成'}), 400
    template_id = params_data.get('template_id') or task.get('doc_type')
    title = params_data.get('title') or f"{task.get('template_name') or '生成文档'}_重新生成"
    inputs = params_data.get('inputs') or {}
    rerun_data = {
        'template_id': template_id,
        'title': title,
        'inputs': inputs,
        'model_id': params_data.get('model_id') or task.get('model_id'),
        'specialty_id': params_data.get('specialty_id') or task.get('specialty_id'),
        'reference_case_ids': [
            item.get('id') for item in (params_data.get('reference_cases') or [])
            if isinstance(item, dict) and item.get('id')
        ],
        'rerun_from': task_id,
    }
    return _create_generation_task(rerun_data)


@bp.route('/download/<int:task_id>')
def download(task_id):
    task = current_app.db.get_generate_task(task_id)
    if not task or task['status'] != 'completed':
        return '文件未就绪', 404

    result_path = task.get('result_path')
    if result_path and os.path.exists(result_path):
        return send_file(result_path, as_attachment=True)

    return '文件不存在', 404


@bp.route('/review/<int:task_id>', methods=['POST'])
def review_generated(task_id):
    """Create a review task for a generated DOCX and jump into the review flow."""
    task = current_app.db.get_generate_task(task_id)
    if not task:
        return jsonify({'error': '生成记录不存在'}), 404
    if task.get('status') != 'completed':
        return jsonify({'error': '只有已完成的生成文档可以送审'}), 400

    result_path = task.get('result_path')
    if not result_path or not os.path.exists(result_path):
        return jsonify({'error': '生成文档文件不存在，无法送审'}), 404
    if not result_path.lower().endswith('.docx'):
        return jsonify({'error': '当前仅支持 DOCX 生成文档送审'}), 400

    from web.tasks import run_review_task

    filename = os.path.basename(result_path)
    payload = request.get_json(silent=True) or {}
    try:
        params_data = json.loads(task.get('params') or '{}')
    except (TypeError, json.JSONDecodeError):
        params_data = {}
    specialty_id = str(task.get('specialty_id') or params_data.get('specialty_id') or '').strip()
    specialty_name = str(task.get('specialty_name') or params_data.get('specialty_name') or '').strip()
    document_kind_name = str(task.get('document_kind_name') or params_data.get('document_kind_name') or '').strip()
    rule_set = f"specialty:{specialty_id}" if specialty_id else "all"
    from web.option_registry import resolve_model_option
    review_model = resolve_model_option(payload.get('model_id'), 'review')
    review_task_id = current_app.db.create_review_task(
        filename=filename,
        filepath=result_path,
        mode='generated_doc',
        rule_set=rule_set,
        model_id=review_model.get('id'),
        model_name=review_model.get('name'),
        specialty_id=specialty_id,
        specialty_name=specialty_name,
        document_kind_name=document_kind_name,
        source_generate_task_id=task_id,
    )

    params_data['linked_review_task_id'] = review_task_id
    params_data['auto_review'] = True
    cursor = current_app.db.conn.cursor()
    cursor.execute(
        'UPDATE generate_tasks SET params = ? WHERE id = ?',
        (json.dumps(params_data, ensure_ascii=False), task_id),
    )
    current_app.db.conn.commit()

    thread = threading.Thread(
        target=run_review_task,
        args=(review_task_id, result_path, rule_set, current_app.db),
    )
    thread.daemon = True
    thread.start()

    return jsonify({'ok': True, 'review_task_id': review_task_id})


@bp.route('/delete/<int:task_id>', methods=['DELETE'])
def delete_task(task_id):
    task = current_app.db.get_generate_task(task_id)
    if not task:
        return jsonify({'error': '生成记录不存在'}), 404
    if task.get('status') in RUNNING_STATUSES:
        return jsonify({'error': '生成任务仍在执行，完成或失败后再删除'}), 400
    deleted = current_app.db.delete_generate_task(task_id)
    return jsonify({'ok': True, 'deleted': deleted})


@bp.route('/delete-all', methods=['DELETE'])
def delete_all_tasks():
    db = current_app.db
    deleted = db.delete_finished_generate_tasks()
    remaining = db.count_generate_tasks()
    return jsonify({'ok': True, 'deleted': deleted, 'remaining': remaining})


def _create_generation_task(data: dict):
    db = current_app.db
    if not data:
        return jsonify({'error': '请输入生成需求'}), 400
    if not data.get('template_id'):
        return jsonify({'error': '请选择文档模板'}), 400

    from templates.template_manager import TemplateManager

    manager = TemplateManager()
    template = manager.get_template(data['template_id'])
    if not template:
        return jsonify({'error': f"模板不存在: {data['template_id']}"}), 400

    inputs = data.get('inputs') or {}
    generation_brief = str(inputs.get('generation_brief') or '').strip()
    supplement_text = str(inputs.get('supplement_doc_text') or '').strip()
    if not generation_brief and not supplement_text:
        return jsonify({'error': '请输入生成需求，或上传补充材料'}), 400
    inputs['dynamic_field_definitions'] = [
        {
            'key': field.key,
            'label': field.label,
            'chapter_keys': field.chapter_keys,
            'placeholder_tokens': field.placeholder_tokens,
        }
        for field in template.input_fields
    ]
    parsed_brief = _parse_generation_brief(generation_brief, _template_brief_labels(template.input_fields))
    _merge_generation_brief_inputs(inputs, parsed_brief, template.input_fields)
    template_dict = manager.serialize_template(template)
    from web.option_registry import build_reference_case_context, get_specialty, resolve_model_option
    model_option = resolve_model_option(data.get('model_id'), 'generate')
    llm_parsed_brief = _try_llm_parse_generation_brief(
        generation_brief=generation_brief,
        supplement_text=supplement_text,
        template=template,
        inputs=inputs,
        model_option=model_option,
    )
    if llm_parsed_brief:
        _merge_generation_brief_inputs(inputs, llm_parsed_brief, template.input_fields)
        parsed_brief = {**parsed_brief, **{key: value for key, value in llm_parsed_brief.items() if value}}
    data['title'] = str(
        parsed_brief.get('title')
        or _derive_title_from_brief(generation_brief, template.name)
        or data.get('title')
    ).strip()
    dynamic_values = inputs.get('dynamic_fields') or {}
    missing_dynamic = [
        field.label
        for field in template.input_fields
        if field.required and not str(dynamic_values.get(field.key, '')).strip()
    ]
    if missing_dynamic:
        return jsonify({'error': f"请填写模板必填项：{'、'.join(missing_dynamic)}"}), 400
    data['inputs'] = inputs
    specialty = get_specialty(data.get('specialty_id'))
    specialty_id = specialty.get('id') if specialty else ''
    specialty_name = specialty.get('name') if specialty else ''
    selected_cases = _resolve_template_reference_cases(
        template_dict,
        data.get('reference_case_ids') if isinstance(data.get('reference_case_ids'), list) else [],
    )
    document_kind_name = str(
        data.get('document_kind_name')
        or (template_dict.get('metadata') or {}).get('document_kind_name')
        or template.name
    ).strip()
    reference_case_context = build_reference_case_context(selected_cases)
    inputs['specialty_id'] = specialty_id
    inputs['specialty_name'] = specialty_name
    inputs['document_kind_name'] = document_kind_name
    inputs['reference_case_ids'] = [case.get('id') for case in selected_cases]
    inputs['reference_case_names'] = [case.get('name') for case in selected_cases]
    inputs['reference_case_context'] = reference_case_context
    if not _has_source_docx(template_dict):
        return jsonify({'error': '该模板没有原始 DOCX 文件，请先在生成模板库中上传模板'}), 400

    params = {
        'title': data['title'],
        'template_id': template.template_id,
        'template_name': template.name,
        'template_version': (template.metadata or {}).get('version', 1),
        'model_id': model_option.get('id'),
        'model_name': model_option.get('name'),
        'specialty_id': specialty_id,
        'specialty_name': specialty_name,
        'document_kind_name': document_kind_name,
        'reference_cases': [
            {
                'id': case.get('id'),
                'name': case.get('name'),
                'doc_type': case.get('doc_type'),
            }
            for case in selected_cases
        ],
        'inputs': inputs,
        'rerun_from': data.get('rerun_from'),
        'auto_review': False,
        'review_result': {
            'reserved': True,
            'message': '生成完成后可一键送入文档审查。',
        },
    }

    new_task_id = db.create_generate_task(
        doc_type=template.doc_type.value if template.doc_type else template.template_id,
        template_name=template.name,
        params=params,
        model_id=model_option.get('id'),
        model_name=model_option.get('name'),
        specialty_id=specialty_id,
        specialty_name=specialty_name,
        document_kind_name=document_kind_name,
        reference_cases=params['reference_cases'],
    )

    thread = threading.Thread(
        target=run_generate_task,
        args=(new_task_id, data, current_app.config['UPLOAD_FOLDER'], db)
    )
    thread.daemon = True
    thread.start()

    return jsonify({'task_id': new_task_id, 'rerun_from': data.get('rerun_from')})


def run_generate_task(task_id, data, upload_folder, db):
    """Execute document generation task."""
    try:
        _wait_if_generation_paused(db, task_id)
        db.update_generate_task(
            task_id,
            progress=5,
            status='processing',
            progress_stage='prepare',
            progress_message='正在准备生成任务',
        )

        from core.pipeline import generate_document_only
        from config.settings import settings
        from llm.client import LLMClientFactory

        db.update_generate_task(
            task_id,
            progress=15,
            progress_stage='prepare',
            progress_message='模板文件已就绪',
        )

        template_id = data.get('template_id')
        inputs = data.get('inputs') or {}
        generation_mode = str(inputs.get('generation_mode') or 'smart')

        llm_client = None
        if generation_mode == 'smart':
            from web.option_registry import resolve_model_option
            model_option = resolve_model_option(data.get('model_id'), 'generate')
            provider = model_option.get('provider') or settings.llm_provider
            model_name = model_option.get('model') or None
            display_name = model_option.get('name') or model_name or settings.llm_model
            client_kwargs = {'model': model_name}
            if model_option.get('base_url'):
                client_kwargs['base_url'] = model_option.get('base_url')
            if model_option.get('api_key'):
                client_kwargs['api_key'] = model_option.get('api_key')
            db.update_generate_task(
                task_id,
                progress=18,
                progress_stage='llm',
                progress_message=f'正在初始化生成模型：{display_name}',
            )
            llm_client = LLMClientFactory.create_client(provider, **client_kwargs)
        else:
            db.update_generate_task(
                task_id,
                progress=18,
                progress_stage='prepare',
                progress_message='已选择模板填充模式，本次不调用大模型',
            )

        def update_chapter_progress(current, total, chapter):
            _wait_if_generation_paused(db, task_id)
            percent = 20 + int((current - 1) / max(total, 1) * 75)
            title = f"{getattr(chapter, 'number', '')} {getattr(chapter, 'title', '')}".strip()
            db.update_generate_task(
                task_id,
                progress=max(20, min(percent, 95)),
                progress_stage='filling',
                progress_message=f"正在处理章节 {current}/{total}：{title}",
                current_section=current,
                total_sections=total,
            )

        output_dir = os.path.join(upload_folder, 'generated')
        result = generate_document_only(
            doc_type=template_id,
            title=data.get('title', ''),
            params={
                "template_id": template_id,
                "inputs": inputs,
                "doc_type": template_id,
                "generator": "template_docx",
            },
            llm_client=llm_client,
            output_dir=output_dir,
            progress_callback=update_chapter_progress,
            control_callback=lambda: _wait_if_generation_paused(db, task_id),
        )

        if result.error:
            raise RuntimeError(result.error)

        _wait_if_generation_paused(db, task_id)
        task = db.get_generate_task(task_id)
        if task:
            params_data = json.loads(task['params']) if isinstance(task['params'], str) else task['params']
            params_data['quality_result'] = result.quality_result
            params_data['generation_meta'] = result.generation_meta
            params_data['review_result'] = {
                'reserved': True,
                'message': '生成完成后可一键送入文档审查。',
                'executed': False,
            }
            cursor = db.conn.cursor()
            cursor.execute(
                'UPDATE generate_tasks SET params = ? WHERE id = ?',
                (json.dumps(params_data, ensure_ascii=False), task_id)
            )
            db.conn.commit()

        _wait_if_generation_paused(db, task_id)
        db.update_generate_task(
            task_id,
            progress=100,
            status='completed',
            progress_stage='completed',
            progress_message='文档生成完成，可一键送入审查',
            result_path=result.generated_path,
            error=None,
            completed_at=beijing_now_str(),
        )

    except GenerationTaskCanceled as e:
        task = db.get_generate_task(task_id)
        if task and task.get('status') != 'canceled':
            db.update_generate_task(
                task_id,
                status='canceled',
                error=str(e),
                progress_stage='canceled',
                progress_message=str(e),
                completed_at=beijing_now_str(),
            )

    except Exception as e:
        import traceback
        traceback.print_exc()
        task = db.get_generate_task(task_id)
        if not task or task.get('status') == 'canceled':
            return
        db.update_generate_task(
            task_id,
            status='failed',
            error=str(e),
            progress_stage='failed',
            progress_message='生成失败',
            completed_at=beijing_now_str(),
        )


def _resolve_template_reference_cases(template: dict, case_ids: list[str]) -> list[dict]:
    selected = set(str(item) for item in (case_ids or []) if item)
    if not selected:
        return []
    metadata = template.get('metadata') or {}
    cases = metadata.get('reference_cases') or []
    return [
        case for case in cases
        if isinstance(case, dict) and str(case.get('id')) in selected
    ]


def _derive_title_from_brief(brief: str, fallback: str) -> str:
    parsed = _parse_generation_brief(brief)
    if parsed.get('title'):
        return parsed['title'][:80]
    if parsed.get('_has_explicit_labels'):
        return fallback or '未命名生成文档'
    for line in str(brief or '').splitlines():
        title = line.strip()
        if title:
            if _split_brief_label_line(title, _default_brief_labels())[0]:
                continue
            title = title.removeprefix('请生成').strip()
            for sep in ('。', '；', ';', '，', ','):
                if sep in title:
                    title = title.split(sep, 1)[0].strip()
            return title[:60] or fallback
    return fallback or '未命名生成文档'


def _brief_label_map() -> dict[str, str]:
    return {
        '文档标题': 'title',
        '标题': 'title',
        '产品名称': 'product_name',
        '项目名称': 'project_name',
        '试验项目': 'test_item',
        '受试对象': 'test_item',
        '试验范围': 'test_scope',
        '使用场景': 'use_scene',
        '背景说明': 'background',
        '关键参数': 'technical_params',
        '产品信息': 'technical_params',
        '产品参数': 'technical_params',
        '技术参数': 'technical_params',
        '补充信息': 'additional_context',
        '补充说明': 'additional_context',
        '引用文件': 'references',
        '依据文件': 'references',
        '生成要求': 'generation_requirements',
        '写作要求': 'generation_requirements',
    }


def _default_brief_labels() -> set[str]:
    return set(_brief_label_map())


def _parse_generation_brief(brief: str, extra_labels=None) -> dict:
    """Parse label-style prompt text from the smart generation input."""
    label_map = _brief_label_map()
    labels = set(label_map)
    for item in extra_labels or []:
        label = str(item or '').strip()
        if label:
            labels.add(label)
            label_map.setdefault(label, f'field::{label}')
    parsed: dict[str, list[str]] = {}
    field_values: dict[str, list[str]] = {}
    current_key = ''
    current_label = ''
    loose_lines: list[str] = []
    for raw_line in str(brief or '').splitlines():
        line = raw_line.strip().strip('；;')
        if not line:
            continue
        label, value = _split_brief_label_line(line, labels)
        if not label and line in labels:
            label = line
            value = ''
        if label:
            current_key = label_map[label]
            current_label = label
            parsed.setdefault(current_key, [])
            field_values.setdefault(label, [])
            if value:
                parsed[current_key].append(value)
                field_values[label].append(value)
            continue
        if current_key:
            loose_key = _classify_loose_brief_line(line)
            if current_key in {'title', 'product_name', 'project_name', 'test_item'} and loose_key and loose_key != current_key:
                parsed.setdefault(loose_key, []).append(line)
                current_key = loose_key if loose_key in {'technical_params', 'references', 'additional_context'} else ''
                current_label = ''
                continue
            parsed.setdefault(current_key, []).append(line)
            if current_label:
                field_values.setdefault(current_label, []).append(line)
        else:
            loose_lines.append(line)

    _merge_loose_brief_lines(parsed, loose_lines)

    result = {key: '\n'.join(value).strip() for key, value in parsed.items() if any(value)}
    result['_field_values'] = {
        key: '\n'.join(value).strip()
        for key, value in field_values.items()
        if any(value)
    }
    result['_has_explicit_labels'] = bool(result['_field_values'])
    if not result.get('references'):
        references = _infer_references_from_brief(brief)
        if references:
            result['references'] = references
    return result


def _template_brief_labels(input_fields) -> list[str]:
    labels = []
    for field in input_fields or []:
        labels.extend([
            getattr(field, 'label', '') or '',
            getattr(field, 'key', '') or '',
            getattr(field, 'placeholder', '') or '',
        ])
        labels.extend(getattr(field, 'placeholder_tokens', []) or [])
    return [str(item).strip() for item in labels if str(item or '').strip()]


def _split_brief_label_line(line: str, labels: set[str]) -> tuple[str, str]:
    normalized_labels = {_normalize_brief_label(label): label for label in labels}
    for separator in ('：', ':'):
        if separator not in line:
            continue
        maybe_label, maybe_value = line.split(separator, 1)
        label = normalized_labels.get(_normalize_brief_label(maybe_label))
        if label:
            return label, maybe_value.strip().strip('；;')
    return '', ''


def _merge_loose_brief_lines(parsed: dict[str, list[str]], lines: list[str]) -> None:
    for line in lines:
        key = _classify_loose_brief_line(line)
        if key:
            parsed.setdefault(key, []).append(line)


def _classify_loose_brief_line(line: str) -> str:
    text = str(line or '').strip()
    if not text:
        return ''
    if any(token in text for token in ('项目编号', '项目名称', '工程项目')) or text.endswith('项目'):
        return 'project_name'
    if any(token in text for token in ('本大纲适用于', '适用于', '试验范围', '共')) and any(token in text for token in ('试验', '审查', '生成')):
        return 'test_scope'
    if any(token in text for token in ('安装于', '使用场景', '飞行中', '工作环境', '应用于', '部署于')):
        return 'use_scene'
    if any(token in text for token in ('为替代', '为满足', '鉴定试验', '设计定型', '背景', '目的')):
        return 'background'
    if _looks_like_reference_line(text):
        return 'references'
    if _looks_like_parameter_line(text):
        return 'technical_params'
    return 'additional_context'


def _looks_like_parameter_line(text: str) -> bool:
    units = ('V', 'W', 'Hz', 'g', 'kg', 'mm', '℃', '%', 'kbps', '台')
    return ('：' in text or ':' in text) and any(unit in text for unit in units)


def _looks_like_reference_line(text: str) -> bool:
    import re

    return bool(re.search(r'(?<![A-Za-z0-9])(?:CCAR-\d+|DO-\d+|RTCA\s+DO-\d+|GJB\s*\d+|Q/[A-Za-z0-9-]+|企业标准)', text, flags=re.I))


def _merge_generation_brief_inputs(inputs: dict, parsed: dict, input_fields=None):
    if not parsed:
        return
    direct_fields = (
        'product_name',
        'project_name',
        'test_item',
        'technical_params',
        'references',
        'generation_requirements',
    )
    for key in direct_fields:
        if parsed.get(key) and not str(inputs.get(key) or '').strip():
            inputs[key] = parsed[key]
    context_parts = []
    for label, key in (
        ('试验范围', 'test_scope'),
        ('使用场景', 'use_scene'),
        ('背景说明', 'background'),
    ):
        value = parsed.get(key)
        if value:
            context_parts.append(f'{label}：{value}')
    if context_parts and not str(inputs.get('additional_context') or '').strip():
        inputs['additional_context'] = '\n'.join(context_parts)
    if parsed.get('background') and not str(inputs.get('background') or '').strip():
        inputs['background'] = parsed['background']
    material_context = str(parsed.get('parsed_material_context') or '').strip()
    if material_context and material_context not in str(inputs.get('parsed_material_context') or ''):
        existing_context = str(inputs.get('parsed_material_context') or '').strip()
        inputs['parsed_material_context'] = (
            f"{existing_context}\n{material_context}".strip()
            if existing_context
            else material_context
        )
    _merge_dynamic_fields_from_brief(inputs, parsed.get('_field_values') or {}, input_fields or [])


def _try_llm_parse_generation_brief(
    *,
    generation_brief: str,
    supplement_text: str,
    template,
    inputs: dict,
    model_option: dict,
) -> dict:
    if str(inputs.get('generation_mode') or 'smart') != 'smart':
        return {}
    if str(inputs.get('enable_llm_parse', 'true')).lower() in {'0', 'false', 'no', 'off'}:
        return {}
    material = "\n\n".join(part for part in (generation_brief, supplement_text) if str(part or '').strip())
    if not material.strip():
        return {}
    try:
        from config.settings import settings
        from llm.client import LLMClientFactory

        provider = model_option.get('provider') or settings.llm_provider
        client_kwargs = {'model': model_option.get('model') or None}
        if model_option.get('base_url'):
            client_kwargs['base_url'] = model_option.get('base_url')
        if model_option.get('api_key'):
            client_kwargs['api_key'] = model_option.get('api_key')
        llm_client = LLMClientFactory.create_client(provider, **client_kwargs)
        prompt = _build_llm_material_parse_prompt(material, template)
        response = llm_client.generate(prompt, system_prompt=(
            "你是技术文档素材结构化助手。只依据用户原文抽取信息，禁止编造。"
            "必须输出一个 JSON 对象，不要输出 Markdown。"
        ))
        parsed = _parse_llm_material_json(getattr(response, 'content', '') or '', template.input_fields)
        if parsed:
            current_app.logger.info("生成素材大模型解析完成，字段数=%s", len(parsed))
        return parsed
    except Exception as exc:
        current_app.logger.warning("生成素材大模型解析失败，使用本地解析结果: %s", exc)
        return {}


def _build_llm_material_parse_prompt(material: str, template) -> str:
    common_fields = [
        {"key": "title", "label": "文档标题"},
        {"key": "product_name", "label": "产品名称"},
        {"key": "project_name", "label": "项目名称"},
        {"key": "test_item", "label": "试验项目/受试对象"},
        {"key": "test_scope", "label": "试验范围"},
        {"key": "use_scene", "label": "使用场景"},
        {"key": "background", "label": "背景说明"},
        {"key": "technical_params", "label": "关键参数/产品信息"},
        {"key": "additional_context", "label": "补充信息"},
        {"key": "references", "label": "引用文件/依据文件"},
        {"key": "generation_requirements", "label": "生成要求"},
    ]
    dynamic_fields = [
        {
            "key": str(getattr(field, "key", "") or ""),
            "label": str(getattr(field, "label", "") or ""),
            "placeholder_tokens": list(getattr(field, "placeholder_tokens", []) or []),
        }
        for field in getattr(template, "input_fields", []) or []
        if str(getattr(field, "key", "") or "").strip()
    ]
    chapters = [
        f"{getattr(chapter, 'number', '')} {getattr(chapter, 'title', '')}".strip()
        for chapter in _flatten_template_chapters_for_prompt(getattr(template, "chapters", []) or [])
    ]
    payload = {
        "template": {
            "name": getattr(template, "name", ""),
            "doc_type": str(getattr(getattr(template, "doc_type", None), "value", "") or ""),
            "common_fields": common_fields,
            "dynamic_fields": dynamic_fields,
            "chapter_headings": chapters[:80],
        },
        "output_schema": {
            "common_fields": {
                "title": "string",
                "product_name": "string",
                "project_name": "string",
                "test_item": "string",
                "test_scope": "string",
                "use_scene": "string",
                "background": "string",
                "technical_params": "string",
                "additional_context": "string",
                "references": "string",
                "generation_requirements": "string"
            },
            "dynamic_fields": {"field_key": "string"},
            "material_sections": [
                {"label": "string", "content": "string"}
            ],
        },
        "rules": [
            "只抽取用户原文中明确出现或可直接归类的信息。",
            "不要补充用户没有提供的型号、编号、参数、结论或标准条款。",
            "优先匹配 template.dynamic_fields；common_fields 只是候选字段，只有确实适合当前文档时才填写。",
            "如果用户素材不适合任何 common_fields 或 dynamic_fields，放入 material_sections，不要强行归类。",
            "无法判断的字段留空字符串或空对象。",
            "references 保留每个引用文件一行；technical_params 保留每个参数一行。",
            "dynamic_fields 只能使用 template.dynamic_fields 中出现的 key。",
        ],
        "user_material": _truncate_for_llm_parse(material),
    }
    return json.dumps(payload, ensure_ascii=False, indent=2)


def _flatten_template_chapters_for_prompt(chapters) -> list:
    flat = []
    for chapter in chapters:
        flat.append(chapter)
        flat.extend(_flatten_template_chapters_for_prompt(getattr(chapter, "sub_chapters", []) or []))
    return flat


def _truncate_for_llm_parse(text: str, max_chars: int = 24000) -> str:
    value = str(text or "").strip()
    if len(value) <= max_chars:
        return value
    head = value[: int(max_chars * 0.7)]
    tail = value[-int(max_chars * 0.3):]
    return f"{head}\n\n……中间内容因长度被省略，以下为末尾内容……\n\n{tail}"


def _parse_llm_material_json(content: str, input_fields) -> dict:
    import re

    raw = str(content or "").strip()
    if raw.startswith("```"):
        raw = re.sub(r"^```(?:json)?\s*", "", raw)
        raw = re.sub(r"\s*```$", "", raw)
    match = re.search(r"\{.*\}", raw, flags=re.S)
    if match:
        raw = match.group(0)
    try:
        data = json.loads(raw)
    except json.JSONDecodeError:
        return {}
    if not isinstance(data, dict):
        return {}
    allowed = {
        "title",
        "product_name",
        "project_name",
        "test_item",
        "test_scope",
        "use_scene",
        "background",
        "technical_params",
        "additional_context",
        "references",
        "generation_requirements",
    }
    common_data = data.get("common_fields") if isinstance(data.get("common_fields"), dict) else data
    parsed = {
        key: str(common_data.get(key) or "").strip()
        for key in allowed
        if str(common_data.get(key) or "").strip()
    }
    allowed_dynamic = {str(getattr(field, "key", "") or "") for field in input_fields or []}
    dynamic = data.get("dynamic_fields") if isinstance(data.get("dynamic_fields"), dict) else {}
    field_values = {}
    for key, value in dynamic.items():
        key = str(key or "").strip()
        value = str(value or "").strip()
        if key and key in allowed_dynamic and value:
            field_values[key] = value
    if field_values:
        parsed["_field_values"] = field_values
    material_context = _format_llm_material_sections(data.get("material_sections"))
    if material_context:
        parsed["parsed_material_context"] = material_context
    return parsed


def _format_llm_material_sections(sections) -> str:
    if not isinstance(sections, list):
        return ""
    lines = []
    for item in sections[:20]:
        if not isinstance(item, dict):
            continue
        label = str(item.get("label") or "").strip()[:80]
        content = str(item.get("content") or "").strip()
        if label and content:
            lines.append(f"{label}：{content}")
    return "\n".join(lines)


def _merge_dynamic_fields_from_brief(inputs: dict, field_values: dict, input_fields) -> None:
    if not field_values or not input_fields:
        return
    dynamic_values = inputs.setdefault('dynamic_fields', {})
    normalized_values = {
        _normalize_brief_label(label): value
        for label, value in field_values.items()
        if value
    }
    for field in input_fields:
        key = str(getattr(field, 'key', '') or '').strip()
        if not key or str(dynamic_values.get(key, '') or '').strip():
            continue
        aliases = {
            str(getattr(field, 'label', '') or ''),
            key,
            str(getattr(field, 'placeholder', '') or ''),
        }
        aliases.update(str(token or '') for token in (getattr(field, 'placeholder_tokens', []) or []))
        for alias in aliases:
            value = normalized_values.get(_normalize_brief_label(alias))
            if value:
                dynamic_values[key] = value
                break


def _normalize_brief_label(label: str) -> str:
    import re

    return re.sub(r'[\s：:（）()《》“”"\'_-]+', '', str(label or '')).lower()


def _infer_references_from_brief(brief: str) -> str:
    import re

    candidates = []
    text = str(brief or '')
    pattern = r'(?<![A-Za-z0-9])(?:CCAR-\d+[A-Z0-9.-]*|DO-\d+[A-Z0-9.-]*|RTCA\s+DO-\d+[A-Z0-9.-]*|GJB\s*\d+(?:\.\d+)?[A-Z]?(?:-\d{4})?|Q/[A-Za-z0-9-]+)'
    for code in re.findall(pattern, text, flags=re.I):
        normalized = re.sub(r'\s+', ' ', code).strip()
        if normalized.upper().startswith('DO-'):
            normalized = f'RTCA {normalized}'
        if normalized not in candidates:
            candidates.append(normalized)
    names = {
        'CCAR-25': '运输类飞机适航标准',
        'RTCA DO-160G': '机载设备环境条件和试验程序',
    }
    rows = []
    for code in candidates:
        lookup_key = code.upper()
        name = next((value for prefix, value in names.items() if lookup_key.startswith(prefix)), '待补充')
        rows.append(f'{code}|{name}|待补充')
    return '\n'.join(rows)


def _has_source_docx(template: dict) -> bool:
    metadata = template.get("metadata") or {}
    return bool(metadata.get("source_docx_path") or metadata.get("source_path"))


def _wait_if_generation_paused(db, task_id: int):
    """Block the worker while paused and stop it cooperatively when canceled."""
    while True:
        task = db.get_generate_task(task_id)
        if not task:
            raise GenerationTaskCanceled("生成任务已删除")
        status = task.get('status')
        if status == 'canceled':
            raise GenerationTaskCanceled("生成任务已停止")
        if status != 'paused':
            return
        time.sleep(1)


def _extract_docx_reference_text(path: Path, max_chars: int = 12000) -> str:
    from docx import Document

    doc = Document(str(path))
    parts = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table_index, table in enumerate(doc.tables, start=1):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            if any(cells):
                rows.append(' | '.join(cells))
        if rows:
            parts.append(f"表格{table_index}：\n" + "\n".join(rows))
    text = "\n".join(parts)
    if len(text) > max_chars:
        text = text[:max_chars] + "\n...（补充材料内容较长，已截取前12000字）"
    return text


def _decorate_generate_task(task: dict) -> dict:
    labels = {
        'pending': ('等待中', 'warning'),
        'processing': ('生成中', 'info'),
        'paused': ('已暂停', 'warning'),
        'generated': ('已生成', 'success'),
        'completed': ('已完成', 'success'),
        'failed': ('失败', 'danger'),
        'review_failed': ('审查未通过', 'danger'),
        'canceled': ('已停止', 'default'),
    }
    label, badge_class = labels.get(task.get('status'), (task.get('status') or '未知', 'default'))
    decorated = dict(task)
    decorated['status_label'] = label
    decorated['badge_class'] = badge_class
    return decorated
