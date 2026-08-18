"""Generation template library routes."""
import json
import os
import uuid
from pathlib import Path

from flask import Blueprint, current_app, jsonify, render_template, request

from templates.docx_template_parser import DocxTemplateParser
from templates.template_manager import TemplateManager
from web.option_registry import build_reference_case_context, get_specialty

bp = Blueprint("template_library", __name__)

DOCUMENT_KINDS_FILE = Path(__file__).parent.parent.parent / "config" / "document_kinds.json"
DEFAULT_DOCUMENT_KINDS = [
    "试验大纲",
    "可靠性分配报告",
    "技术说明书",
    "产品规范",
    "验证报告",
    "需求文档",
    "通用特性文档",
]


def _normalize_document_kind(value: str) -> str:
    return str(value or "").strip()


def _load_custom_document_kinds() -> list[str]:
    if not DOCUMENT_KINDS_FILE.exists():
        return []
    try:
        with open(DOCUMENT_KINDS_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
    except (json.JSONDecodeError, OSError):
        return []
    items = data.get("document_kinds", data if isinstance(data, list) else [])
    kinds = []
    for item in items:
        name = _normalize_document_kind(item.get("name") if isinstance(item, dict) else item)
        if name and name not in kinds:
            kinds.append(name)
    return kinds


def _save_custom_document_kinds(kinds: list[str]) -> None:
    clean = []
    for item in kinds:
        name = _normalize_document_kind(item)
        if name and name not in clean:
            clean.append(name)
    DOCUMENT_KINDS_FILE.parent.mkdir(parents=True, exist_ok=True)
    with open(DOCUMENT_KINDS_FILE, "w", encoding="utf-8") as f:
        json.dump({"document_kinds": clean}, f, ensure_ascii=False, indent=2)


def _template_document_kind_counts(manager: TemplateManager) -> dict[str, int]:
    counts = {}
    for template in manager.list_template_dicts():
        name = _normalize_document_kind((template.get("metadata") or {}).get("document_kind_name"))
        if not name:
            continue
        counts[name] = counts.get(name, 0) + 1
    return counts


def _serialize_document_kinds(manager: TemplateManager) -> list[dict]:
    custom_kinds = _load_custom_document_kinds()
    template_counts = _template_document_kind_counts(manager)
    ordered = []
    for name in DEFAULT_DOCUMENT_KINDS + custom_kinds + list(template_counts.keys()):
        clean = _normalize_document_kind(name)
        if clean and clean not in ordered:
            ordered.append(clean)
    return [
        {
            "name": name,
            "template_count": template_counts.get(name, 0),
            "is_default": name in DEFAULT_DOCUMENT_KINDS,
            "is_custom": name in custom_kinds,
            "can_delete": name not in DEFAULT_DOCUMENT_KINDS and template_counts.get(name, 0) == 0,
        }
        for name in ordered
    ]


def _validate_input_fields(items):
    seen = set()
    for index, item in enumerate(items or [], start=1):
        if not isinstance(item, dict):
            return f"第 {index} 个输入字段格式不正确"
        key = str(item.get("key") or "").strip()
        label = str(item.get("label") or "").strip()
        field_type = str(item.get("type") or item.get("field_type") or "text")
        if not key or not label:
            return f"第 {index} 个输入字段的字段键和名称不能为空"
        if key in seen:
            return f"输入字段键不能重复：{key}"
        if field_type not in {"text", "textarea", "number", "date", "select"}:
            return f"输入字段类型不支持：{field_type}"
        if field_type == "select" and not [
            value for value in item.get("options", []) if str(value).strip()
        ]:
            return f"下拉字段“{label}”至少需要一个选项"
        seen.add(key)
    return ""


@bp.route("/")
def index():
    return render_template("template_library.html", active_page="template_library")


@bp.route("/api/templates")
def api_templates():
    manager = TemplateManager()
    return jsonify({"templates": manager.list_template_dicts()})


@bp.route("/api/document-kinds")
def api_document_kinds():
    manager = TemplateManager()
    return jsonify({"document_kinds": _serialize_document_kinds(manager)})


@bp.route("/api/document-kinds", methods=["POST"])
def api_create_document_kind():
    data = request.get_json() or {}
    name = _normalize_document_kind(data.get("name"))
    if not name:
        return jsonify({"error": "请填写技术文档类型名称"}), 400
    manager = TemplateManager()
    existing = [item["name"] for item in _serialize_document_kinds(manager)]
    if name in existing:
        return jsonify({"error": "该技术文档类型已存在"}), 400
    custom_kinds = _load_custom_document_kinds()
    custom_kinds.append(name)
    _save_custom_document_kinds(custom_kinds)
    return jsonify({"ok": True, "document_kinds": _serialize_document_kinds(manager)})


@bp.route("/api/document-kinds/rename", methods=["POST"])
def api_rename_document_kind():
    data = request.get_json() or {}
    old_name = _normalize_document_kind(data.get("old_name"))
    new_name = _normalize_document_kind(data.get("new_name"))
    if not old_name or not new_name:
        return jsonify({"error": "请填写原类型和新类型名称"}), 400
    if old_name == new_name:
        return jsonify({"ok": True, "document_kinds": _serialize_document_kinds(TemplateManager())})

    manager = TemplateManager()
    existing = [item["name"] for item in _serialize_document_kinds(manager)]
    if old_name not in existing:
        return jsonify({"error": "原技术文档类型不存在"}), 404
    if new_name in existing:
        return jsonify({"error": "新技术文档类型已存在"}), 400

    custom_kinds = _load_custom_document_kinds()
    if old_name in custom_kinds:
        custom_kinds = [new_name if item == old_name else item for item in custom_kinds]
    elif old_name not in DEFAULT_DOCUMENT_KINDS:
        custom_kinds.append(new_name)
    else:
        return jsonify({"error": "默认技术文档类型暂不支持重命名"}), 400
    _save_custom_document_kinds(custom_kinds)

    for template in manager.list_template_dicts():
        metadata = dict(template.get("metadata") or {})
        if _normalize_document_kind(metadata.get("document_kind_name")) != old_name:
            continue
        metadata["document_kind_name"] = new_name
        manager.update_template(template["id"], {"metadata": metadata})

    refreshed = TemplateManager()
    return jsonify({"ok": True, "document_kinds": _serialize_document_kinds(refreshed)})


@bp.route("/api/document-kinds/delete", methods=["POST"])
def api_delete_document_kind():
    data = request.get_json() or {}
    name = _normalize_document_kind(data.get("name"))
    if not name:
        return jsonify({"error": "请填写要删除的技术文档类型"}), 400
    manager = TemplateManager()
    counts = _template_document_kind_counts(manager)
    if name in DEFAULT_DOCUMENT_KINDS:
        return jsonify({"error": "默认技术文档类型不可删除"}), 400
    if counts.get(name, 0):
        return jsonify({"error": f"该类型已有 {counts[name]} 个模板使用，请先调整模板类型后再删除"}), 400
    custom_kinds = [item for item in _load_custom_document_kinds() if item != name]
    _save_custom_document_kinds(custom_kinds)
    return jsonify({"ok": True, "document_kinds": _serialize_document_kinds(manager)})


@bp.route("/api/templates/<template_id>")
def api_template_detail(template_id: str):
    manager = TemplateManager()
    template = manager.get_template(template_id)
    if not template:
        return jsonify({"error": f"模板不存在: {template_id}"}), 404
    return jsonify(manager.serialize_template(template))


@bp.route("/api/parse", methods=["POST"])
def api_parse_template():
    file = request.files.get("file")
    if not file:
        return jsonify({"error": "请上传 DOCX 模板文件"}), 400
    if not file.filename.lower().endswith(".docx"):
        return jsonify({"error": "仅支持 DOCX 模板文件"}), 400

    template_id = f"template_{uuid.uuid4().hex[:8]}"
    template_dir = Path(current_app.config["UPLOAD_FOLDER"]) / "templates" / template_id
    template_dir.mkdir(parents=True, exist_ok=True)
    path = template_dir / "source.docx"
    file.save(path)

    parsed = DocxTemplateParser().parse(str(path))
    specialty = get_specialty(request.form.get("specialty_id"))
    parsed["id"] = template_id
    parsed["metadata"] = {
        "template_asset_id": template_id,
        "source_filename": file.filename,
        "source_path": os.path.relpath(path, Path(current_app.root_path).parent),
        "source_docx_path": os.path.relpath(path, Path(current_app.root_path).parent),
    }
    if specialty:
        parsed["metadata"]["specialty_id"] = specialty.get("id")
        parsed["metadata"]["specialty_name"] = specialty.get("name")
    return jsonify(parsed)


def _validate_template_asset_metadata(metadata: dict) -> str:
    specialty_id = str((metadata or {}).get("specialty_id") or "").strip()
    document_kind_name = str((metadata or {}).get("document_kind_name") or "").strip()
    if not specialty_id or not get_specialty(specialty_id):
        return "请选择模板所属专业"
    if not document_kind_name:
        return "请填写技术文档类型"
    return ""


@bp.route("/api/templates", methods=["POST"])
def api_create_template():
    data = request.get_json() or {}
    name = str(data.get("name", "")).strip()
    if not name:
        return jsonify({"error": "模板名称不能为空"}), 400
    chapters = data.get("chapters") or []
    if not chapters:
        return jsonify({"error": "模板章节不能为空"}), 400
    field_error = _validate_input_fields(data.get("input_fields") or [])
    if field_error:
        return jsonify({"error": field_error}), 400
    metadata = data.get("metadata") or {}
    metadata_error = _validate_template_asset_metadata(metadata)
    if metadata_error:
        return jsonify({"error": metadata_error}), 400

    manager = TemplateManager()
    template = manager.create_template(
        name=name,
        description=data.get("description", ""),
        chapters=chapters,
        metadata=metadata,
        source_type=data.get("source_type", "uploaded_docx"),
        template_id=data.get("id") or None,
        input_fields=data.get("input_fields") or [],
    )
    return jsonify({"ok": True, "template": manager.serialize_template(template)})


@bp.route("/api/templates/<template_id>", methods=["PUT"])
def api_update_template(template_id: str):
    data = request.get_json() or {}
    field_error = _validate_input_fields(data.get("input_fields") or [])
    if field_error:
        return jsonify({"error": field_error}), 400
    metadata_error = _validate_template_asset_metadata(data.get("metadata") or {})
    if metadata_error:
        return jsonify({"error": metadata_error}), 400
    manager = TemplateManager()
    template = manager.update_template(template_id, data)
    if not template:
        return jsonify({"error": "模板不存在或内置模板不可编辑"}), 404
    return jsonify({"ok": True, "template": manager.serialize_template(template)})


@bp.route("/api/templates/<template_id>/source", methods=["POST"])
def api_replace_template_source(template_id: str):
    file = request.files.get("file")
    if not file:
        return jsonify({"error": "请上传新的 DOCX 模板文件"}), 400
    if not file.filename.lower().endswith(".docx"):
        return jsonify({"error": "仅支持 DOCX 模板文件"}), 400

    manager = TemplateManager()
    template = manager.get_template(template_id)
    if not template or template.source_type == "built_in":
        return jsonify({"error": "模板不存在或内置模板不可替换源文件"}), 404

    template_dir = Path(current_app.config["UPLOAD_FOLDER"]) / "templates" / template_id
    template_dir.mkdir(parents=True, exist_ok=True)
    path = template_dir / "source.docx"
    file.save(path)

    parsed = DocxTemplateParser().parse(str(path))
    metadata = {
        "template_asset_id": template_id,
        "source_filename": file.filename,
        "source_path": os.path.relpath(path, Path(current_app.root_path).parent),
        "source_docx_path": os.path.relpath(path, Path(current_app.root_path).parent),
        "specialty_id": (template.metadata or {}).get("specialty_id"),
        "specialty_name": (template.metadata or {}).get("specialty_name"),
        "document_kind_name": (template.metadata or {}).get("document_kind_name"),
    }
    updated = manager.replace_template_source(template_id, parsed, metadata)
    if not updated:
        return jsonify({"error": "模板不存在或不可替换源文件"}), 404
    return jsonify({"ok": True, "template": manager.serialize_template(updated)})


@bp.route("/api/templates/<template_id>/reference-cases", methods=["POST"])
def api_upload_reference_case(template_id: str):
    file = request.files.get("file")
    if not file:
        return jsonify({"error": "请上传优秀案例 DOCX 文件"}), 400
    if not file.filename.lower().endswith(".docx"):
        return jsonify({"error": "仅支持 DOCX 案例文件"}), 400

    manager = TemplateManager()
    template = manager.get_template(template_id)
    if not template or template.source_type == "built_in":
        return jsonify({"error": "请先将模板复制或保存为用户模板，再维护优秀案例"}), 400

    case_id = f"case_{uuid.uuid4().hex[:8]}"
    case_dir = Path(current_app.config["UPLOAD_FOLDER"]) / "templates" / template_id / "reference_cases"
    case_dir.mkdir(parents=True, exist_ok=True)
    path = case_dir / f"{case_id}.docx"
    file.save(path)

    relative_path = os.path.relpath(path, Path(current_app.root_path).parent)
    case = {
        "id": case_id,
        "name": str(request.form.get("name") or file.filename.rsplit(".", 1)[0]).strip(),
        "doc_type": str(request.form.get("doc_type") or template.metadata.get("document_kind_name") or "").strip(),
        "scenario": str(request.form.get("scenario") or "").strip(),
        "file_name": file.filename,
        "file_path": relative_path,
        "features": _extract_reference_case_features(path),
    }

    metadata = dict(template.metadata or {})
    cases = [item for item in metadata.get("reference_cases", []) if isinstance(item, dict)]
    cases.append(case)
    metadata["reference_cases"] = cases
    updated = manager.update_template(template_id, {"metadata": metadata})
    if not updated:
        return jsonify({"error": "保存优秀案例失败"}), 400
    return jsonify({"ok": True, "case": case, "template": manager.serialize_template(updated)})


@bp.route("/api/templates/<template_id>/reference-cases/<case_id>", methods=["DELETE"])
def api_delete_reference_case(template_id: str, case_id: str):
    manager = TemplateManager()
    template = manager.get_template(template_id)
    if not template or template.source_type == "built_in":
        return jsonify({"error": "模板不存在或不可编辑"}), 404
    metadata = dict(template.metadata or {})
    cases = [item for item in metadata.get("reference_cases", []) if isinstance(item, dict)]
    target = next((item for item in cases if item.get("id") == case_id), None)
    metadata["reference_cases"] = [item for item in cases if item.get("id") != case_id]
    updated = manager.update_template(template_id, {"metadata": metadata})
    if target and target.get("file_path"):
        try:
            (Path(current_app.root_path).parent / target["file_path"]).unlink(missing_ok=True)
        except OSError:
            pass
    return jsonify({"ok": True, "template": manager.serialize_template(updated)})


@bp.route("/api/templates/<template_id>", methods=["DELETE"])
def api_delete_template(template_id: str):
    manager = TemplateManager()
    if not manager.delete_template(template_id):
        return jsonify({"error": "模板不存在"}), 404
    return jsonify({"ok": True})


def _extract_reference_case_features(path: Path) -> dict:
    text, headings, tables = _extract_docx_outline(path)
    terms = _pick_terms(text)
    return {
        "chapter_structure": headings[:18],
        "terminology": terms[:18],
        "writing_style": _infer_writing_style(text),
        "format_features": _infer_format_features(headings, tables),
        "context_preview": build_reference_case_context([{
            "name": path.stem,
            "features": {
                "chapter_structure": headings[:12],
                "terminology": terms[:12],
                "writing_style": _infer_writing_style(text),
                "format_features": _infer_format_features(headings, tables),
            },
            "file_path": str(path),
        }])[:3000],
    }


def _extract_docx_outline(path: Path) -> tuple[str, list[str], int]:
    from docx import Document

    doc = Document(str(path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    headings = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = getattr(paragraph.style, "name", "") or ""
        if style_name.lower().startswith("heading") or style_name.startswith("标题"):
            headings.append(text)
            continue
        if len(text) <= 40 and any(text.startswith(prefix) for prefix in ("1", "2", "3", "4", "5", "一", "二", "三")):
            headings.append(text)
    return "\n".join(paragraphs[:120]), headings, len(doc.tables)


def _pick_terms(text: str) -> list[str]:
    candidates = [
        "作动器", "伺服阀", "液压源", "试验件", "位移", "压力", "载荷", "响应时间",
        "控制指令", "验收准则", "判据", "环境试验", "可靠性", "接口", "性能指标",
    ]
    return [term for term in candidates if term in text]


def _infer_writing_style(text: str) -> list[str]:
    styles = []
    if "应" in text:
        styles.append("技术要求多采用“应……”的强制性表述")
    if "试验" in text:
        styles.append("围绕试验条件、步骤、记录和判据组织内容")
    if "表" in text:
        styles.append("使用表格归纳指标、条件或记录项")
    return styles or ["保持正式工程技术文档语气，按模板结构展开"]


def _infer_format_features(headings: list[str], table_count: int) -> list[str]:
    features = []
    if headings:
        features.append("章节标题层级清晰，生成时应保持同类编号结构")
    if table_count:
        features.append(f"包含 {table_count} 个表格，可参考其指标/步骤/判据组织方式")
    return features or ["按模板原有标题、段落和表格格式输出"]
